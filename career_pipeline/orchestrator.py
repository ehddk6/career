from datetime import datetime
"""전체 파이프라인 조정. prepare와 finalize로 나뉘며, profile/posting/matching/research/finalize 흐름을 제어합니다."""
from dataclasses import asdict, replace
from hashlib import sha256
import json
from pathlib import Path
import re

import yaml
from docx import Document

from .conflicts import (
    apply_overrides,
    conflict_override_key,
    detect_conflicts,
)
from .extractors import extract_path
from .facts import METRIC, _normalize, extract_fact_claims
from .inventory import build_inventory
from .candidate_selection import generate_and_select_candidates
from .artifacts import write_final_artifact_manifest
from .character_count import count_characters
from .matching import match_questions, render_matches_markdown
from .posting_loader import (
    PostingSourceError,
    load_posting_source,
    write_posting_snapshot,
)
from .copyeditor_adapter import copyedit_responses
from .contract_builder import refresh_run_interview_contract
from .cost_limit import CostLimitExceeded, CostTracker
from .model_policy import ModelTier, choose_tier, resolve_model, resolve_role_model
from .rigorous_selection import (
    RigorousSelectionError,
    run_rigorous_selection,
    subprocess_model_runner,
)
from .posting_parser import parse_posting, reconcile_questions, render_posting_analysis
from .profile_refresh import refresh_profile
from .profile_schema import (
    ExperienceLedger,
    ProfileValidationError,
    ledger_to_dict,
    load_ledger,
)
from .prompt_contracts import (
    INTERVIEW_CONTRACT_NAME,
    prompt_contract_context,
    validate_run_prompt_contracts,
)
from .quality_profiles import get_quality_profile, legacy_rigorous_profile
from .question_requirements import (
    build_question_requirement_map,
    validate_question_requirement_map,
)
from .quality import (
    STRICT_MIN_ANSWER_SCORE,
    STRICT_MIN_AVERAGE_SCORE,
    QualityIssue,
    score_answer_quality,
    validate_answer_quality,
    validate_interview_pack,
    validate_matching_gate,
    validate_posting_gate,
    validate_profile_gate,
)
from .questions import extract_questions
from .models import DraftResponse, ExperienceClaimRef, Question, ValidationIssue
from .patina_adapter import (
    HumanizationResult,
    PatinaScoreResult,
    humanize_text,
    score_text,
)
from .research_evidence import (
    DEFAULT_RESEARCH_METHOD,
    REQUIRED_RESEARCH_POLICY,
    load_research_execution,
    load_research_claims,
    official_domains_for_target,
    validate_research_execution,
    validate_research_evidence,
)
from .rendering import render_draft_docx, render_draft_markdown
from .state import resolve_run_dir, write_json, write_state
from .style_diagnostics import diagnose_responses
from .source_policy import is_evidence_path
from .validation import referenced_claim_values, validate_draft
from .writing_guidance import attach_writing_guidance


def _inventory_markdown(records) -> str:
    lines = ["# 자료 목록", "", "| 상태 | 파일 | 사유 |", "|---|---|---|"]
    lines.extend(
        f"| {item.status} | {item.relative_path} | {item.reason} |"
        for item in records
    )
    return "\n".join(lines) + "\n"


def _conflict_markdown(conflicts, claims) -> str:
    lines = ["# 충돌 검사", ""]
    if not conflicts:
        lines.append("- 충돌 없음")
    for number, conflict in enumerate(conflicts, 1):
        lines.extend(
            [
                f"## 충돌 {number}: {conflict.field}",
                f"값: {', '.join(conflict.values)}",
            ]
        )
        for index in conflict.claim_indexes:
            claim = claims[index]
            lines.append(f"- `{claim.source_path}`: {claim.context}")
        lines.append(
            f"- override key: `{conflict_override_key(conflict, claims)}`"
        )
        lines.append(
            "- 확인 질문: 실제 제출에 사용할 값과 근거 파일을 지정해 주세요."
        )
    return "\n".join(lines) + "\n"


def _load_overrides(path: Path) -> dict[str, str]:
    if not path.exists():
        return {}
    payload = yaml.safe_load(path.read_text(encoding="utf-8")) or {}
    if not isinstance(payload, dict) or not all(
        isinstance(key, str) and isinstance(value, str)
        for key, value in payload.items()
    ):
        raise ValueError("fact_overrides.yaml must be a string-to-string mapping")
    return payload


def _load_draft_responses(path: Path) -> tuple[list[DraftResponse], list[ValidationIssue]]:
    """Load the externally authored draft contract without leaking parser errors."""
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError, UnicodeDecodeError) as error:
        return [], [
            ValidationIssue("invalid_draft_json", 0, f"draft.json을 읽을 수 없습니다: {error}")
        ]
    if not isinstance(payload, list):
        return [], [
            ValidationIssue("invalid_draft_shape", 0, "draft.json 최상위 값은 배열이어야 합니다.")
        ]

    responses: list[DraftResponse] = []
    issues: list[ValidationIssue] = []
    for position, item in enumerate(payload, start=1):
        if not isinstance(item, dict):
            issues.append(
                ValidationIssue("invalid_draft_entry", 0, f"{position}번 항목은 객체여야 합니다.")
            )
            continue
        question_index = item.get("question_index")
        answer = item.get("answer")
        evidence_paths = item.get("evidence_paths", [])
        experience_refs = item.get("experience_refs", [])
        research_refs = item.get("research_refs", [])
        if isinstance(question_index, bool) or not isinstance(question_index, int):
            issues.append(
                ValidationIssue("invalid_question_index", 0, f"{position}번 항목의 question_index는 정수여야 합니다.")
            )
            continue
        if not isinstance(answer, str):
            issues.append(
                ValidationIssue("invalid_answer", question_index, "answer는 문자열이어야 합니다.")
            )
            continue
        if not isinstance(evidence_paths, list) or not all(
            isinstance(value, str) for value in evidence_paths
        ):
            issues.append(
                ValidationIssue("invalid_evidence_paths", question_index, "evidence_paths는 문자열 배열이어야 합니다.")
            )
            continue
        if not isinstance(research_refs, list) or not all(
            isinstance(value, str) for value in research_refs
        ):
            issues.append(
                ValidationIssue("invalid_research_refs", question_index, "research_refs는 문자열 배열이어야 합니다.")
            )
            continue
        if not isinstance(experience_refs, list):
            issues.append(
                ValidationIssue("invalid_experience_refs", question_index, "experience_refs는 배열이어야 합니다.")
            )
            continue
        parsed_refs: list[ExperienceClaimRef] = []
        for reference in experience_refs:
            if not isinstance(reference, dict):
                issues.append(
                    ValidationIssue("invalid_experience_ref", question_index, "experience_refs 항목은 객체여야 합니다.")
                )
                break
            experience_id = reference.get("experience_id")
            claim_fields = reference.get("claim_fields", [])
            claim_ids = reference.get("claim_ids", [])
            if (
                not isinstance(experience_id, str)
                or not isinstance(claim_fields, list)
                or not all(isinstance(field, str) for field in claim_fields)
                or not isinstance(claim_ids, list)
                or not all(isinstance(claim_id, str) for claim_id in claim_ids)
            ):
                issues.append(
                    ValidationIssue("invalid_experience_ref", question_index, "experience_refs 항목 형식이 올바르지 않습니다.")
                )
                break
            parsed_refs.append(
                ExperienceClaimRef(
                    experience_id,
                    tuple(claim_fields),
                    tuple(claim_ids),
                )
            )
        else:
            responses.append(
                DraftResponse(
                    question_index,
                    answer,
                    tuple(evidence_paths),
                    tuple(parsed_refs),
                    tuple(research_refs),
                )
            )
    return responses, issues


def _hydrate_claim_evidence_paths(
    responses: list[DraftResponse], ledger: ExperienceLedger
) -> None:
    """Reconcile evidence paths from exact claim IDs without trusting model paths.

    In the rigorous contract, claim IDs are the authoritative reference.  A
    model may omit ``evidence_paths`` or return a display path that is not the
    ledger's canonical source path.  Keeping that path would create a false
    ``unknown_evidence`` failure even though the claim itself is exact.  When
    valid claim IDs are present, rebuild the path list solely from their
    verified evidence; preserve model paths only when no claim can supply one
    so an ungrounded response still fails closed.
    """
    by_experience = {item.experience_id: item for item in ledger.experiences}
    hydrated: list[DraftResponse] = []
    for response in responses:
        claim_paths: set[str] = set()
        for reference in response.experience_refs:
            experience = by_experience.get(reference.experience_id)
            if experience is None:
                continue
            claims = {claim.claim_id: claim for claim in experience.claims}
            for claim_id in reference.claim_ids:
                claim = claims.get(claim_id)
                if claim is None:
                    continue
                claim_paths.update(evidence.source_path for evidence in claim.evidence)
        if not response.experience_refs and response.research_refs:
            # Research evidence is tracked by research_refs.  Do not let a
            # model copy official URLs into the personal evidence channel.
            paths: set[str] = set()
        else:
            paths = claim_paths or set(response.evidence_paths)
        hydrated.append(replace(response, evidence_paths=tuple(sorted(paths))))
    responses[:] = hydrated


def _replace_interview_contract_claims(
    run_dir: Path, responses: list[DraftResponse]
) -> None:
    """Replace the JSON claim audit with the exact current response references."""
    contract_path = run_dir / INTERVIEW_CONTRACT_NAME
    if contract_path.is_file():
        payload = json.loads(contract_path.read_text(encoding="utf-8"))
        if isinstance(payload, dict):
            payload["submitted_claims"] = [
                {
                    "question_index": response.question_index,
                    "experience_ids": sorted(
                        {
                            reference.experience_id
                            for reference in response.experience_refs
                            if reference.experience_id
                        }
                    ),
                    "experience_claim_ids": sorted(
                        {
                            claim_id
                            for reference in response.experience_refs
                            for claim_id in reference.claim_ids
                            if claim_id
                        }
                    ),
                    "research_claim_ids": sorted(set(response.research_refs)),
                    "status": "CONFIRMED",
                }
                for response in sorted(responses, key=lambda item: item.question_index)
            ]
            write_json(contract_path, payload)


def _link_final_claims_to_interview_pack(
    run_dir: Path, responses: list[DraftResponse]
) -> None:
    """Attach the selected draft's evidence IDs to JSON and Markdown interview blocks.

    The interview pack is prepared before rigorous selection, so its original
    evidence lines can refer to the incumbent rather than the selected
    candidate. Preserve the authored practice answers while replacing the JSON
    claim audit and adding deterministic Markdown linkage lines.
    """
    contract_path = run_dir / INTERVIEW_CONTRACT_NAME
    refresh_inputs = (
        "00_채용공고분석.json",
        "02_확정경험원장.json",
        "03_경험직무매칭.json",
        "04_공식근거.json",
    )
    if contract_path.is_file() and all((run_dir / name).is_file() for name in refresh_inputs):
        refresh_run_interview_contract(
            run_dir,
            [
                {
                    "question_index": response.question_index,
                    "answer": response.answer,
                    "evidence_paths": list(response.evidence_paths),
                    "experience_refs": [
                        {
                            "experience_id": reference.experience_id,
                            "claim_fields": list(reference.claim_fields),
                            "claim_ids": list(reference.claim_ids),
                        }
                        for reference in response.experience_refs
                    ],
                    "research_refs": list(response.research_refs),
                }
                for response in responses
            ],
        )
    else:
        _replace_interview_contract_claims(run_dir, responses)

    path = run_dir / "08_면접대비팩.md"
    if not path.exists():
        return
    text = path.read_text(encoding="utf-8")
    original_text = text
    marker_re = re.compile(r"(?m)^##\s*문항\s+(\d+)\s*$")
    markers = list(marker_re.finditer(text))
    if not markers:
        return
    by_index = {response.question_index: response for response in responses}
    rebuilt = [text[: markers[0].start()]]
    for offset, marker in enumerate(markers):
        index = int(marker.group(1))
        response = by_index.get(index)
        end = markers[offset + 1].start() if offset + 1 < len(markers) else len(text)
        block = text[marker.start():end]
        block = re.sub(r"(?m)^- 최종 제출본 근거 ID:.*\n?", "", block)
        if response is None:
            rebuilt.append(block)
            continue
        ids: list[str] = []
        for reference in response.experience_refs:
            ids.append(reference.experience_id)
            ids.extend(reference.claim_ids)
        ids.extend(response.research_refs)
        ids = list(dict.fromkeys(ids))
        if ids:
            line = "- 최종 제출본 근거 ID: " + ", ".join(ids)
            insertion = block.find("\n")
            insertion = len(block) if insertion < 0 else insertion + 1
            block = block[:insertion] + "\n" + line + "\n" + block[insertion:]
        rebuilt.append(block)
    text = "".join(rebuilt)
    if text != original_text:
        path.write_text(text, encoding="utf-8")


def _write_research_execution_template(run_dir: Path) -> None:
    path = run_dir / "04_리서치실행.json"
    if path.exists():
        return
    write_json(
        path,
        {
            "policy": REQUIRED_RESEARCH_POLICY,
            "skill_name": DEFAULT_RESEARCH_METHOD,
            "mode": "ordinary-online",
            "searched_at": "",
            "status": "pending",
            "queries": [],
            "source_families": [],
            "verified_claim_ids": [],
        },
    )


def _resolve_voice_sample(
    explicit: Path | None,
    state: dict,
    run_dir: Path,
) -> Path | None:
    candidates = [
        explicit,
        Path(state["patina_voice_sample"]) if state.get("patina_voice_sample") else None,
        Path(state["root"]) / ".career_profile" / "voice_sample.txt"
        if state.get("root")
        else None,
        run_dir / "voice_sample.txt",
    ]
    for candidate in candidates:
        if candidate is None:
            continue
        path = candidate.resolve()
        if path.is_file():
            if path.stat().st_size > 20_000:
                raise ValueError("Patina voice sample must be 20KB or smaller")
            paragraphs = [
                item.strip()
                for item in re.split(r"\r?\n\s*\r?\n", path.read_text(encoding="utf-8"))
                if item.strip()
            ]
            if not 1 <= len(paragraphs) <= 3:
                raise ValueError("Patina voice sample must contain 1-3 paragraphs")
            return path
        if explicit is not None and path == explicit.resolve():
            raise FileNotFoundError(f"Patina voice sample not found: {path}")
    return None


def _confirmed_ledger(ledger: ExperienceLedger) -> ExperienceLedger:
    experiences = tuple(
        replace(
            experience,
            claims=tuple(
                claim for claim in experience.claims if claim.status == "confirmed"
            ),
        )
        for experience in ledger.experiences
        if experience.status == "confirmed"
    )
    return replace(ledger, experiences=experiences)


def _blocked_v2_state(
    run_dir: Path,
    root: Path,
    target: str,
    draft: Path,
    posting: str | None,
    status: str,
    stage: str,
    issues: list[QualityIssue],
    questions: list[Question] | tuple[Question, ...] = (),
) -> dict:
    state = {
        "status": status,
        "quality_mode": "v2",
        "strict_quality": True,
        "blocked_stage": stage,
        "issues": [asdict(issue) for issue in issues],
        "run_dir": str(run_dir),
        "root": str(root),
        "target": target,
        "draft": str(draft),
        "posting": posting,
        "questions": [asdict(question) for question in questions],
        "conflict_count": sum(
            issue.code == "conflicting_profile_claim" for issue in issues
        ),
    }
    attach_writing_guidance(root, run_dir, state)
    write_state(run_dir, state)
    return state


def _prepare_v2(
    *,
    root: Path,
    target: str,
    draft: Path,
    posting: str | None,
    run_dir: Path,
    profile: Path,
    official_domains: tuple[str, ...],
    research_domains: tuple[str, ...],
    official_source: bool,
    questions: list[Question],
) -> dict:
    try:
        ledger = load_ledger(profile)
    except (OSError, ProfileValidationError) as error:
        return _blocked_v2_state(
            run_dir,
            root,
            target,
            draft,
            posting,
            "blocked_profile",
            "profile",
            [QualityIssue("invalid_profile", str(error), str(profile))],
            questions,
        )

    review = refresh_profile(root, ledger)
    review_issues = [
        QualityIssue(
            "stale_profile_evidence" if item.status == "stale" else "missing_profile_evidence",
            f"{item.experience_id}: {item.reason}",
            item.source_path,
        )
        for item in review.items
        if item.status != "unchanged"
    ]
    selected_ids = {
        item.experience_id
        for item in ledger.experiences
        if item.status == "confirmed"
    }
    profile_issues = validate_profile_gate(
        ledger, selected_experience_ids=selected_ids
    )
    all_profile_issues = review_issues + profile_issues
    if all_profile_issues:
        status = (
            "blocked_conflict"
            if any(issue.code == "conflicting_profile_claim" for issue in all_profile_issues)
            else "blocked_profile"
        )
        return _blocked_v2_state(
            run_dir,
            root,
            target,
            draft,
            posting,
            status,
            "profile",
            all_profile_issues,
            questions,
        )

    confirmed = _confirmed_ledger(ledger)
    write_json(run_dir / "02_확정경험원장.json", ledger_to_dict(confirmed))

    if not posting:
        return _blocked_v2_state(
            run_dir,
            root,
            target,
            draft,
            posting,
            "blocked_posting",
            "posting",
            [QualityIssue("missing_posting", "채용공고 입력이 필요합니다.")],
            questions,
        )
    try:
        loaded = load_posting_source(
            posting,
            official_source=official_source,
            official_domains=official_domains,
        )
        write_posting_snapshot(run_dir, loaded)
        analysis = parse_posting(loaded, target=target)
    except (OSError, PostingSourceError) as error:
        return _blocked_v2_state(
            run_dir,
            root,
            target,
            draft,
            posting,
            "blocked_posting",
            "posting",
            [QualityIssue("invalid_posting_source", str(error))],
            questions,
        )

    write_json(run_dir / "00_채용공고분석.json", asdict(analysis))
    (run_dir / "00_채용공고분석.md").write_text(
        render_posting_analysis(analysis), encoding="utf-8"
    )
    posting_issues = validate_posting_gate(analysis)
    reconciliation = reconcile_questions(analysis.questions, tuple(questions))
    posting_issues.extend(
        QualityIssue(
            f"question_{item.reason}",
            f"문항 {item.index}: 공고={item.posting_value!r}, 초안={item.draft_value!r}",
            "00_채용공고분석.md",
        )
        for item in reconciliation.mismatches
    )
    if posting_issues:
        return _blocked_v2_state(
            run_dir,
            root,
            target,
            draft,
            posting,
            "blocked_posting",
            "posting",
            posting_issues,
            questions,
        )

    matches = match_questions(confirmed, analysis, reconciliation.questions)
    write_json(
        run_dir / "03_경험직무매칭.json",
        [asdict(item) for item in matches],
    )
    (run_dir / "03_경험직무매칭.md").write_text(
        render_matches_markdown(matches), encoding="utf-8"
    )
    question_requirement_map = build_question_requirement_map(
        reconciliation.questions,
        target=target,
        posting=asdict(analysis),
        matches=matches,
    )
    write_json(run_dir / "05_문항전략.json", question_requirement_map)
    matching_issues = validate_matching_gate(matches)
    if matching_issues:
        return _blocked_v2_state(
            run_dir,
            root,
            target,
            draft,
            posting,
            "blocked_profile",
            "matching",
            matching_issues,
            questions,
        )
    research_domains = official_domains_for_target(
        target, official_domains + research_domains
    )
    official_evidence_path = run_dir / "04_공식근거.json"
    if not official_evidence_path.exists():
        write_json(official_evidence_path, [])
    _write_research_execution_template(run_dir)
    state = {
        "status": "ready_for_research",
        "quality_mode": "v2",
        "strict_quality": True,
        "run_dir": str(run_dir),
        "root": str(root),
        "target": target,
        "draft": str(draft),
        "posting": posting,
        "profile": str(profile),
        "posting_snapshot_id": analysis.source.content_sha256,
        "official_research_domains": list(research_domains),
        "research_policy": REQUIRED_RESEARCH_POLICY,
        "research_method_default": DEFAULT_RESEARCH_METHOD,
        "research_method_enforced": False,
        "questions": [asdict(question) for question in reconciliation.questions],
        "question_requirement_map": "05_문항전략.json",
        "selected_experience_ids": [
            item.recommended.experience_id
            for item in matches
            if item.recommended is not None
        ],
        "conflict_count": 0,
    }
    attach_writing_guidance(root, run_dir, state)
    write_state(run_dir, state)
    return state


def prepare_run(
    root: Path,
    target: str,
    draft: Path,
    posting: str | None,
    run_name: str | None,
    resume: Path | None = None,
    *,
    profile: Path | None = None,
    official_domains: tuple[str, ...] = (),
    research_domains: tuple[str, ...] = (),
    official_source: bool = False,
) -> dict:
    root = root.resolve()
    draft = draft.resolve()
    run_dir = resolve_run_dir(root, target, run_name, resume)
    inventory = build_inventory(root)
    draft_record = next(
        item for item in inventory if item.path.resolve() == draft
    )
    if draft_record.status == "failed":
        raise PermissionError(
            "대상 초안을 읽을 수 없습니다. 초안 파일을 닫고 다시 실행해 주세요. "
            f"원인: {draft_record.reason}"
        )
    documents = []
    for index, source in enumerate(inventory):
        if source.status != "use":
            continue
        try:
            documents.append(extract_path(source))
        except Exception as error:
            inventory[index] = replace(
                source, status="failed", reason=f"{type(error).__name__}: {error}"
            )

    if draft_record.extension == ".docx":
        draft_document = Document(draft_record.path)
        question_blocks = [
            " ".join(paragraph.text.split())
            for paragraph in draft_document.paragraphs
            if paragraph.text.strip()
        ]
        for table in draft_document.tables:
            for row in table.rows:
                for cell in row.cells:
                    question_blocks.extend(
                        " ".join(line.split())
                        for line in cell.text.splitlines()
                        if line.strip()
                    )
        questions = extract_questions(tuple(question_blocks))
    else:
        questions = extract_questions(extract_path(draft_record).paragraphs)
    (run_dir / "01_자료목록.md").write_text(
        _inventory_markdown(inventory), encoding="utf-8"
    )
    if profile is not None:
        return _prepare_v2(
            root=root,
            target=target,
            draft=draft,
            posting=posting,
            run_dir=run_dir,
            profile=profile.resolve(),
            official_domains=official_domains,
            research_domains=research_domains,
            official_source=official_source,
            questions=questions,
        )

    fact_documents = [
        document
        for document in documents
        if is_evidence_path(document.source.relative_path)
    ]
    claims = extract_fact_claims(fact_documents)
    overrides = _load_overrides(run_dir / "fact_overrides.yaml")
    accepted = apply_overrides(claims, overrides)
    conflicts = detect_conflicts(accepted)

    fact_payload = []
    for claim in claims:
        item = asdict(claim)
        item["tokens"] = sorted(claim.tokens)
        fact_payload.append(item)
    write_json(run_dir / "02_사실원장.json", fact_payload)
    (run_dir / "03_충돌검사.md").write_text(
        _conflict_markdown(conflicts, accepted), encoding="utf-8"
    )

    state = {
        "status": "blocked_conflict" if conflicts else "ready_for_research",
        "quality_mode": "legacy",
        "run_dir": str(run_dir),
        "root": str(root),
        "target": target,
        "draft": str(draft),
        "posting": posting,
        "research_policy": REQUIRED_RESEARCH_POLICY,
        "research_method_default": DEFAULT_RESEARCH_METHOD,
        "research_method_enforced": False,
        "questions": [asdict(question) for question in questions],
        "conflict_count": len(conflicts),
    }
    _write_research_execution_template(run_dir)
    attach_writing_guidance(root, run_dir, state)
    write_state(run_dir, state)
    return state


def _write_review_report(
    run_dir: Path,
    questions: list[Question],
    responses: list[DraftResponse],
    *,
    v2: bool,
    issues: list[ValidationIssue],
) -> None:
    response_by_index = {item.question_index: item for item in responses}
    review_lines = ["# 자기소개서 검토보고서", ""]
    count_rows: list[dict[str, object]] = []
    if v2:
        status = "통과" if not issues else "실패"
        review_lines.extend(
            [
                f"- 경험 원장: {status}",
                f"- 공고 공식성: {status}",
                f"- 경험·문항 매칭: {status}",
                f"- stale 근거: {'없음' if not issues else '검토 필요'}",
            ]
        )
    for question in questions:
        response = response_by_index.get(question.index)
        if response is None:
            continue
        actual_count = count_characters(response.answer, question.count_mode)
        if question.character_limit:
            if question.character_limit <= 800:
                target_min = round(question.character_limit * 5 / 6)
                target_max = round(question.character_limit * 11 / 12)
            else:
                target_min = round(question.character_limit * 0.75)
                target_max = round(question.character_limit * 0.9)
            target_status = "PASS" if target_min <= actual_count <= target_max else "REVIEW_REQUIRED"
        else:
            target_min = target_max = None
            target_status = "NOT_APPLICABLE"
        count_rows.append(
            {
                "question_index": question.index,
                "answer_sha256": sha256(response.answer.encode("utf-8")).hexdigest(),
                "actual_count": actual_count,
                "hard_limit": question.character_limit,
                "target_min": target_min,
                "target_max": target_max,
                "count_mode": question.count_mode,
                "newline_policy": "counted_as_stored",
                "hard_limit_status": "PASS" if not question.character_limit or actual_count <= question.character_limit else "FAIL",
                "target_status": target_status,
                "headroom": question.character_limit - actual_count if question.character_limit else None,
                "metric_type": "FORMAT_CHECK",
                "is_applicant_fact_or_claim": False,
            }
        )
        review_lines.append(
            f"- 문항 {question.index}: {actual_count}/"
            f"{question.character_limit or '미지정'}자 "
            f"({'공백 제외' if question.count_mode == 'spaces_excluded' else '공백 포함'}), "
            f"경험 근거 {len(response.experience_refs)}개, "
            f"공식 근거 {len(response.research_refs)}개"
        )
    if issues:
        review_lines.extend(["", "## 검증 이슈", ""])
        review_lines.extend(
            f"- `{issue.code}` 문항 {issue.question_index}: {issue.message}"
            for issue in issues
        )
    else:
        review_lines.extend(
            ["- 블라인드: 통과", "- 타기관명: 통과", "- 빈 답변: 없음"]
        )
    (run_dir / "07_자기소개서_검토보고서.md").write_text(
        "\n".join(review_lines) + "\n", encoding="utf-8"
    )
    write_json(
        run_dir / "07_글자수검증.json",
        {
            "schema_version": 1,
            "count_function": "career_pipeline.character_count.count_characters",
            "target_policy": "limit<=800: 5/6..11/12, limit>800: 0.75..0.90",
            "rows": count_rows,
        },
    )


def _incumbent_from_markdown(
    path: Path,
    baseline: list[DraftResponse],
) -> list[DraftResponse]:
    text = path.read_text(encoding="utf-8")
    matches = list(re.finditer(r"(?m)^#{2,3}\s*(?:문항\s*)?(\d+)\s*$", text))
    if not matches:
        raise ValueError("incumbent markdown has no numbered headings")
    answers: dict[int, str] = {}
    for index, match in enumerate(matches):
        end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
        answer = text[match.end():end].strip()
        answer = re.sub(r"(?m)^`?글자\s*수\s*:.*$", "", answer).strip()
        answers[int(match.group(1))] = answer
    baseline_by_index = {item.question_index: item for item in baseline}
    if set(answers) != set(baseline_by_index):
        raise ValueError("incumbent question set does not match the run")
    return [
        replace(baseline_by_index[index], answer=answers[index])
        for index in sorted(answers)
    ]


def finalize_run(
    run_dir: Path,
    *,
    copyedit: bool = False,
    copyeditor_timeout_ms: int = 180_000,
    humanize: bool = False,
    patina_backend: str = "codex-cli",
    patina_timeout_ms: int = 180_000,
    patina_max_retries: int = 1,
    patina_voice_sample: Path | None = None,
    patina_ai_threshold: int = 30,
    patina_score: bool = True,
    postprocess: str | None = None,
    postprocess_tier: ModelTier | None = None,
    postprocess_timeout_ms: int | None = None,
    postprocess_runner=None,
    max_model_calls: int | None = None,
    max_postprocess_calls: int = 1,
    max_stage_seconds: float | None = None,
    selection_mode: str = "single",
    incumbent_path: Path | None = None,
    rigorous_runner=None,
    rigorous_timeout_ms: int = 300_000,
    quality_profile: str | None = None,
) -> dict:
    run_dir = run_dir.resolve()
    state = json.loads((run_dir / "run.json").read_text(encoding="utf-8"))
    previous_status = str(state.get("status", ""))
    legacy_copyedit = postprocess is None and copyedit
    effective_postprocess = postprocess
    if effective_postprocess is None:
        effective_postprocess = "always" if copyedit else "never"
    if effective_postprocess not in {"auto", "always", "never"}:
        raise ValueError("postprocess must be auto, always, or never")
    if selection_mode not in {"single", "rigorous"}:
        raise ValueError("selection_mode must be single or rigorous")
    resolved_profile = get_quality_profile(quality_profile) if quality_profile else None
    if resolved_profile is not None:
        selection_mode = resolved_profile.selection_mode
    if selection_mode == "rigorous" and (copyedit or humanize):
        raise ValueError("rigorous selection requires copyedit and patina disabled")
    if selection_mode == "rigorous":
        effective_postprocess = "never"
    rigorous_profile = (
        resolved_profile if selection_mode == "rigorous" else None
    )
    if selection_mode == "rigorous" and rigorous_profile is None:
        rigorous_profile = legacy_rigorous_profile()
    required_selection_calls = (
        rigorous_profile.max_selection_calls if rigorous_profile is not None else 0
    )
    if (
        selection_mode == "rigorous"
        and max_model_calls is not None
        and max_model_calls < required_selection_calls
    ):
        raise ValueError(
            f"{rigorous_profile.name} selection requires max_model_calls >= "
            f"{required_selection_calls}"
        )
    state["selection_mode"] = selection_mode
    state["quality_profile"] = (
        resolved_profile.name
        if resolved_profile is not None
        else ("legacy_rigorous" if selection_mode == "rigorous" else "legacy_single")
    )
    state["postprocess_policy"] = effective_postprocess
    state["max_model_calls"] = max_model_calls
    state["max_postprocess_calls"] = max_postprocess_calls
    state["max_stage_seconds"] = max_stage_seconds
    state["legacy_patina"] = bool(humanize)
    if not humanize:
        # A finalize rerun defines the current execution mode.  Do not let a
        # legacy opt-in from an older run keep affecting audits or manifests.
        state["patina_attempted"] = False
        state["patina_score_attempted"] = False
        state["patina_applied"] = False
        state["patina_status"] = "disabled"
        state["patina_score_enabled"] = False
        state["patina_voice_sample_used"] = None
        state.pop("patina_summary", None)
    if state.get("status") in {
        "blocked",
        "blocked_profile",
        "blocked_posting",
        "blocked_conflict",
    }:
        raise ValueError("준비 단계의 차단 이슈를 먼저 해결해야 합니다.")
    # Prevent a previous successful run from remaining falsely complete if a
    # rerun fails before its new final manifest is committed.
    state["status"] = "finalizing"
    state["final_artifact"] = None
    write_state(run_dir, state)

    v2 = state.get("quality_mode") == "v2"
    required = [
        "04_기업직무조사.md",
        "05_문항전략.md",
        "08_면접대비팩.md",
        "draft.json",
    ]
    required.extend(
        [
            "00_채용공고분석.json",
            "02_확정경험원장.json",
            "03_경험직무매칭.json",
        ]
        if v2
        else ["02_사실원장.json"]
    )
    if v2 and state.get("strict_quality", False):
        required.append("04_공식근거.json")
    if state.get("research_policy") == REQUIRED_RESEARCH_POLICY:
        required.append("04_리서치실행.json")
    missing = [name for name in required if not (run_dir / name).exists()]
    if missing:
        raise FileNotFoundError(
            f"missing synthesis artifacts: {', '.join(missing)}"
        )

    questions = [Question(**item) for item in state["questions"]]
    responses, draft_issues = _load_draft_responses(run_dir / "draft.json")
    if draft_issues:
        _write_review_report(run_dir, questions, responses, v2=v2, issues=draft_issues)
        state.update(
            status="blocked_validation",
            validation_issues=[asdict(item) for item in draft_issues],
        )
        if v2:
            state["blocked_stage"] = "finalize"
            state["issues"] = [asdict(item) for item in draft_issues]
        write_state(run_dir, state)
        return state

    if selection_mode == "rigorous" and incumbent_path is not None:
        try:
            responses = _incumbent_from_markdown(incumbent_path.resolve(), responses)
        except (OSError, ValueError) as error:
            state.update(
                status="blocked_validation",
                blocked_stage="rigorous_incumbent",
                validation_issues=[{
                    "code": "invalid_incumbent",
                    "question_index": 0,
                    "message": str(error),
                }],
            )
            write_state(run_dir, state)
            return state

    ledger = None
    research_claims = ()
    job_terms: tuple[str, ...] = ()
    if v2:
        ledger = load_ledger(run_dir / "02_확정경험원장.json")
        posting_payload = json.loads(
            (run_dir / "00_채용공고분석.json").read_text(encoding="utf-8")
        )
        job_terms = tuple(
            posting_payload.get("duties", []) + posting_payload.get("competencies", [])
        )
        known_sources = {
            evidence.source_path
            for experience in ledger.experiences
            for claim in experience.claims
            for evidence in claim.evidence
        }
    else:
        fact_data = json.loads(
            (run_dir / "02_사실원장.json").read_text(encoding="utf-8")
        )
        known_sources = {item["source_path"] for item in fact_data}
    issues = validate_draft(
        questions,
        responses,
        state["target"],
        known_sources,
        profile_ledger=ledger,
        require_experience_refs=v2,
    )
    if v2 and state.get("strict_quality", False):
        write_json(
            run_dir / "10_품질점수.json",
            [
                {
                    "question_index": question.index,
                    "score": asdict(
                        score_answer_quality(
                            question,
                            next(
                                response.answer
                                for response in responses
                                if response.question_index == question.index
                            ),
                            state["target"],
                            job_terms=job_terms,
                        )
                    ),
                }
                for question in questions
                if any(
                    response.question_index == question.index
                    for response in responses
                )
            ],
        )
        issues.extend(
            validate_answer_quality(
                questions,
                responses,
                state["target"],
                job_terms=job_terms,
                minimum_score=STRICT_MIN_ANSWER_SCORE,
                average_minimum_score=STRICT_MIN_AVERAGE_SCORE,
            )
        )
        try:
            research_claims = load_research_claims(run_dir / "04_공식근거.json")
        except (OSError, ValueError, TypeError, json.JSONDecodeError) as error:
            issues.append(
                ValidationIssue(
                    "invalid_research_evidence",
                    0,
                    f"공식 근거 JSON을 읽을 수 없습니다: {error}",
                )
            )
        else:
            issues.extend(
                validate_research_evidence(
                    questions,
                    responses,
                    research_claims,
                    allowed_domains=tuple(
                        state.get("official_research_domains", [])
                    ),
                )
            )

    if state.get("research_policy") == REQUIRED_RESEARCH_POLICY:
        try:
            research_execution = load_research_execution(
                run_dir / "04_리서치실행.json"
            )
        except (OSError, ValueError, TypeError, json.JSONDecodeError) as error:
            issues.append(
                ValidationIssue(
                    "invalid_research_execution",
                    0,
                    f"기업조사 실행 기록을 읽을 수 없습니다: {error}",
                )
            )
        else:
            issues.extend(
                validate_research_execution(research_execution, research_claims)
            )

    research = (run_dir / "04_기업직무조사.md").read_text(encoding="utf-8")
    if "http://" not in research and "https://" not in research:
        issues.append(
            ValidationIssue(
                "missing_research_link", 0, "공식 조사 링크가 없습니다."
            )
        )

    interview = (run_dir / "08_면접대비팩.md").read_text(encoding="utf-8")
    for section in ("1분 자기소개", "꼬리질문", "압박질문", "근거"):
        if section not in interview:
            issues.append(
                ValidationIssue(
                    "missing_interview_section",
                    0,
                    f"면접팩 누락: {section}",
                )
            )
    if v2 and ledger is not None:
        allowed_values = referenced_claim_values(responses, ledger)
        for match in METRIC.finditer(interview):
            normalized, _ = _normalize(match.group("number"), match.group("unit"))
            if normalized not in allowed_values:
                issues.append(
                    ValidationIssue(
                        "unapproved_interview_metric",
                        0,
                        f"면접팩의 승인되지 않은 수치: {match.group(0)}",
                    )
                )
        if state.get("strict_quality", False):
            issues.extend(
                validate_interview_pack(
                    interview,
                    questions,
                    responses,
                    allowed_metric_values=allowed_values,
                )
            )

    if previous_status in {"blocked_selection", "complete"} and selection_mode == "rigorous":
        # A failed final contract revalidation may have left the sidecar linked
        # to the attempted selection. Restore the incumbent contract before
        # preflight; the selected candidate replaces it atomically later.
        _replace_interview_contract_claims(run_dir, responses)
    contract_report = validate_run_prompt_contracts(
        run_dir,
        target=state["target"],
        responses=responses,
    )
    if contract_report.enabled:
        state["prompt_contracts"] = contract_report.to_dict()
        issues.extend(
            ValidationIssue(issue.code, 0, issue.message)
            for issue in contract_report.issues
            if issue.severity == "HARD_FAIL"
        )

    _write_review_report(
        run_dir, questions, responses, v2=v2, issues=issues
    )
    if selection_mode == "rigorous" and issues:
        # The incumbent is an anonymous comparison candidate, not the final
        # submission.  Its response-level failures must remain visible to the
        # judges instead of aborting independent candidate generation.  Run-
        # level failures (question_index == 0) still fail closed.
        state["incumbent_preflight_issues"] = [
            asdict(item) for item in issues if item.question_index != 0
        ]
        issues = [item for item in issues if item.question_index == 0]
    if issues:
        state.update(
            status="blocked_validation",
            validation_issues=[asdict(item) for item in issues],
        )
        if v2:
            state["blocked_stage"] = "finalize"
            state["issues"] = [asdict(item) for item in issues]
        write_state(run_dir, state)
        return state

    if legacy_copyedit:
        copyedited, copyeditor_report = copyedit_responses(
            responses,
            target_org=state["target"],
            job_terms=job_terms,
            timeout_ms=copyeditor_timeout_ms,
        )
        copyedit_issues = validate_draft(
            questions,
            copyedited,
            state["target"],
            known_sources,
            profile_ledger=ledger,
            require_experience_refs=v2,
        )
        if v2 and state.get("strict_quality", False):
            copyedit_issues.extend(
                validate_answer_quality(
                    questions,
                    copyedited,
                    state["target"],
                    job_terms=job_terms,
                    minimum_score=STRICT_MIN_ANSWER_SCORE,
                    average_minimum_score=STRICT_MIN_AVERAGE_SCORE,
                )
            )
            copyedit_issues.extend(
                validate_research_evidence(
                    questions,
                    copyedited,
                    research_claims,
                    allowed_domains=tuple(
                        state.get("official_research_domains", [])
                    ),
                )
            )
        if copyedit_issues:
            copyeditor_report.append(
                {
                    "question_index": 0,
                    "status": "fallback_validation",
                    "message": "; ".join(
                        issue.code for issue in copyedit_issues
                    ),
                }
            )
            state["copyeditor_status"] = "fallback_validation"
            state["copyeditor_applied"] = False
        else:
            responses = copyedited
            state["copyeditor_applied"] = any(
                item.get("status") == "copyedited"
                for item in copyeditor_report
            )
            copyeditor_fallback = any(
                str(item.get("status", "")).startswith("fallback_")
                for item in copyeditor_report
            )
            state["copyeditor_status"] = (
                "copyedited"
                if state["copyeditor_applied"]
                else "fallback" if copyeditor_fallback else "unchanged"
            )
            write_json(
                run_dir / "draft_copyedited.json",
                [asdict(response) for response in responses],
            )
        state["copyeditor_attempted"] = True
        write_json(run_dir / "09_copyeditor_report.json", copyeditor_report)
    else:
        # Record explicit disablement so a stale fallback report cannot affect
        # a later audit of the same run directory.
        write_json(
            run_dir / "09_copyeditor_report.json",
            [
                {
                    "question_index": question.index,
                    "status": "disabled",
                    "message": "copyeditor disabled by explicit finalize option",
                    "applied_rules": [],
                    "change_ratio": 0.0,
                }
                for question in questions
            ],
        )
        state["copyeditor_attempted"] = False
        state["copyeditor_applied"] = False
        state["copyeditor_status"] = "disabled"

    postprocess_attempted = False
    postprocess_applied = False
    postprocess_model_id: str | None = None
    postprocess_budget_blocked = False
    model_unconfigured = False
    selected_source = "draft"
    postprocess_report: list[dict[str, object]] = []
    call_tracker = CostTracker(
        max_model_calls
        if max_model_calls is not None
        else 1_000_000 if humanize else max_postprocess_calls,
        max_postprocess_calls=max_postprocess_calls,
        max_stage_seconds=max_stage_seconds,
    )
    if not legacy_copyedit and not humanize:
        diagnostics = diagnose_responses(responses)
        write_json(
            run_dir / "09_style_diagnostics.json",
            [item.to_dict() for item in diagnostics],
        )
        target_responses = (
            responses
            if effective_postprocess == "always"
            else [
                response
                for response, diagnostic in zip(responses, diagnostics)
                if effective_postprocess == "auto" and diagnostic.should_rewrite
            ]
        )
        should_call = bool(target_responses)
        if effective_postprocess == "never":
            postprocess_report = [
                {
                    "question_index": item.question_index,
                    "status": "disabled",
                    "style_risk_score": diagnostic.style_risk_score,
                    "style_reasons": list(diagnostic.style_reasons),
                }
                for item, diagnostic in zip(responses, diagnostics)
            ]
        elif not should_call:
            postprocess_report = [
                {
                    "question_index": item.question_index,
                    "status": "skipped_style_pass",
                    "style_risk_score": diagnostic.style_risk_score,
                    "style_reasons": list(diagnostic.style_reasons),
                }
                for item, diagnostic in zip(responses, diagnostics)
            ]
        else:
            selected_diagnostics = [
                diagnostic
                for diagnostic in diagnostics
                if effective_postprocess == "always" or diagnostic.should_rewrite
            ]
            model = choose_tier(selected_diagnostics, postprocess_tier)
            postprocess_tier = model.tier
            postprocess_model_id = model.model_id
            model_unconfigured = postprocess_runner is None and model.model_id is None
            if model_unconfigured:
                # Never let the CLI silently select an unspecified external model.
                call_tracker.budget = 0
            try:
                call_tracker.record_call(
                    "copyeditor",
                    stage="postprocess",
                    model_tier=model.tier,
                    model_id=model.model_id,
                )
            except CostLimitExceeded as error:
                postprocess_budget_blocked = not model_unconfigured
                postprocess_report = [
                    {
                        "question_index": item.question_index,
                        "status": "skipped_model_unconfigured" if model_unconfigured else "fallback_budget_exceeded",
                        "message": "실제 모델 ID가 설정되지 않아 외부 호출을 건너뛰었습니다." if model_unconfigured else str(error),
                    }
                    for item in target_responses
                ]
            else:
                postprocess_attempted = True
                effective_timeout_ms = postprocess_timeout_ms or copyeditor_timeout_ms
                if max_stage_seconds is not None:
                    effective_timeout_ms = min(
                        effective_timeout_ms,
                        max(1, int(max_stage_seconds * 1000)),
                    )
                kwargs = {
                    "target_org": state["target"],
                    "job_terms": job_terms,
                    "diagnostics_by_index": {
                        diagnostic.question_index: diagnostic.style_reasons
                        for diagnostic in selected_diagnostics
                    },
                    "timeout_ms": effective_timeout_ms,
                    "model_tier": model.tier,
                    "model_id": model.model_id,
                }
                if postprocess_runner is not None:
                    kwargs["runner"] = postprocess_runner
                within_stage_limit = True
                try:
                    edited, postprocess_report = copyedit_responses(
                        target_responses,
                        **kwargs,
                    )
                    call_status = "complete"
                except Exception as error:  # external runner failures are safe fallbacks
                    edited = target_responses
                    postprocess_report = [
                        {
                            "question_index": item.question_index,
                            "status": "fallback_backend_error",
                            "message": str(error),
                        }
                        for item in target_responses
                    ]
                    call_status = "failed"
                finally:
                    within_stage_limit = call_tracker.finish_call(status=call_status)
                if not within_stage_limit:
                    edited = target_responses
                    postprocess_report.append(
                        {
                            "question_index": 0,
                            "status": "fallback_stage_timeout",
                            "message": f"max_stage_seconds={max_stage_seconds}",
                        }
                    )
                elif any(
                    str(item.get("status", "")).startswith("fallback")
                    for item in postprocess_report
                    if isinstance(item, dict)
                ):
                    call_tracker.set_last_status("fallback")
                edited_by_index = {item.question_index: item for item in edited}
                target_indexes = {response.question_index for response in target_responses}
                original_responses = responses
                candidate = [
                    edited_by_index.get(item.question_index, item)
                    if item.question_index in target_indexes else item
                    for item in responses
                ]
                changed = any(
                    candidate_item.answer != original_item.answer
                    for candidate_item, original_item in zip(candidate, original_responses)
                )
                candidate_issues = []
                if changed:
                    candidate_issues.extend(
                        validate_draft(
                            questions,
                            candidate,
                            state["target"],
                            known_sources,
                            profile_ledger=ledger,
                            require_experience_refs=v2,
                        )
                    )
                    if v2 and state.get("strict_quality", False):
                        candidate_issues.extend(
                            validate_answer_quality(
                                questions,
                                candidate,
                                state["target"],
                                job_terms=job_terms,
                                minimum_score=STRICT_MIN_ANSWER_SCORE,
                                average_minimum_score=STRICT_MIN_AVERAGE_SCORE,
                            )
                        )
                        candidate_issues.extend(
                            validate_research_evidence(
                                questions,
                                candidate,
                                research_claims,
                                allowed_domains=tuple(state.get("official_research_domains", [])),
                            )
                        )
                bad_indexes = {
                    issue.question_index
                    for issue in candidate_issues
                    if issue.question_index in target_indexes
                }
                if any(issue.question_index == 0 for issue in candidate_issues):
                    bad_indexes = set(target_indexes)
                if bad_indexes:
                    messages: dict[int, list[str]] = {
                        question_index: [] for question_index in bad_indexes
                    }
                    for issue in candidate_issues:
                        for question_index in bad_indexes:
                            if issue.question_index in {0, question_index}:
                                messages[question_index].append(issue.code)
                    candidate = [
                        original if item.question_index in bad_indexes else item
                        for item, original in zip(candidate, original_responses)
                    ]
                    for question_index, codes in messages.items():
                        postprocess_report.append(
                            {
                                "question_index": question_index,
                                "status": "fallback_validation",
                                "message": "; ".join(codes),
                            }
                        )
                    call_tracker.set_last_status("fallback_validation")
                responses = candidate
                postprocess_applied = any(
                    item.answer != original.answer
                    for item, original in zip(responses, original_responses)
                )
                selected_source = "copyedited" if postprocess_applied else "draft"
                if postprocess_applied:
                    write_json(
                        run_dir / "draft_copyedited.json",
                        [asdict(response) for response in responses],
                    )
            state["postprocess_attempted"] = postprocess_attempted
            state["postprocess_applied"] = postprocess_applied
            state["postprocess_status"] = (
                "copyedited" if postprocess_applied
                else "fallback" if postprocess_attempted
                else "model_unconfigured" if model_unconfigured
                else "budget_exceeded" if postprocess_budget_blocked else "not_needed"
            )
            state["copyeditor_attempted"] = postprocess_attempted
            state["copyeditor_applied"] = postprocess_applied
            state["copyeditor_status"] = state["postprocess_status"]
            write_json(run_dir / "09_copyeditor_report.json", postprocess_report)
            state["postprocess_tier"] = postprocess_tier
            state["postprocess_model_id"] = postprocess_model_id
        if not should_call:
            state["postprocess_attempted"] = False
            state["postprocess_applied"] = False
            state["postprocess_status"] = (
                "model_unconfigured" if model_unconfigured
                else "budget_exceeded" if postprocess_budget_blocked
                else "not_needed" if effective_postprocess == "auto" else "disabled"
            )
            state["copyeditor_attempted"] = False
            state["copyeditor_applied"] = False
            state["copyeditor_status"] = state["postprocess_status"]
            write_json(run_dir / "09_copyeditor_report.json", postprocess_report)
        if model_unconfigured:
            call_tracker.budget = (
                max_model_calls
                if max_model_calls is not None
                else max_postprocess_calls
            )
        state["model_calls"] = call_tracker.to_dict()

    if humanize:
        legacy_timeout_ms = patina_timeout_ms
        if max_stage_seconds is not None:
            legacy_timeout_ms = min(
                legacy_timeout_ms,
                max(1, int(max_stage_seconds * 1000)),
            )

        def tracked_humanize(text: str, **kwargs) -> HumanizationResult:
            try:
                call_tracker.record_call("patina", stage="legacy_patina")
            except CostLimitExceeded:
                return HumanizationResult(text, "fallback_budget_exceeded", "model call budget exceeded")
            kwargs["timeout_ms"] = legacy_timeout_ms
            if postprocess_runner is not None:
                kwargs["runner"] = postprocess_runner
            try:
                result = humanize_text(text, **kwargs)
                call_tracker.finish_call(
                    status="fallback" if result.status.startswith("fallback") else "complete"
                )
                return result
            except Exception:
                call_tracker.finish_call(status="failed")
                raise

        def tracked_score(text: str, **kwargs) -> PatinaScoreResult:
            try:
                call_tracker.record_call("patina", stage="legacy_patina")
            except CostLimitExceeded:
                return PatinaScoreResult(None, "score_unavailable", "model call budget exceeded")
            kwargs["timeout_ms"] = legacy_timeout_ms
            if postprocess_runner is not None:
                kwargs["runner"] = postprocess_runner
            try:
                result = score_text(text, **kwargs)
                call_tracker.finish_call(
                    status="fallback" if result.score is None else "complete"
                )
                return result
            except Exception:
                call_tracker.finish_call(status="failed")
                raise

        voice_sample = _resolve_voice_sample(patina_voice_sample, state, run_dir)
        state["patina_backend"] = patina_backend
        state["patina_max_retries"] = patina_max_retries
        state["patina_ai_threshold"] = patina_ai_threshold
        state["patina_score_enabled"] = patina_score
        state["patina_voice_sample_used"] = str(voice_sample) if voice_sample else None
        humanized, patina_report = generate_and_select_candidates(
            responses,
            questions,
            state["target"],
            job_terms=job_terms,
            backend=patina_backend,
            timeout_ms=legacy_timeout_ms,
            voice_sample=voice_sample,
            max_retries=patina_max_retries,
            scorer=tracked_score if patina_score else None,
            ai_score_threshold=patina_ai_threshold,
            conditional_rewrite=patina_score,
            rewriter=tracked_humanize,
        )
        state["patina_attempted"] = any(
            bool(item.get("patina_attempted")) for item in patina_report
        )
        state["patina_score_attempted"] = any(
            bool(item.get("patina_score_attempted"))
            for item in patina_report
        )
        post_issues = validate_draft(
            questions,
            humanized,
            state["target"],
            known_sources,
            profile_ledger=ledger,
            require_experience_refs=v2,
        )
        if v2 and state.get("strict_quality", False):
            post_issues.extend(
                validate_answer_quality(
                    questions,
                    humanized,
                    state["target"],
                    job_terms=job_terms,
                    minimum_score=STRICT_MIN_ANSWER_SCORE,
                    average_minimum_score=STRICT_MIN_AVERAGE_SCORE,
                )
            )
            post_issues.extend(
                validate_research_evidence(
                    questions,
                    humanized,
                    research_claims,
                    allowed_domains=tuple(
                        state.get("official_research_domains", [])
                    ),
                )
            )
        if post_issues:
            patina_report.append(
                {
                    "question_index": 0,
                    "status": "fallback_validation",
                    "message": "; ".join(issue.code for issue in post_issues),
                    "backend": patina_backend,
                }
            )
            state["patina_status"] = "fallback_validation"
            state["patina_applied"] = False
        else:
            responses = humanized
            selected_variants = {
                str(item["selected_variant"]) for item in patina_report
            }
            fallback_count = sum(
                str(candidate["status"]).startswith("fallback_")
                for item in patina_report
                for candidate in item["candidates"]
            )
            if selected_variants.difference({"original", "copyedited"}):
                state["patina_status"] = "humanized"
            elif selected_variants == {"copyedited"} and all(
                item.get("ai_score_gate") == "passed"
                for item in patina_report
            ):
                state["patina_status"] = "not_needed"
            elif any(
                item.get("ai_score_gate") in {"failed", "unavailable"}
                for item in patina_report
            ):
                state["patina_status"] = "fallback_score_gate"
            elif fallback_count:
                state["patina_status"] = "fallback"
            else:
                state["patina_status"] = "unchanged"
            applied_count = sum(
                str(item["selected_variant"]) not in {"original", "copyedited"}
                for item in patina_report
            )
            state["patina_attempted"] = any(
                bool(
                    item.get(
                        "patina_attempted",
                        item.get("selected_variant") not in {"original", "copyedited"},
                    )
                )
                for item in patina_report
            )
            state["patina_score_attempted"] = any(
                bool(item.get("patina_score_attempted"))
                for item in patina_report
            )
            state["patina_applied"] = applied_count > 0
            state["patina_summary"] = {
                "attempted_questions": sum(
                    bool(
                        item.get(
                            "patina_attempted",
                            item.get("selected_variant")
                            not in {"original", "copyedited"},
                        )
                    )
                    for item in patina_report
                ),
                "applied_questions": applied_count,
                "baseline_selected_questions": len(patina_report) - applied_count,
                "fallback_candidates": fallback_count,
                "score_passed_questions": sum(
                    item.get("ai_score_gate") == "passed"
                    for item in patina_report
                ),
                "score_unavailable_questions": sum(
                    item.get("ai_score_gate") == "unavailable"
                    for item in patina_report
                ),
                "headroom_met_questions": sum(
                    bool(item.get("headroom_target_met"))
                    for item in patina_report
                ),
            }
            write_json(
                run_dir / "draft_humanized.json",
                [asdict(response) for response in responses],
            )
        write_json(run_dir / "09_patina_report.json", patina_report)
        write_json(run_dir / "09_초안후보평가.json", patina_report)
        write_json(
            run_dir / "10_품질점수.json",
            [
                {
                    "question_index": item["question_index"],
                    "selected_variant": item["selected_variant"],
                    "score": item["selected_score"],
                    "score_semantics": {
                        "metric_type": "INTERNAL_EVALUATION",
                        "is_applicant_fact_or_claim": False,
                        "allowed_use": "후보 품질 비교",
                        "prohibited_use": "지원자 경험 성과 수치로 인용",
                    },
                }
                for item in patina_report
                if "selected_score" in item
            ],
        )
        _write_review_report(run_dir, questions, responses, v2=v2, issues=[])

    if selection_mode == "rigorous":
        posting_frozen = json.loads(
            (run_dir / "00_채용공고분석.json").read_text(encoding="utf-8")
        )
        research_frozen = []
        research_path = run_dir / "04_공식근거.json"
        if research_path.exists():
            research_frozen = json.loads(research_path.read_text(encoding="utf-8"))
        frozen_packet = {
            "target": state["target"],
            "posting": posting_frozen,
            "experience_ledger": ledger_to_dict(ledger) if ledger is not None else None,
            "research_claims": research_frozen,
        }
        question_requirement_path = run_dir / "05_문항전략.json"
        if question_requirement_path.is_file():
            question_requirement_map = json.loads(
                question_requirement_path.read_text(encoding="utf-8")
            )
        else:
            question_requirement_map = build_question_requirement_map(
                questions,
                target=state["target"],
                posting=posting_frozen,
            )
        frozen_packet["question_requirement_map"] = question_requirement_map
        contract_context = prompt_contract_context(run_dir)
        if contract_context is not None:
            frozen_packet["prompt_contracts"] = contract_context

        def validate_rigorous_candidate(candidate: list[DraftResponse]):
            _hydrate_claim_evidence_paths(candidate, ledger)
            candidate_issues = validate_draft(
                questions,
                candidate,
                state["target"],
                known_sources,
                profile_ledger=ledger,
                require_experience_refs=v2,
            )
            if v2 and state.get("strict_quality", False):
                # Qualitative preferences (density, repeated experience,
                # narrative strength) are decided by the three independent
                # judges.  The deterministic gate remains limited to facts,
                # exact claim references, question limits and official-source
                # linkage so that a style heuristic cannot eliminate every
                # candidate before blind selection.
                candidate_issues.extend(validate_research_evidence(
                    questions,
                    candidate,
                    research_claims,
                    allowed_domains=tuple(state.get("official_research_domains", [])),
                ))
                candidate_issues.extend(
                    validate_question_requirement_map(
                        candidate,
                        question_requirement_map,
                        target=state["target"],
                        enforce_preferred_range=(
                            rigorous_profile is not None
                            and rigorous_profile.name == "max_quality"
                        ),
                    )
                )
            return candidate_issues

        sol_model = resolve_model("sol").model_id
        role_models = {
            role: resolve_role_model(role).model_id or sol_model
            for role in ("generation", "judge", "synthesis", "comparison")
        }
        try:
            rigorous_result = run_rigorous_selection(
                run_dir,
                questions=questions,
                incumbent=tuple(responses),
                frozen_packet=frozen_packet,
                model_id=sol_model,
                validate_candidate=validate_rigorous_candidate,
                runner=rigorous_runner or subprocess_model_runner,
                max_calls=max_model_calls or required_selection_calls,
                timeout_ms=rigorous_timeout_ms,
                quality_profile=rigorous_profile,
                stage_models=role_models,
                resume_from_checkpoint=(
                    previous_status in {"blocked_selection", "complete"}
                ),
            )
        except (OSError, ValueError, RigorousSelectionError) as error:
            state.update(
                status="blocked_selection",
                blocked_stage="rigorous",
                rigorous_selection={
                    "status": "failed",
                    "selection_mode": "rigorous",
                    "model_id": sol_model,
                    "error": str(error),
                },
            )
            write_state(run_dir, state)
            return state
        responses = list(rigorous_result.responses)
        _link_final_claims_to_interview_pack(run_dir, responses)
        final_contract_report = validate_run_prompt_contracts(
            run_dir,
            target=state["target"],
            responses=responses,
        )
        if final_contract_report.enabled:
            state["prompt_contracts"] = final_contract_report.to_dict()
        final_contract_failures = [
            issue
            for issue in final_contract_report.issues
            if issue.severity == "HARD_FAIL"
        ]
        if final_contract_failures:
            state.update(
                status="blocked_selection",
                blocked_stage="prompt_contracts",
                validation_issues=[asdict(issue) for issue in final_contract_failures],
                rigorous_selection={
                    **rigorous_result.metadata,
                    "status": "failed_contract_revalidation",
                },
            )
            write_state(run_dir, state)
            return state
        state["rigorous_selection"] = rigorous_result.metadata
        selected_source = "rigorous"
        final_diagnostics = diagnose_responses(responses)
        write_json(
            run_dir / "09_style_diagnostics.json",
            [item.to_dict() for item in final_diagnostics],
        )
        write_json(
            run_dir / "09_copyeditor_report.json",
            [
                {
                    "question_index": item.question_index,
                    "status": "rigorous_integrated_style_pass",
                    "style_risk_score": diagnostic.style_risk_score,
                    "style_reasons": list(diagnostic.style_reasons),
                }
                for item, diagnostic in zip(responses, final_diagnostics)
            ],
        )
        write_json(
            run_dir / "10_품질점수.json",
            [
                {
                    "question_index": question.index,
                    "score": asdict(
                        score_answer_quality(
                            question,
                            next(
                                response.answer
                                for response in responses
                                if response.question_index == question.index
                            ),
                            state["target"],
                            job_terms=job_terms,
                        )
                    ),
                    "score_semantics": {
                        "metric_type": "INTERNAL_EVALUATION",
                        "is_applicant_fact_or_claim": False,
                        "allowed_use": "후보 품질 비교",
                        "prohibited_use": "지원자 경험 성과 수치로 인용",
                    },
                }
                for question in questions
            ],
        )
    else:
        state["rigorous_selection"] = {
            "status": "not_run", "selection_mode": "single", "hard_fail": False,
        }

    final_diagnostics = diagnose_responses(responses)
    write_json(
        run_dir / "09_style_diagnostics.json",
        [item.to_dict() for item in final_diagnostics],
    )
    write_json(
        run_dir / "10_품질점수.json",
        [
            {
                "question_index": question.index,
                "score": asdict(
                    score_answer_quality(
                        question,
                        next(item.answer for item in responses if item.question_index == question.index),
                        state["target"],
                        job_terms=job_terms,
                    )
                ),
                "score_semantics": {
                    "metric_type": "INTERNAL_EVALUATION",
                    "is_applicant_fact_or_claim": False,
                    "allowed_use": "후보 품질 비교",
                    "prohibited_use": "지원자 경험 성과 수치로 인용",
                },
            }
            for question in questions
        ],
    )
    state["model_calls"] = call_tracker.to_dict()
    output_paths = (
        run_dir / "06_자기소개서.md",
        run_dir / "06_자기소개서.docx",
        run_dir / "draft_final.json",
    )
    if any(path.is_symlink() for path in output_paths):
        raise ValueError("최종 산출물 경로에 심볼릭 링크가 있어 안전하게 저장할 수 없습니다.")
    markdown = render_draft_markdown(questions, responses)
    (run_dir / "06_자기소개서.md").write_text(markdown, encoding="utf-8")
    render_draft_docx(
        questions, responses, run_dir / "06_자기소개서.docx"
    )
    write_json(
        run_dir / "draft_final.json",
        [asdict(response) for response in responses],
    )
    state["status"] = "complete"
    state["finished_at"] = datetime.now().isoformat()
    state["run_duration_seconds"] = round(
        (datetime.fromisoformat(state["finished_at"])
         - datetime.fromisoformat(state.get("started_at", state["finished_at"]))).total_seconds(),
        2,
    )
    state.pop("validation_issues", None)
    state.pop("issues", None)
    state.pop("blocked_stage", None)
    state["final_artifact"] = write_final_artifact_manifest(
        run_dir,
        selected_source=(
            "legacy_patina" if humanize and state.get("patina_applied")
            else "copyedited" if state.get("copyeditor_applied") else selected_source
        ),
        postprocess_attempted=bool(state.get("postprocess_attempted", False)),
        postprocess_applied=bool(state.get("postprocess_applied", False)),
        model_tier=state.get("postprocess_tier"),
        model_id=state.get("postprocess_model_id"),
        validation={"status": "passed", "issues": []},
        selection=state.get("rigorous_selection"),
    )
    write_state(run_dir, state)
    return state
