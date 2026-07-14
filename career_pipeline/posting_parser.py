"""채용공고 파서. HTML/DOCX에서 기관명, 직무, 문항 등을 추출합니다. 레이블 없는 일반 텍스트도 처리합니다."""
from dataclasses import dataclass
from html.parser import HTMLParser
from io import BytesIO
import re

from docx import Document
from pypdf import PdfReader

from .extractors import split_text_blocks
from .models import Question
from .posting_loader import PostingSourceError
from .posting_schema import LoadedPosting, PostingAnalysis
from .questions import extract_questions


SECTION_MARKERS = {
    "organization": ("기관명", "채용기관", "회사명"),
    "role": ("채용분야", "모집분야", "직무"),
    "locations": ("근무지", "근무지역"),
    "duties": ("담당업무", "직무내용", "주요업무"),
    "competencies": ("필요역량", "필요 역량", "직무역량"),
    "requirements": ("지원자격", "응시자격"),
    "preferences": ("우대사항", "가점사항"),
    "questions": ("자기소개서", "지원서 문항"),
    "constraints": ("유의사항", "블라인드", "작성 시 유의"),
}
_BLOCK_TAGS = {
    "address",
    "article",
    "br",
    "div",
    "h1",
    "h2",
    "h3",
    "h4",
    "h5",
    "h6",
    "li",
    "p",
    "section",
    "td",
    "th",
    "tr",
}
_IGNORED_TAGS = {"script", "style", "noscript"}


class _VisibleTextParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.blocks: list[str] = []
        self._parts: list[str] = []
        self._ignored_depth = 0

    def _flush(self) -> None:
        text = " ".join(" ".join(self._parts).split())
        if text:
            self.blocks.append(text)
        self._parts.clear()

    def handle_starttag(self, tag: str, attrs) -> None:
        tag = tag.lower()
        if tag in _IGNORED_TAGS:
            self._ignored_depth += 1
            return
        if not self._ignored_depth and tag in _BLOCK_TAGS:
            self._flush()

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        if tag in _IGNORED_TAGS:
            self._ignored_depth = max(0, self._ignored_depth - 1)
            return
        if not self._ignored_depth and tag in _BLOCK_TAGS:
            self._flush()

    def handle_data(self, data: str) -> None:
        if not self._ignored_depth and data.strip():
            self._parts.append(data)

    def close(self) -> None:
        super().close()
        self._flush()


def _charset(content_type: str) -> str:
    match = re.search(r"charset=([\w-]+)", content_type, re.IGNORECASE)
    return match.group(1) if match else "utf-8"


def _posting_blocks(loaded: LoadedPosting) -> tuple[str, ...]:
    if loaded.extension == ".html":
        try:
            text = loaded.content.decode(_charset(loaded.metadata.content_type))
        except (LookupError, UnicodeDecodeError):
            text = loaded.content.decode("utf-8", errors="replace")
        parser = _VisibleTextParser()
        parser.feed(text)
        parser.close()
        return tuple(parser.blocks)
    if loaded.extension == ".docx":
        document = Document(BytesIO(loaded.content))
        blocks: list[str] = []
        for paragraph in document.paragraphs:
            text = " ".join(paragraph.text.split())
            if text:
                blocks.append(text)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for line in cell.text.splitlines():
                        text = " ".join(line.split())
                        if text:
                            blocks.append(text)
        return tuple(blocks)
    if loaded.extension == ".pdf":
        blocks = []
        for page in PdfReader(BytesIO(loaded.content)).pages:
            blocks.extend(split_text_blocks(page.extract_text() or ""))
        return tuple(blocks)
    if loaded.extension == ".txt":
        return tuple(
            split_text_blocks(loaded.content.decode("utf-8-sig", errors="replace"))
        )
    raise PostingSourceError(f"unsupported posting extension: {loaded.extension}")


def _section_name(block: str) -> str | None:
    normalized = re.sub(r"[\s:：]+", "", block)
    for section, markers in SECTION_MARKERS.items():
        if any(normalized == re.sub(r"[\s:：]+", "", marker) for marker in markers):
            return section
    return None


def _split_labeled_block(block: str) -> tuple[str | None, str]:
    for section, markers in SECTION_MARKERS.items():
        for marker in markers:
            pattern = rf"^\s*{re.escape(marker)}\s*[:：\-]\s*(.+)$"
            match = re.match(pattern, block)
            if match:
                return section, match.group(1).strip()
    return None, block


def _target_defaults(target: str) -> tuple[str, str]:
    target = " ".join(target.split())
    if not target:
        return "", ""
    match = re.match(
        r"^(.+?(?:공사|공단|은행|기금|진흥원|조정원|서비스|농협|금고|기관|회사|HUG|HF|NPS|NH|IBK))\s*(.*)$",
        target,
    )
    if match:
        return match.group(1).strip(), match.group(2).strip()
    parts = target.split(maxsplit=1)
    if len(parts) == 2:
        return parts[0].strip(), parts[1].strip()
    return "", target


@dataclass(frozen=True)
class QuestionMismatch:
    index: int
    reason: str
    posting_value: str | int | None
    draft_value: str | int | None


@dataclass(frozen=True)
class QuestionReconciliation:
    questions: tuple[Question, ...]
    mismatches: tuple[QuestionMismatch, ...]

    @property
    def matched(self) -> bool:
        return not self.mismatches


def _normalized_prompt(prompt: str) -> str:
    return "".join(prompt.split()).rstrip(".!?")


def reconcile_questions(
    posting_questions: tuple[Question, ...],
    draft_questions: tuple[Question, ...],
) -> QuestionReconciliation:
    if not posting_questions:
        return QuestionReconciliation(draft_questions, ())
    if not draft_questions:
        return QuestionReconciliation(posting_questions, ())

    mismatches: list[QuestionMismatch] = []
    if len(posting_questions) != len(draft_questions):
        mismatches.append(
            QuestionMismatch(
                0, "question_count_mismatch", len(posting_questions), len(draft_questions)
            )
        )
    for index, (posting, draft) in enumerate(
        zip(posting_questions, draft_questions), 1
    ):
        if _normalized_prompt(posting.prompt) != _normalized_prompt(draft.prompt):
            mismatches.append(
                QuestionMismatch(index, "prompt_mismatch", posting.prompt, draft.prompt)
            )
        if posting.character_limit != draft.character_limit:
            mismatches.append(
                QuestionMismatch(
                    index,
                    "character_limit_mismatch",
                    posting.character_limit,
                    draft.character_limit,
                )
            )
        if posting.minimum_character_limit != draft.minimum_character_limit:
            mismatches.append(
                QuestionMismatch(
                    index,
                    "minimum_character_limit_mismatch",
                    posting.minimum_character_limit,
                    draft.minimum_character_limit,
                )
            )
        if posting.count_mode != draft.count_mode:
            mismatches.append(
                QuestionMismatch(
                    index,
                    "character_count_mode_mismatch",
                    posting.count_mode,
                    draft.count_mode,
                )
            )
    return QuestionReconciliation(posting_questions, tuple(mismatches))




def _fallback_section_assignment(blocks: tuple[str, ...], target: str) -> dict[str, list[str]]:
    """Assign blocks to sections when no section labels are present."""
    sections = {name: [] for name in SECTION_MARKERS}
    if not blocks:
        return sections
    target_organization, target_role = _target_defaults(target)
    organization = blocks[0].strip()
    if organization.lower() in {"공고", "채용공고", "모집공고"} and target_organization:
        organization = target_organization
    sections["organization"].append(organization)
    remaining = list(blocks[1:])
    role = ""
    for idx, block in enumerate(remaining):
        if any(kw in block for kw in ("기금", "사업", "직무", "부문", "팀")):
            role = block.strip()
            remaining = remaining[idx + 1:]
            break
    if not role and target_role:
        role = target_role
    if role:
        sections["role"].append(role)
    duty_keywords = ("업무", "관리", "접수", "안내", "지원", "조사", "검토", "운영")
    competency_keywords = ("역량", "자격", "요건", "능력", "소통")
    for block in remaining:
        if any(kw in block for kw in duty_keywords):
            sections["duties"].append(block.strip())
        elif any(kw in block for kw in competency_keywords):
            sections["competencies"].append(block.strip())
        elif not sections["duties"]:
            sections["duties"].append(block.strip())
        else:
            sections["competencies"].append(block.strip())
    return sections


def parse_posting(loaded: LoadedPosting, *, target: str) -> PostingAnalysis:
    blocks = _posting_blocks(loaded)
    sections: dict[str, list[str]] = {name: [] for name in SECTION_MARKERS}
    current: str | None = None
    unclassified: list[str] = []
    for block in blocks:
        labeled_section, labeled_value = _split_labeled_block(block)
        if labeled_section:
            current = labeled_section
            if labeled_value:
                sections[current].append(labeled_value)
            continue
        section = _section_name(block)
        if section:
            current = section
            continue
        if current is None:
            continue
        sections[current].append(block)
    if not sections["organization"] and not sections["role"] and not sections["duties"] and blocks:
        sections = _fallback_section_assignment(blocks, target)

    organization_values = sections["organization"]
    role_values = sections["role"]
    if len(organization_values) > 1:
        unclassified.extend(organization_values[1:])
    if len(role_values) > 1:
        unclassified.extend(role_values[1:])

    organization = organization_values[0] if organization_values else ""
    role = role_values[0] if role_values else ""
    organization_inferred = not organization
    role_inferred = not role or role in {"근무기간", "채용인원", "근무 본부점"}
    target_organization, target_role = _target_defaults(target)
    if not organization and target_organization:
        organization = target_organization
    if (not role or role in {"근무기간", "채용인원", "근무 본부점"}) and target_role:
        role = target_role
    duties = tuple(sections["duties"])
    uncertainties = [f"unclassified: {value}" for value in unclassified]
    if not organization or organization_inferred:
        uncertainties.append("organization")
    if not role or role_inferred:
        uncertainties.append("role")
    if not duties:
        uncertainties.append("duties")

    return PostingAnalysis(
        schema_version=1,
        target=target,
        source=loaded.metadata,
        organization=organization,
        role=role,
        locations=tuple(sections["locations"]),
        duties=duties,
        competencies=tuple(sections["competencies"]),
        requirements=tuple(sections["requirements"]),
        preferences=tuple(sections["preferences"]),
        questions=tuple(extract_questions(tuple(sections["questions"]))),
        constraints=tuple(sections["constraints"]),
        uncertainties=tuple(uncertainties),
    )


def render_posting_analysis(analysis: PostingAnalysis) -> str:
    source_location = analysis.source.location
    if source_location.startswith(("https://", "http://")):
        source_display = f"[{source_location}]({source_location})"
    else:
        source_display = f"`{source_location}`"
    lines = [
        "# 채용공고 분석",
        "",
        "## 출처",
        "",
        f"- 위치: {source_display}",
        f"- 공식성: `{analysis.source.official_status}`",
        f"- SHA-256: `{analysis.source.content_sha256}`",
        "",
        "## 지원 대상",
        "",
        f"- {analysis.target}",
        "",
    ]
    sections = (
        ("담당업무", analysis.duties),
        ("필요 역량", analysis.competencies),
        ("지원자격", analysis.requirements),
        ("우대사항", analysis.preferences),
    )
    for title, values in sections:
        lines.extend([f"## {title}", ""])
        lines.extend([f"- {value}" for value in values] or ["- 없음"])
        lines.append("")
    lines.extend(["## 자기소개서 문항", ""])
    if analysis.questions:
        lines.extend(
            f"- {question.index}. {question.prompt} ({question.character_limit or '미지정'}자, "
            f"{'공백 제외' if question.count_mode == 'spaces_excluded' else '공백 포함'})"
            for question in analysis.questions
        )
    else:
        lines.append("- 없음")
    lines.extend(["", "## 제약", ""])
    lines.extend([f"- {value}" for value in analysis.constraints] or ["- 없음"])
    lines.extend(["", "## 불확실성", ""])
    lines.extend([f"- {value}" for value in analysis.uncertainties] or ["- 없음"])
    return "\n".join(lines).rstrip() + "\n"
