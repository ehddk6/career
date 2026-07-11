"""DOCX/Markdown 렌더링. 표준 비즈니스 문서 형식으로 자소서를 생성합니다."""
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .character_count import count_characters
from .models import DraftResponse, Question


BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x66, 0x66, 0x66)


def _constraint_text(question: Question, answer: str = "") -> str:
    mode = "공백 제외" if question.count_mode == "spaces_excluded" else "공백 포함"
    current = count_characters(answer, question.count_mode)
    return (
        f"제한: {question.character_limit or '미지정'}자 ({mode})"
        f" · 현재: {current}자"
    )


def _set_font(style, name: str, size: float, color=None, bold=None) -> None:
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:ascii"), name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = color
    if bold is not None:
        style.font.bold = bold


def _configure_docx(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    _set_font(normal, "Calibri", 11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    heading = document.styles["Heading 1"]
    _set_font(heading, "Calibri", 16, BLUE, True)
    heading.paragraph_format.space_before = Pt(16)
    heading.paragraph_format.space_after = Pt(8)
    heading.paragraph_format.keep_with_next = True

    title = document.styles.add_style("Application Title", WD_STYLE_TYPE.PARAGRAPH)
    _set_font(title, "Calibri", 22, DARK_BLUE, True)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(12)
    title.paragraph_format.keep_with_next = True

    constraint = document.styles.add_style("Constraint", WD_STYLE_TYPE.PARAGRAPH)
    _set_font(constraint, "Calibri", 9, MUTED)
    constraint.paragraph_format.space_before = Pt(0)
    constraint.paragraph_format.space_after = Pt(6)
    constraint.paragraph_format.keep_with_next = True


def render_draft_markdown(
    questions: list[Question], responses: list[DraftResponse]
) -> str:
    by_index = {item.question_index: item.answer for item in responses}
    chunks = ["# 자기소개서", ""]
    for question in questions:
        chunks.extend(
            [
                f"## {question.index}. {question.prompt}",
                _constraint_text(question, by_index.get(question.index, "")),
                "",
                by_index.get(question.index, ""),
                "",
            ]
        )
    return "\n".join(chunks)


def render_draft_docx(
    questions: list[Question], responses: list[DraftResponse], output: Path
) -> None:
    by_index = {item.question_index: item.answer for item in responses}
    document = Document()
    _configure_docx(document)
    title = document.add_paragraph("자기소개서", style="Application Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for question in questions:
        document.add_heading(f"{question.index}. {question.prompt}", level=1)
        document.add_paragraph(
            _constraint_text(question, by_index.get(question.index, "")),
            style="Constraint",
        )
        document.add_paragraph(by_index.get(question.index, ""))
    document.save(output)
