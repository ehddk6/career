"""DOCX, PDF, XLSX, TXT, MD 파일에서 텍스트를 추출합니다. 블록 단위로 분리하여 paragraph 리스트를 반환합니다."""
from pathlib import Path
import re

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

from .models import ExtractedDocument, SourceRecord


def split_text_blocks(text: str) -> list[str]:
    blocks = []
    for line in re.split(r"[\r\n]+", text):
        clean = line.strip()
        if not clean:
            continue
        blocks.extend(
            sentence.strip()
            for sentence in re.split(r"(?<=[.!?])\s+", clean)
            if sentence.strip()
        )
    return blocks


def _extract_docx(path: Path) -> list[str]:
    document = Document(path)
    paragraphs = []
    for paragraph in document.paragraphs:
        paragraphs.extend(split_text_blocks(paragraph.text))
    for table_index, table in enumerate(document.tables, 1):
        for row_index, row in enumerate(table.rows, 1):
            for cell_index, cell in enumerate(row.cells, 1):
                for block in split_text_blocks(cell.text):
                    paragraphs.append(
                        f"[표 {table_index} R{row_index}C{cell_index}] {block}"
                    )
    return paragraphs


def _extract_pdf(path: Path) -> list[str]:
    paragraphs = []
    for page in PdfReader(path).pages:
        text = (page.extract_text() or "").strip()
        if text:
            paragraphs.extend(split_text_blocks(text))
    return paragraphs


def _extract_xlsx(path: Path) -> list[str]:
    workbook = load_workbook(path, read_only=False, data_only=False)
    paragraphs = []
    for sheet in workbook.worksheets:
        paragraphs.append(f"[시트] {sheet.title}")
        for row in sheet.iter_rows():
            for cell in row:
                if cell.value not in (None, ""):
                    for block in split_text_blocks(str(cell.value)):
                        paragraphs.append(
                            f"[{sheet.title}!{cell.coordinate}] {block}"
                        )
    return paragraphs


def extract_path(source: SourceRecord) -> ExtractedDocument:
    if source.extension == ".docx":
        paragraphs = _extract_docx(source.path)
    elif source.extension == ".pdf":
        paragraphs = _extract_pdf(source.path)
    elif source.extension == ".xlsx":
        paragraphs = _extract_xlsx(source.path)
    elif source.extension in {".txt", ".md"}:
        paragraphs = split_text_blocks(
            source.path.read_text(encoding="utf-8-sig")
        )
    else:
        raise ValueError(f"unsupported extension: {source.extension}")

    clean = tuple(paragraph for paragraph in paragraphs if paragraph)
    return ExtractedDocument(source, "\n".join(clean), clean)
