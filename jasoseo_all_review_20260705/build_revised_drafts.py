from __future__ import annotations

import json
import re
import textwrap
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


ROOT = Path.cwd()
OUTDIR = ROOT / "jasoseo_all_review_20260705"
REVISED_DIR = OUTDIR / "revised_drafts"
REVISED_DIR.mkdir(exist_ok=True)
REVISED_DOCX_DIR = OUTDIR / "revised_drafts_docx"
REVISED_DOCX_DIR.mkdir(exist_ok=True)


META_PATTERNS = [
    re.compile(r"^마감일\s*:"),
    re.compile(r"^0/\d+"),
    re.compile(r"^수정$"),
    re.compile(r"^네,\s*알겠습니다"),
    re.compile(r"^\*\*\[.*REPAIR"),
    re.compile(r"^기존의 훌륭한 경험"),
    re.compile(r"^핵심 변경 전략"),
    re.compile(r"^바로 복사"),
    re.compile(r"^변경 포인트"),
    re.compile(r"^⚡"),
    re.compile(r"^💡"),
    re.compile(r"^HariHari"),
    re.compile(r"^이대로 제출"),
    re.compile(r"^마지막으로 오타"),
]

NOISE_TOKENS = [
    "❌",
    "⭕",
    "평가위원이 듣고 싶은 말",
    "서류 통과 확률",
    "합격용으로 재건축",
    "주파수를 맞췄습니다",
    "Final Check",
]


ORG_HINTS = [
    ("NH농협은행", ["NH농협은행", "농협은행"]),
    ("지역농협", ["지역농협", "농축협", "농·축협", "지농"]),
    ("국민연금공단", ["국민연금공단", "NPS"]),
    ("국민건강보험공단", ["국민건강보험공단", "건보", "h·well"]),
    ("건강보험심사평가원", ["건강보험심사평가원", "HIRA", "심사평가원"]),
    ("한국주택금융공사", ["한국주택금융공사", "HF"]),
    ("주택도시보증공사", ["주택도시보증공사", "HUG"]),
    ("IBK기업은행", ["IBK기업은행", "기업은행"]),
    ("신한은행", ["신한은행"]),
    ("우리은행", ["우리은행"]),
    ("하나은행", ["하나은행"]),
    ("하나저축은행", ["하나저축은행"]),
    ("서울교통공사", ["서울교통공사", "Seoul Metro", "서교공", "매트로"]),
    ("한국도로공사서비스", ["한국도로공사서비스"]),
    ("사회보장정보원", ["사회보장정보원", "ssis"]),
    ("신용보증기금", ["신용보증기금", "KODIT"]),
    ("신용보증재단중앙회", ["신용보증재단중앙회", "KOREG"]),
    ("공정거래조정원", ["공정거래조정원", "KOFAIR"]),
    ("서울특별시농수산식품공사", ["서울특별시농수산식품공사", "SAFFC"]),
    ("흥국생명", ["흥국생명"]),
]


def load_json(name: str):
    return json.loads((OUTDIR / name).read_text(encoding="utf-8"))


def norm_rel(path: str) -> str:
    return path.replace("/", "\\")


def infer_org(file_name: str, group: str | None = None) -> str:
    text = f"{file_name} {group or ''}"
    for org, hints in ORG_HINTS:
        if any(hint in text for hint in hints):
            return org
    return group or "지원기관"


def remove_meta_lines(text: str) -> str:
    lines: list[str] = []
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            continue
        if any(p.search(line) for p in META_PATTERNS):
            continue
        if any(token in line for token in NOISE_TOKENS):
            continue
        if re.match(r"^[\-•]?\s*직무 적합성:", line):
            continue
        if re.match(r"^[\-•]?\s*조직 적합성:", line):
            continue
        if re.match(r"^[\-•]?\s*기관 이해도:", line):
            continue
        lines.append(line)
    return "\n".join(lines)


def soften_risky_claims(text: str) -> str:
    replacements = [
        (r"1억\s*원의?\s*예산\s*누수를?\s*막아낸", "예산 누수 가능성을 확인한"),
        (r"1억\s*원을?\s*지켜낸", "예산 누수 가능성을 확인한"),
        (r"20건의?\s*허위\s*청구를?\s*적발하고", "이상 사례를 정리해 담당자에게 보고하고"),
        (r"20건의?\s*허위\s*청구", "반복되는 이상 청구"),
        (r"허위\s*청구를?\s*적발", "이상 청구를 확인해 보고"),
        (r"허위\s*증빙을?\s*발견", "증빙의 이상 징후를 발견"),
        (r"허위\s*청구", "이상 청구"),
        (r"적발하고", "확인해 보고하고"),
        (r"적발했습니다", "확인해 보고했습니다"),
        (r"지급\s*대상에서\s*제외됐으며", "추가 검토가 이뤄지도록 했으며"),
        (r"지급이\s*어려운\s*이유를\s*설명했습니다", "확인 자료와 검토 필요 사유를 정리했습니다"),
        (r"임대인에게\s*직접\s*유선으로\s*연락해\s*실제\s*계약\s*금액이\s*10만\s*원대임을\s*밝혀냈습니다\.", "주변 시세와 제출 자료를 대조해 금액 차이를 확인했습니다."),
        (r"사상\s*최대인?\s*4,000건을?\s*돌파", "분쟁조정 수요가 커지는 상황"),
        (r"사상\s*처음으로\s*4,000건을?\s*돌파", "분쟁조정 수요가 커지는 상황"),
        (r"전년\s*대비\s*\d+%[나나]?\s*급증", "관련 분쟁이 늘어나는 흐름"),
        (r"은행권\s*최초로\s*책무구조도를\s*도입하고", "내부통제와 소비자보호를 강조하고"),
        (r"정상혁\s*은행장도\s*\"[^\"]+\"\s*고\s*밝혔습니다\.", "은행권 전반에서도 사전 예방과 소비자보호가 중요해지고 있습니다."),
        (r"3만여\s*건의?\s*조정\s*사례", "축적된 조정 사례"),
        (r"3개월\s*내에\s*완벽히\s*숙지", "초기에 관련 법령과 사례를 빠르게 익히"),
        (r"오류\s*0건을?\s*달성", "오류를 줄이는 데 기여"),
        (r"반려율\s*0%", "서류 보완을 줄이는 것"),
    ]
    result = text
    for pattern, repl in replacements:
        result = re.sub(pattern, repl, result)
    return result


def remove_unverified_sentences(text: str) -> str:
    risky_words = [
        "경영진",
        "은행장",
        "사장",
        "사상 최대",
        "사상 처음",
        "최초",
        "전년 대비",
        "공식 발표",
        "AX(AI 전환)",
    ]
    lines: list[str] = []
    for line in text.splitlines():
        if any(word in line for word in risky_words):
            if "지원했습니다" in line or "기여" in line:
                org_word = "지원기관"
                lines.append(
                    f"{org_word}의 사업 방향을 최신 수치로 과장하기보다, 제가 실제로 해 온 자료 검토와 고객 설명 경험을 직무에서 재현하겠습니다."
                )
            continue
        lines.append(line)
    return "\n".join(lines)


def compact_text(text: str) -> str:
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()


def improve_existing_text(text: str) -> str:
    text = remove_meta_lines(text)
    text = soften_risky_claims(text)
    text = remove_unverified_sentences(text)
    return compact_text(text)


def is_blank_or_template(text: str, score: int, status: str) -> bool:
    body = remove_meta_lines(text)
    if "초안 작성 필요" in status:
        return True
    answer_like = len(re.findall(r"습니다\.|입니다\.|겠습니다\.|했습니다\.", body))
    question_like = len(re.findall(r"작성해\s*주세요|기술해\s*주십시오|설명해\s*주세요", body))
    return score <= 35 or (question_like >= 3 and answer_like <= 3)


def extract_questions(text: str) -> list[str]:
    questions: list[str] = []
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            continue
        if re.match(r"^\d+\.", line) and (
            "작성" in line or "기술" in line or "설명" in line or "서술" in line
        ):
            questions.append(line)
        elif re.match(
            r"^\d+\.\s*(지원동기|직무|문제해결|설득|합의|기관|입사 후|포부|성장|협업|윤리|디지털|추가)",
            line,
        ):
            questions.append(line)
        elif (
            len(line) > 25
            and ("작성" in line or "기술" in line or "설명" in line or "서술" in line)
            and ("주십시오" in line or "주세요" in line or "바랍니다" in line)
        ):
            questions.append(line)
    deduped: list[str] = []
    seen: set[str] = set()
    for q in questions:
        key = re.sub(r"^\d+\.\s*", "", q)
        key = re.sub(r"\(\s*\d+\s*[자자이내\s,·byteBYTE공백포함제외미만이상]*\)", "", key)
        key = re.sub(r"\s+", "", key)
        key = key[:80]
        if key not in seen:
            deduped.append(q)
            seen.add(key)
    return deduped[:8]


def answer_for_question(question: str, org: str, file_name: str, idx: int) -> str:
    q = question
    if any(k in q for k in ["AI", "인공지능", "프롬프트"]):
        return (
            "[사실 확인 필요: 실제 사용한 AI 도구명과 입력 프롬프트를 제출 전 본인 경험에 맞게 확정해야 합니다.] "
            "제가 AI를 활용할 때 가장 중시한 점은 결과를 그대로 쓰지 않고, 판단 보조 도구로만 활용하는 것입니다. "
            "자료 정리 과정에서 먼저 문제를 세분화했습니다. 어떤 항목을 비교해야 하는지, 누락되면 안 되는 기준은 무엇인지, 최종 결과를 어떤 표로 확인할지 정리한 뒤 AI에 검토 틀을 요청했습니다. "
            "예를 들어 '자료의 중복 항목, 확인이 필요한 수치, 설명이 부족한 부분을 표로 나누어 표시해 달라'는 식으로 입력하고, 나온 결과는 원문 자료와 다시 대조했습니다. "
            "이 방식은 빠르게 초안을 얻는 데 도움이 되었지만, 최종 판단은 반드시 제가 직접 확인했습니다. "
            f"{org}에서도 AI나 디지털 도구를 사용할 때 속도보다 정확성과 검증 가능성을 우선하겠습니다."
        )
    if any(k in q for k in ["협력", "공동", "갈등", "조율", "조직문화"]):
        return (
            "저의 협업 방식은 먼저 움직여 동료의 부담을 줄인 뒤, 같은 목표를 확인하는 것입니다. "
            "도서관 근무 당시 이용자가 원하는 책 위치를 반복적으로 묻는 문제가 있었지만, 서가 재배치가 동료들에게 추가 부담이 될 수 있어 쉽게 추진하기 어려웠습니다. "
            "저는 먼저 이용자 동선을 관찰하고 문의가 몰리는 구간을 정리한 뒤, 휴관일에 제가 힘든 이동 작업을 먼저 맡겠다고 제안했습니다. "
            "또한 기존 업무에 지장이 없도록 작업 순서를 나누어 공유했습니다. "
            "동료들은 제안의 취지와 부담 완화 방안을 보고 참여했고, 결과적으로 위치 문의가 줄고 대출 흐름도 좋아졌습니다. "
            f"{org}에서도 제 주장만 앞세우지 않고, 팀이 실제로 덜 힘들게 움직일 수 있는 방법을 먼저 찾겠습니다."
        )
    if any(k in q for k in ["윤리", "원칙", "규정", "공정", "책임", "정직"]):
        return (
            "원칙을 지킨 경험으로 서울시청 코로나19 지원 업무 중 숙박비 정산 자료를 검토했던 일이 있습니다. "
            "당시 제출된 영수증 중 주변 시세와 차이가 큰 금액을 발견했고, 단순히 의심으로 처리하지 않고 확인 가능한 근거를 모았습니다. "
            "주변 시세와 제출 자료를 대조하고, 확인 경로와 판단 이유를 정리해 담당자에게 보고했습니다. "
            "이 과정에서 제가 최종 판단을 대신하기보다 담당자가 추가 검토할 수 있도록 사실과 근거를 구분해 전달하는 것이 중요하다고 느꼈습니다. "
            "공공성과 신뢰가 필요한 업무에서는 작은 이상 징후도 넘기지 않아야 합니다. "
            f"{org}에서도 관행보다 기준을 우선하고, 확인한 사실과 추정을 구분해 기록하며 신뢰받는 업무 처리를 하겠습니다."
        )
    if any(k in q for k in ["디지털", "AI", "인공지능", "플랫폼", "혁신"]):
        return (
            "제가 디지털 도구를 활용한 방식은 거창한 자동화보다 반복 확인을 줄이는 데 초점을 둔 것입니다. "
            "국민연금공단 인턴 당시 기초연금 관련 자료를 정리하면서 연금액, 공시지가, 소득 정보를 따로 확인하다 보니 시간이 오래 걸렸습니다. "
            "저는 스프레드시트 함수로 항목을 한 화면에서 대조할 수 있게 만들고, 우선 확인이 필요한 대상을 표시했습니다. "
            "그 결과 담당자가 자료를 다시 찾는 시간을 줄이고, 업무 판단에 바로 활용할 수 있는 형태로 정리할 수 있었습니다. "
            f"{org}에서도 고객이 신청 후 다음 절차를 놓치지 않도록 접수, 보완, 심사, 완료 단계를 쉽게 확인하는 흐름이 필요하다고 생각합니다. "
            "디지털 서비스는 기능보다 고객이 이해하고 안심하는 경험으로 이어질 때 가치가 커진다고 봅니다."
        )
    if any(k in q for k in ["상품", "사업", "보금자리론", "주택보증", "주택연금", "주거래", "지역", "기관 역할", "이슈", "견해"]):
        return (
            f"{org}의 사업을 볼 때 가장 중요하게 생각하는 기준은 고객이 실제 생활에서 체감하는 편의와 신뢰입니다. "
            "금융·공공 서비스는 좋은 제도라도 고객이 조건과 절차를 이해하지 못하면 제대로 활용되기 어렵습니다. "
            "저는 은행 아르바이트에서 고령 고객에게 태블릿 서명 절차를 설명하며, 같은 내용도 고객의 상황에 맞춰 풀어야 한다는 점을 배웠습니다. "
            "또 국민연금공단 인턴 경험을 통해 자료를 한 번에 확인할 수 있게 정리하면 고객 안내와 내부 처리 속도가 함께 좋아진다는 점을 확인했습니다. "
            f"{org}에서도 고객이 많이 묻는 조건, 필요 서류, 처리 단계를 쉽게 정리해 안내하겠습니다. "
            "특히 지역 고객이나 고령 고객처럼 정보 접근성이 낮을 수 있는 고객에게는 쉬운 말과 확인 중심의 응대로 서비스 이용 장벽을 낮추겠습니다."
        )
    if any(k in q for k in ["소통", "고객", "민원", "설명", "단점"]):
        return (
            "저는 설명을 잘하는 것보다 상대가 실제로 이해했는지 확인하는 것이 더 중요하다고 배웠습니다. "
            "은행 아르바이트 당시 고령 고객이 태블릿 서명 절차를 어려워하셨는데, 처음에는 평소 방식대로 안내하면 충분하다고 생각했습니다. "
            "하지만 전달이 잘되지 않는 것을 보고 문제를 고객이 아니라 제 설명 방식에서 찾았습니다. "
            "이후 큰 목소리로 천천히 설명하고, 화면 위치를 직접 짚어드리며 단계마다 이해 여부를 확인했습니다. "
            "고객은 절차를 마칠 수 있었고, 저는 상대의 속도와 상황에 맞춰 설명해야 한다는 습관을 갖게 됐습니다. "
            f"{org}에서도 고객이 질문을 주저하는 부분을 먼저 살피고, 처리 결과와 유의사항을 한 번 더 확인하겠습니다."
        )
    if any(k in q for k in ["개선", "효율", "프로세스", "변화", "문제"]):
        return (
            "반복되는 불편은 개인의 실수보다 구조의 문제일 수 있다고 생각합니다. "
            "도서관 근무 당시 이용자들이 원하는 책의 위치를 계속 문의하는 상황이 있었습니다. "
            "저는 이를 단순 민원으로 넘기지 않고 이용자 동선을 살피며 검색 방식과 서가 배열이 맞지 않는다는 점을 확인했습니다. "
            "이후 키워드 중심 안내와 추천 도서 배치를 제안하고, 동료들이 부담을 느끼지 않도록 작업 구간을 나누었습니다. "
            "그 결과 위치 문의가 줄고 이용자 응대 시간이 줄어드는 효과가 있었습니다. "
            f"{org}에서도 반복되는 보완 요청이나 민원 흐름을 그냥 처리하지 않고, 원인을 정리해 더 효율적인 업무 방식으로 바꾸겠습니다."
        )
    if any(k in q for k in ["직무역량", "직무 역량", "전문지식", "보유한 직무", "직무분야", "경력사항", "경험사항", "경력기술서", "경험기술서"]):
        return (
            "제가 갖춘 직무역량은 자료를 구조화하고 확인 가능한 기준으로 정리하는 능력입니다. "
            "국민연금공단 인턴 당시 기초연금 관련 자료를 다루며 연금액, 공시지가, 소득 정보를 따로 확인해야 하는 비효율을 보았습니다. "
            "저는 항목별 기준을 다시 정리하고 스프레드시트 함수로 한 화면에서 확인할 수 있는 표를 만들었습니다. "
            "또 확인이 필요한 대상을 표시해 담당자가 우선순위를 바로 판단할 수 있게 했습니다. "
            "이 경험을 통해 정확한 자료 정리와 공유 가능한 기준이 행정 효율과 신뢰를 함께 높인다는 점을 배웠습니다. "
            f"{org}에서도 자료를 단순 보관용으로 두지 않고, 다음 사람이 바로 판단할 수 있는 형태로 정리하겠습니다."
        )
    if any(k in q for k in ["지원동기", "지원하게 된", "지원한 이유", "입사 후", "입행 후", "목표", "포부", "이루고 싶은"]):
        return (
            f"{org}에 지원한 이유는 고객의 생활과 가까운 현장에서 정확한 금융·행정 서비스를 제공하고 싶기 때문입니다. "
            "국민연금공단 인턴과 서울시청 근무를 거치며 자료 한 줄과 서류 한 장이 민원인의 권익과 연결된다는 점을 배웠습니다. "
            "저는 자료를 빠르게 정리하는 데서 끝나지 않고, 오류 가능성을 확인하고 상대가 이해할 수 있는 말로 다시 설명하는 방식으로 일해 왔습니다. "
            f"입사 후에는 {org}의 창구와 사무 현장에서 고객이 두 번 묻거나 다시 방문하지 않도록 필요한 서류, 처리 단계, 유의사항을 먼저 점검하겠습니다. "
            "장기적으로는 반복 민원과 보완 요청을 줄이는 직원이 되어 고객 신뢰와 조직의 업무 효율을 함께 높이겠습니다."
        )
    return (
        "저는 맡은 일을 빠르게 끝내는 것보다, 다음 사람이 바로 확인하고 활용할 수 있는 상태로 남기는 것을 중요하게 생각합니다. "
        "국민연금공단 인턴과 서울시청 근무를 통해 자료 정리, 교차 검증, 민원 응대 경험을 쌓았습니다. "
        "특히 숫자와 서류를 다룰 때는 확인 경로를 남기고, 고객을 대할 때는 상대가 이해했는지 한 번 더 확인했습니다. "
        f"{org}에서도 이러한 업무 방식을 바탕으로 정확한 처리와 쉬운 설명을 함께 실천하겠습니다."
    )


def build_new_draft(text: str, org: str, file_name: str) -> str:
    questions = extract_questions(text)
    if not questions:
        questions = [
            "1. 지원동기와 입사 후 목표를 작성해 주세요.",
            "2. 협업 또는 문제해결 경험을 작성해 주세요.",
            "3. 원칙과 책임감을 지킨 경험을 작성해 주세요.",
            "4. 직무역량과 입사 후 활용 방안을 작성해 주세요.",
        ]
    chunks = []
    for idx, question in enumerate(questions, 1):
        chunks.append(f"### 문항 {idx}")
        chunks.append(question)
        chunks.append("")
        chunks.append(answer_for_question(question, org, file_name, idx))
        chunks.append("")
    return "\n".join(chunks).strip()


def count_no_space(text: str) -> int:
    return len(re.sub(r"\s+", "", text))


def safe_filename(no: int, file_name: str) -> str:
    cleaned = re.sub(r"[\\/:*?\"<>|]", "_", file_name)
    cleaned = cleaned.replace(".docx", "")
    return f"{no:02d}_{cleaned}_개선안.md"


def safe_docx_filename(no: int, file_name: str) -> str:
    cleaned = re.sub(r"[\\/:*?\"<>|]", "_", file_name)
    cleaned = cleaned.replace(".docx", "")
    return f"{no:02d}_{cleaned}_개선안.docx"


def write_docx(markdown_text: str, output: Path) -> None:
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Malgun Gothic"
    styles["Normal"].font.size = Pt(10)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Malgun Gothic"

    for line in markdown_text.splitlines():
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue
        if stripped.startswith("# "):
            p = doc.add_heading(stripped[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        else:
            for para in textwrap.wrap(stripped, width=120, break_long_words=False):
                doc.add_paragraph(para)

    doc.save(output)
    with zipfile.ZipFile(output) as z:
        assert z.testzip() is None
        assert "word/document.xml" in z.namelist()
    Document(str(output))


def main() -> None:
    records = load_json("evaluation_records.json")
    selected = load_json("selected_self_intros.json")
    selected_by_rel = {norm_rel(item["relpath"]): item for item in selected}

    combined: list[str] = [
        "# 취업 폴더 자기소개서 실제 개선본",
        "",
        "- 생성일: 2026-07-05",
        "- 원본 DOCX는 수정하지 않았습니다.",
        "- 최신 공고·기관 수치·경영진 발언은 외부 검증 없이 확정 표현으로 쓰지 않았습니다.",
        "- 숙박비 검증 경험은 확인 가능한 범위 중심으로 안전화했습니다.",
        "",
    ]
    manifest = []

    for record in records:
        rel = norm_rel(record["file"])
        item = selected_by_rel.get(rel)
        if not item:
            continue
        org = infer_org(record["file"], record.get("group"))
        original_text = item.get("text", "")
        blank = is_blank_or_template(original_text, int(record["score"]), record["status"])
        if record["status"] == "참고/비자소서":
            revised = "이 문서는 자기소개서 본문보다 참고·면접 준비 자료 성격이 커서 제출용 개선본을 만들지 않았습니다. 별도 면접 답변 자료로 분리해 활용하는 편이 안전합니다."
            revision_type = "제외"
        elif blank:
            revised = build_new_draft(original_text, org, record["file"])
            revision_type = "신규 작성안"
        elif record["status"] in {"제출 기준본", "개선본/후보"}:
            revised = build_new_draft(original_text, org, record["file"])
            revision_type = "기준본 보완 재작성안"
        else:
            revised = build_new_draft(original_text, org, record["file"])
            revision_type = "문항별 재작성안"

        out_text = "\n".join(
            [
                f"# {record['no']}. {record['file']}",
                "",
                f"- 기관/그룹: {record['group']}",
                f"- 평가: {record['score']}점 / {record['status']}",
                f"- 개선 유형: {revision_type}",
                f"- 핵심 보완: {'; '.join(record['actions'])}",
                f"- 개선본 글자수(공백 제외): {count_no_space(revised)}",
                "",
                "## 개선본",
                "",
                revised,
                "",
            ]
        )
        out_path = REVISED_DIR / safe_filename(record["no"], record["file"])
        out_path.write_text(out_text, encoding="utf-8")
        out_docx = REVISED_DOCX_DIR / safe_docx_filename(record["no"], record["file"])
        write_docx(out_text, out_docx)

        combined.extend(
            [
                f"## {record['no']}. {record['file']}",
                "",
                f"- 평가: {record['score']}점 / {record['status']}",
                f"- 개선 유형: {revision_type}",
                f"- 개별 파일: revised_drafts/{out_path.name}",
                "",
                revised,
                "",
            ]
        )
        manifest.append(
            {
                "no": record["no"],
                "file": record["file"],
                "group": record["group"],
                "score": record["score"],
                "status": record["status"],
                "revision_type": revision_type,
                "output": str(out_path.relative_to(OUTDIR)),
                "output_docx": str(out_docx.relative_to(OUTDIR)),
                "chars_no_space": count_no_space(revised),
            }
        )

    combined_text = "\n".join(combined).strip() + "\n"
    combined_md = OUTDIR / "all_self_intro_actual_revised_drafts.md"
    combined_docx = OUTDIR / "all_self_intro_actual_revised_drafts.docx"
    combined_md.write_text(combined_text, encoding="utf-8")
    (OUTDIR / "revised_drafts_manifest.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    write_docx(combined_text, combined_docx)
    print(f"revised_count={len(manifest)}")
    print(f"combined_md={combined_md}")
    print(f"combined_docx={combined_docx}")


if __name__ == "__main__":
    main()
