from __future__ import annotations

import json
import os
import re
import shutil
import subprocess
import sys
from dataclasses import asdict
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from career_pipeline.copyeditor_adapter import copyedit_responses
from career_pipeline.models import DraftResponse, Question
from career_pipeline.patina_adapter import humanize_text, score_text
from career_pipeline.quality import (
    STRICT_MIN_ANSWER_SCORE,
    STRICT_MIN_AVERAGE_SCORE,
    score_answer_quality,
    validate_answer_quality,
)

from process_legacy_self_intros import JOB_TERMS, RISK_TERMS, infer_org


ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "final_gate_20260705"
COPYEDITOR_REPORT = OUT_DIR / "09_copyeditor_report_all_20260705.json"
PATINA_REPORT = OUT_DIR / "09_patina_report_all_20260705.json"
QUALITY_REPORT = OUT_DIR / "11_final_gate_quality_audit_all_20260705.json"
SUMMARY_MD = OUT_DIR / "최종게이트_전체보고서_20260705.md"
FINAL_MD_DIR = OUT_DIR / "final_gate_drafts"
FINAL_DOCX_DIR = OUT_DIR / "final_gate_drafts_docx"
FINAL_COMBINED_DOCX = OUT_DIR / "all_self_intro_final_gate_complete_20260705.docx"
STATE_JSON = OUT_DIR / "run_state.json"

USER_NPM_BIN = Path.home() / "AppData" / "Roaming" / "npm"
if USER_NPM_BIN.exists():
    os.environ["PATH"] = str(USER_NPM_BIN) + os.pathsep + os.environ.get("PATH", "")

ORG_HINTS = [
    ("NH농협은행", ("NH농협은행", "농협은행")),
    ("지역농협", ("지역농협", "농축협", "농·축협", "지농")),
    ("국민연금공단", ("국민연금공단", "NPS")),
    ("국민건강보험공단", ("국민건강보험공단", "건보", "h·well")),
    ("건강보험심사평가원", ("건강보험심사평가원", "HIRA", "심사평가원")),
    ("한국주택금융공사", ("한국주택금융공사", "HF")),
    ("주택도시보증공사", ("주택도시보증공사", "HUG")),
    ("IBK기업은행", ("IBK기업은행", "기업은행")),
    ("신한은행", ("신한은행",)),
    ("우리은행", ("우리은행",)),
    ("하나저축은행", ("하나저축은행",)),
    ("하나은행", ("하나은행",)),
    ("서울교통공사", ("서울교통공사", "Seoul Metro", "서교공", "매트로")),
    ("한국도로공사서비스", ("한국도로공사서비스",)),
    ("사회보장정보원", ("사회보장정보원", "ssis", "서시서")),
    ("신용보증기금", ("신용보증기금", "KODIT")),
    ("신용보증재단중앙회", ("신용보증재단중앙회", "KOREG")),
    ("한국공정거래조정원", ("공정거래조정원", "KOFAIR")),
    ("서울특별시농수산식품공사", ("서울특별시농수산식품공사", "SAFFC")),
    ("흥국생명", ("흥국생명",)),
    ("해양박물관", ("해양박물관",)),
]

GATE_JOB_TERMS = {
    **JOB_TERMS,
    "NH농협은행": ("고객 상담", "금융상품 설명", "서류 검토", "창구 업무", "신뢰"),
    "지역농협": ("조합원", "지역 고객", "창구 업무", "서류 검토", "생활금융"),
    "건강보험심사평가원": ("심사", "평가", "자료 검토", "기록 정확성", "공정성"),
    "한국주택금융공사": ("주택금융", "상담", "서류 검토", "공공성", "고객 설명"),
    "주택도시보증공사": ("보증", "심사", "리스크", "서류 검토", "주거 안정"),
    "신한은행": ("고객 신뢰", "금융소비자 보호", "상담", "내부통제", "서류"),
    "하나저축은행": ("서민금융", "기업금융", "서류 검토", "리스크", "상담"),
    "서울교통공사": ("시민 안전", "고객 응대", "현장 규정", "상황 대응", "서비스"),
    "사회보장정보원": ("복지", "정보시스템", "데이터 정확성", "사용자 안내", "공공"),
    "신용보증재단중앙회": ("소상공인", "보증", "지역", "자료 검토", "재단"),
    "흥국생명": ("보험", "고객", "상품 설명", "영업기획", "자료"),
    "해양박물관": ("관람객 안내", "전시 운영", "교육", "자료 정리", "안전"),
}


def infer_gate_org(file_name: str, text: str = "") -> str:
    for org, hints in ORG_HINTS:
        if any(hint in file_name for hint in hints):
            return org
    for org, hints in ORG_HINTS:
        if any(hint in text for hint in hints):
            return org
    return infer_org(file_name, text)


def collect_files() -> list[Path]:
    folders = [
        ROOT / "submission_ready_drafts",
        ROOT / "legacy_submission_ready_drafts",
        ROOT / "supplemental_submission_ready_drafts",
    ]
    files: list[Path] = []
    for folder in folders:
        files.extend(sorted(folder.glob("*.md")))
    return files


def body_text(markdown: str) -> str:
    return markdown.split("## 자기소개서 본문", 1)[-1].strip()


def parse_questions(markdown: str) -> tuple[list[Question], list[DraftResponse]]:
    body = body_text(markdown)
    parts = re.split(r"(?m)^###\s+문항\s+(\d+)\s*$", body)
    questions: list[Question] = []
    responses: list[DraftResponse] = []
    if len(parts) == 1:
        clean = re.sub(r"^#.*$", "", body, flags=re.MULTILINE).strip()
        questions.append(Question(1, "자기소개서 본문", None))
        responses.append(DraftResponse(1, clean, ("final_gate",)))
        return questions, responses

    for i in range(1, len(parts), 2):
        index = int(parts[i])
        section = parts[i + 1].strip()
        lines = [line.strip() for line in section.splitlines() if line.strip()]
        if not lines:
            continue
        prompt = lines[0]
        answer = "\n".join(lines[1:]).strip()
        limit_values = [int(item) for item in re.findall(r"(\d{2,4})\s*자", prompt)]
        limit_values += [
            int(item)
            for item in re.findall(r"(\d{2,4})\s*byte", prompt, re.IGNORECASE)
        ]
        limit_values = [value for value in limit_values if value >= 150]
        limit = max(limit_values) if limit_values else None
        questions.append(Question(index, prompt, limit))
        responses.append(DraftResponse(index, answer, ("final_gate",)))
    return questions, responses


def replace_answers(markdown: str, responses: list[DraftResponse]) -> str:
    response_by_index = {item.question_index: item.answer for item in responses}
    parts = re.split(r"(?m)^###\s+문항\s+(\d+)\s*$", markdown)
    if len(parts) == 1:
        return markdown
    rebuilt = [parts[0]]
    for i in range(1, len(parts), 2):
        index = int(parts[i])
        section = parts[i + 1].strip()
        lines = [line.strip() for line in section.splitlines() if line.strip()]
        prompt = lines[0] if lines else "자기소개서 본문"
        rebuilt.append(f"### 문항 {index}\n{prompt}\n\n{response_by_index.get(index, '')}\n\n")
    return "".join(rebuilt).strip() + "\n"


def safe_name(path: Path) -> str:
    value = path.stem
    value = re.sub(r'[<>:"/\\|?*]+', "_", value)
    if len(value) <= 120:
        return value
    return value[:120].rstrip()


def write_docx(path: Path, markdown: str) -> None:
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10.5)
    for line in markdown.splitlines():
        if line.startswith("# "):
            heading = document.add_heading(line[2:], level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("## "):
            document.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            document.add_heading(line[4:], level=3)
        elif line.startswith("- "):
            document.add_paragraph(line[2:])
        elif line.strip():
            document.add_paragraph(line)
    document.save(path)


def load_json(path: Path) -> list[dict]:
    if not path.exists():
        return []
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except json.JSONDecodeError:
        return []
    return data if isinstance(data, list) else []


def save_json(path: Path, data: list[dict] | dict) -> None:
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def patina_doctor() -> dict:
    try:
        patina_command = shutil.which("patina.cmd") or shutil.which("patina") or "patina"
        completed = subprocess.run(
            [patina_command, "doctor", "--json"],
            text=True,
            encoding="utf-8",
            errors="replace",
            capture_output=True,
            timeout=90,
        )
    except (OSError, subprocess.SubprocessError) as error:
        return {"ok": False, "message": str(error)}
    payload: dict = {"returncode": completed.returncode}
    try:
        payload["json"] = json.loads(completed.stdout)
    except json.JSONDecodeError:
        payload["stdout"] = completed.stdout
    payload["stderr"] = completed.stderr
    payload["ok"] = completed.returncode == 0
    return payload


def quality_audit(
    relpath: str,
    org: str,
    questions: list[Question],
    responses: list[DraftResponse],
) -> dict:
    job_terms = GATE_JOB_TERMS.get(org, GATE_JOB_TERMS["지원기관"])
    rows: list[dict] = []
    for question in questions:
        response = next(item for item in responses if item.question_index == question.index)
        peer_answers = tuple(item.answer for item in responses if item.question_index != question.index)
        score = score_answer_quality(
            question,
            response.answer,
            org,
            job_terms=job_terms,
            peer_answers=peer_answers,
        )
        rows.append(
            {
                "question_index": question.index,
                "score": asdict(score),
                "chars_no_space": len(re.sub(r"\s+", "", response.answer)),
            }
        )
    validation = validate_answer_quality(
        questions,
        responses,
        org,
        job_terms=job_terms,
        minimum_score=STRICT_MIN_ANSWER_SCORE,
        average_minimum_score=STRICT_MIN_AVERAGE_SCORE,
    )
    risk_hits = sorted({term for term in RISK_TERMS if term in "\n".join(item.answer for item in responses)})
    average = round(sum(row["score"]["total"] for row in rows) / len(rows), 1) if rows else 0
    recommendation = (
        "제출권장"
        if average >= 90 and not validation and not risk_hits
        else "보완필요"
    )
    return {
        "file": relpath,
        "organization": org,
        "question_count": len(questions),
        "average_score": average,
        "recommendation": recommendation,
        "validation_issues": [asdict(issue) for issue in validation],
        "risk_terms_remaining": risk_hits,
        "questions": rows,
    }


def run_gate_for_file(path: Path, *, prefer_existing_final: bool = False) -> tuple[dict, dict, dict]:
    markdown_path = path
    if prefer_existing_final:
        existing_final = FINAL_MD_DIR / f"{safe_name(path)}_final_gate.md"
        if existing_final.exists():
            markdown_path = existing_final
    markdown = markdown_path.read_text(encoding="utf-8")
    org = infer_gate_org(path.name, markdown)
    questions, responses = parse_questions(markdown)
    job_terms = GATE_JOB_TERMS.get(org, GATE_JOB_TERMS["지원기관"])
    copy_rows: list[dict] = []
    patina_rows: list[dict] = []
    final_responses: list[DraftResponse] = []
    copyedited_responses, copy_report_items = copyedit_responses(
        responses,
        target_org=org,
        job_terms=job_terms,
        timeout_ms=180_000,
    )
    copyedited_by_index = {item.question_index: item for item in copyedited_responses}
    copy_report_by_index = {
        int(item["question_index"]): item for item in copy_report_items
    }
    for response in responses:
        question = next(item for item in questions if item.index == response.question_index)
        copy_report_item = copy_report_by_index[response.question_index]
        protected_terms = tuple(
            term for term in (org, *job_terms) if term and term in response.answer
        )
        copy_applied = copy_report_item["status"] == "copyedited"
        after_copy = copyedited_by_index[response.question_index].answer
        score_before = score_text(
            after_copy,
            threshold=30,
            backend="codex-cli",
            timeout_ms=90_000,
            max_retries=0,
        )
        patina_attempted = score_before.score is not None
        patina_applied = False
        selected_variant = "copyedited" if copy_applied else "original"
        selected_text = after_copy
        humanize_status = "not_needed"
        humanize_message = ""
        if score_before.score is not None and score_before.score > 30:
            humanized = humanize_text(
                after_copy,
                character_limit=question.character_limit,
                count_mode=question.count_mode,
                backend="codex-cli",
                timeout_ms=180_000,
                profile="formal",
                tone="professional",
                voice_sample=None,
                protected_terms=protected_terms,
                max_retries=0,
            )
            humanize_status = humanized.status
            humanize_message = humanized.message
            if humanized.status in {"humanized", "humanized_compacted"}:
                score_after = score_text(
                    humanized.text,
                    threshold=30,
                    backend="codex-cli",
                    timeout_ms=90_000,
                    max_retries=0,
                )
                if score_after.score is not None and score_after.score <= 30:
                    selected_text = humanized.text
                    patina_applied = True
                    selected_variant = "patina"
                    score_before = score_after
        final_responses.append(
            DraftResponse(
                response.question_index,
                selected_text,
                response.evidence_paths,
                response.experience_refs,
                response.research_refs,
            )
        )
        copy_rows.append(
            {
                "question_index": response.question_index,
                "copyeditor_attempted": True,
                "copyeditor_applied": copy_applied,
                "status": copy_report_item["status"],
                "message": copy_report_item["message"],
                "change_ratio": copy_report_item["change_ratio"],
                "applied_rules": copy_report_item["applied_rules"],
            }
        )
        patina_rows.append(
            {
                "question_index": response.question_index,
                "patina_attempted": patina_attempted,
                "patina_applied": patina_applied,
                "score_status": score_before.status,
                "selected_ai_score": score_before.score,
                "ai_score_gate": "passed"
                if score_before.score is not None and score_before.score <= 30
                else "blocked",
                "selected_variant": selected_variant,
                "patina_status": humanize_status,
                "message": humanize_message or score_before.message,
            }
        )
    output_markdown = replace_answers(markdown, final_responses)
    rel = str(path.relative_to(ROOT))
    output_path = FINAL_MD_DIR / f"{safe_name(path)}_final_gate.md"
    docx_path = FINAL_DOCX_DIR / f"{safe_name(path)}_final_gate.docx"
    output_path.write_text(output_markdown, encoding="utf-8")
    write_docx(docx_path, output_markdown)
    quality = quality_audit(rel, org, questions, final_responses)
    quality["final_file"] = str(output_path.relative_to(ROOT))
    quality["final_docx_file"] = str(docx_path.relative_to(ROOT))
    copy_report = {
        "file": rel,
        "organization": org,
        "question_count": len(questions),
        "copyeditor_attempted": True,
        "copyeditor_applied_count": sum(1 for row in copy_rows if row["copyeditor_applied"]),
        "items": copy_rows,
    }
    patina_report = {
        "file": rel,
        "organization": org,
        "question_count": len(questions),
        "patina_attempted": True,
        "patina_applied_count": sum(1 for row in patina_rows if row["patina_applied"]),
        "items": patina_rows,
    }
    return copy_report, patina_report, quality


def write_combined_docx(quality_rows: list[dict]) -> None:
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10.5)
    heading = document.add_heading("전체 자기소개서 최종 게이트 통합본", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(f"대상: {len(quality_rows)}개")
    document.add_paragraph("게이트: im-ai-copyeditor 안전 검증, Patina 점수/필요시 재작성, 제출품질 재감사")
    for index, row in enumerate(quality_rows, 1):
        document.add_page_break()
        document.add_heading(f"{index:02d}. {row['file']}", level=1)
        text = (ROOT / row["final_file"]).read_text(encoding="utf-8")
        for line in text.splitlines():
            if line.startswith("# "):
                document.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                document.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                document.add_heading(line[4:], level=3)
            elif line.startswith("- "):
                document.add_paragraph(line[2:])
            elif line.strip():
                document.add_paragraph(line)
    document.save(FINAL_COMBINED_DOCX)


def write_summary(copy_reports: list[dict], patina_reports: list[dict], quality_rows: list[dict]) -> None:
    total_questions = sum(row["question_count"] for row in quality_rows)
    copy_applied = sum(row["copyeditor_applied_count"] for row in copy_reports)
    patina_applied = sum(row["patina_applied_count"] for row in patina_reports)
    copy_items = [item for report in copy_reports for item in report["items"]]
    patina_items = [item for report in patina_reports for item in report["items"]]
    copy_backend_ok = sum(
        1 for item in copy_items if item["status"] != "fallback_backend_error"
    )
    patina_scored = sum(
        1 for item in patina_items if item["selected_ai_score"] is not None
    )
    patina_blocked = sum(
        1 for item in patina_items if item["selected_ai_score"] is None
    )
    patina_scores = [
        item["selected_ai_score"]
        for report in patina_reports
        for item in report["items"]
        if item["selected_ai_score"] is not None
    ]
    max_score = max(patina_scores) if patina_scores else None
    pass_count = sum(1 for row in quality_rows if row["recommendation"] == "제출권장")
    lines = [
        "# 전체 자기소개서 최종 게이트 보고서",
        "",
        "- 실행일: 2026-07-05",
        f"- 대상 문서: {len(quality_rows)}개",
        f"- 대상 문항: {total_questions}개",
        f"- 제출권장 재감사: {pass_count}/{len(quality_rows)}개",
        f"- im-ai-copyeditor 백엔드 완료 문항: {copy_backend_ok}/{total_questions}개",
        f"- im-ai-copyeditor 적용 문항: {copy_applied}/{total_questions}개",
        f"- Patina 점수 측정 문항: {patina_scored}/{total_questions}개",
        f"- Patina 점수 미측정 문항: {patina_blocked}/{total_questions}개",
        f"- Patina 적용 문항: {patina_applied}/{total_questions}개",
        f"- Patina 최고 점수: {max_score if max_score is not None else '측정 실패'}",
        f"- 최종 통합 DOCX: `{FINAL_COMBINED_DOCX.name}`",
        "",
        "## 판정",
        "",
    ]
    if (
        pass_count == len(quality_rows)
        and copy_backend_ok == total_questions
        and patina_scored == total_questions
        and max_score is not None
        and max_score <= 30
    ):
        lines.append("- 결론: 전체 문서가 최종 게이트를 통과했습니다.")
    else:
        lines.append("- 결론: 일부 문서 또는 Patina 점수 측정에 추가 확인이 필요합니다.")
    lines.extend(
        [
            "",
            "## 문서별 결과",
            "",
            "| 번호 | 문서 | 제출평가 | 평균 | Patina 최고 | copyeditor 적용 | Patina 적용 |",
            "|---:|---|---|---:|---:|---:|---:|",
        ]
    )
    patina_by_file = {row["file"]: row for row in patina_reports}
    copy_by_file = {row["file"]: row for row in copy_reports}
    for index, row in enumerate(quality_rows, 1):
        patina_items = patina_by_file[row["file"]]["items"]
        scores = [item["selected_ai_score"] for item in patina_items if item["selected_ai_score"] is not None]
        file_max = max(scores) if scores else None
        copy_count = copy_by_file[row["file"]]["copyeditor_applied_count"]
        patina_count = patina_by_file[row["file"]]["patina_applied_count"]
        lines.append(
            f"| {index} | {row['file']} | {row['recommendation']} | {row['average_score']} | "
            f"{file_max if file_max is not None else '측정 실패'} | {copy_count} | {patina_count} |"
        )
    SUMMARY_MD.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> int:
    OUT_DIR.mkdir(exist_ok=True)
    FINAL_MD_DIR.mkdir(exist_ok=True)
    FINAL_DOCX_DIR.mkdir(exist_ok=True)
    files = collect_files()
    doctor = patina_doctor()
    state = {"total_files": len(files), "patina_doctor": doctor}
    save_json(STATE_JSON, state)
    copy_reports = load_json(COPYEDITOR_REPORT)
    patina_reports = load_json(PATINA_REPORT)
    quality_rows = load_json(QUALITY_REPORT)
    completed = {row["file"] for row in quality_rows}
    for index, path in enumerate(files, 1):
        rel = str(path.relative_to(ROOT))
        if rel in completed:
            continue
        print(f"[{index}/{len(files)}] {rel}", flush=True)
        copy_report, patina_report, quality = run_gate_for_file(path)
        copy_reports.append(copy_report)
        patina_reports.append(patina_report)
        quality_rows.append(quality)
        save_json(COPYEDITOR_REPORT, copy_reports)
        save_json(PATINA_REPORT, patina_reports)
        save_json(QUALITY_REPORT, quality_rows)
        state.update({"completed_files": len(quality_rows), "last_file": rel})
        save_json(STATE_JSON, state)
    write_combined_docx(quality_rows)
    write_summary(copy_reports, patina_reports, quality_rows)
    pass_count = sum(1 for row in quality_rows if row["recommendation"] == "제출권장")
    print(json.dumps({"processed": len(quality_rows), "submission_recommended": pass_count}, ensure_ascii=False))
    return 0 if pass_count == len(quality_rows) else 1


def rebuild_quality_only() -> int:
    quality_rows: list[dict] = []
    for path in collect_files():
        rel = str(path.relative_to(ROOT))
        final_name = f"{safe_name(path)}_final_gate.md"
        final_path = FINAL_MD_DIR / final_name
        markdown_path = final_path if final_path.exists() else path
        markdown = markdown_path.read_text(encoding="utf-8")
        org = infer_gate_org(path.name, markdown)
        questions, responses = parse_questions(markdown)
        if markdown_path.exists():
            write_docx(FINAL_DOCX_DIR / f"{safe_name(path)}_final_gate.docx", markdown)
        quality = quality_audit(rel, org, questions, responses)
        quality["final_file"] = str(markdown_path.relative_to(ROOT))
        docx_name = f"{safe_name(path)}_final_gate.docx"
        docx_path = FINAL_DOCX_DIR / docx_name
        quality["final_docx_file"] = str(docx_path.relative_to(ROOT)) if docx_path.exists() else ""
        quality_rows.append(quality)
    save_json(QUALITY_REPORT, quality_rows)
    copy_reports = load_json(COPYEDITOR_REPORT)
    patina_reports = load_json(PATINA_REPORT)
    write_combined_docx(quality_rows)
    write_summary(copy_reports, patina_reports, quality_rows)
    pass_count = sum(1 for row in quality_rows if row["recommendation"] == "제출권장")
    print(json.dumps({"processed": len(quality_rows), "submission_recommended": pass_count}, ensure_ascii=False))
    return 0 if pass_count == len(quality_rows) else 1


def rerun_copyeditor_only() -> int:
    """Retry copyeditor without pretending old Patina scores are still current."""
    OUT_DIR.mkdir(exist_ok=True)
    FINAL_MD_DIR.mkdir(exist_ok=True)
    FINAL_DOCX_DIR.mkdir(exist_ok=True)
    files = collect_files()
    copy_by_file = {row["file"]: row for row in load_json(COPYEDITOR_REPORT)}
    patina_by_file = {row["file"]: row for row in load_json(PATINA_REPORT)}
    targets = []
    for path in files:
        rel = str(path.relative_to(ROOT))
        report = copy_by_file.get(rel)
        if report is None or any(
            item.get("status") == "fallback_backend_error"
            for item in report.get("items", [])
        ):
            targets.append(path)

    save_json(
        STATE_JSON,
        {
            "mode": "rerun_copyeditor_only",
            "total_files": len(files),
            "target_files": len(targets),
            "completed_files": 0,
        },
    )
    for index, path in enumerate(targets, 1):
        rel = str(path.relative_to(ROOT))
        final_path = FINAL_MD_DIR / f"{safe_name(path)}_final_gate.md"
        markdown_path = final_path if final_path.exists() else path
        markdown = markdown_path.read_text(encoding="utf-8")
        org = infer_gate_org(path.name, markdown)
        questions, responses = parse_questions(markdown)
        edited, items = copyedit_responses(
            responses,
            target_org=org,
            job_terms=GATE_JOB_TERMS.get(org, GATE_JOB_TERMS["지원기관"]),
            timeout_ms=180_000,
        )
        applied_count = sum(item.get("status") == "copyedited" for item in items)
        copy_by_file[rel] = {
            "file": rel,
            "organization": org,
            "question_count": len(questions),
            "copyeditor_attempted": True,
            "copyeditor_applied_count": applied_count,
            "items": [
                {
                    **item,
                    "copyeditor_attempted": True,
                    "copyeditor_applied": item.get("status") == "copyedited",
                }
                for item in items
            ],
        }
        if applied_count:
            output_markdown = replace_answers(markdown, edited)
            final_path.write_text(output_markdown, encoding="utf-8")
            write_docx(
                FINAL_DOCX_DIR / f"{safe_name(path)}_final_gate.docx",
                output_markdown,
            )
            existing_patina = patina_by_file.get(rel)
            if existing_patina:
                for item in existing_patina.get("items", []):
                    item["score_status"] = "stale_after_copyeditor"
                    item["selected_ai_score"] = None
                    item["ai_score_gate"] = "blocked"
                    item["message"] = "copyeditor 적용 후 Patina 재측정 필요"
        save_json(COPYEDITOR_REPORT, [copy_by_file[str(item.relative_to(ROOT))] for item in files if str(item.relative_to(ROOT)) in copy_by_file])
        save_json(PATINA_REPORT, [patina_by_file[str(item.relative_to(ROOT))] for item in files if str(item.relative_to(ROOT)) in patina_by_file])
        save_json(
            STATE_JSON,
            {
                "mode": "rerun_copyeditor_only",
                "total_files": len(files),
                "target_files": len(targets),
                "completed_files": index,
                "last_file": rel,
            },
        )

    result = rebuild_quality_only()
    return result


def _report_has_blocked_model_gate(report: dict, *, kind: str) -> bool:
    if kind == "copyeditor":
        return any(
            item.get("status") == "fallback_backend_error"
            for item in report.get("items", [])
        )
    if kind == "patina":
        return any(
            item.get("selected_ai_score") is None
            or item.get("ai_score_gate") == "blocked"
            for item in report.get("items", [])
        )
    return True


def rerun_model_blocked_only() -> int:
    OUT_DIR.mkdir(exist_ok=True)
    FINAL_MD_DIR.mkdir(exist_ok=True)
    FINAL_DOCX_DIR.mkdir(exist_ok=True)
    files = collect_files()
    copy_reports = load_json(COPYEDITOR_REPORT)
    patina_reports = load_json(PATINA_REPORT)
    quality_rows = load_json(QUALITY_REPORT)
    copy_by_file = {row["file"]: row for row in copy_reports}
    patina_by_file = {row["file"]: row for row in patina_reports}
    quality_by_file = {row["file"]: row for row in quality_rows}

    targets: list[Path] = []
    for path in files:
        rel = str(path.relative_to(ROOT))
        copy_report = copy_by_file.get(rel)
        patina_report = patina_by_file.get(rel)
        if (
            copy_report is None
            or patina_report is None
            or _report_has_blocked_model_gate(copy_report, kind="copyeditor")
            or _report_has_blocked_model_gate(patina_report, kind="patina")
        ):
            targets.append(path)

    state = load_json(STATE_JSON)
    save_json(
        STATE_JSON,
        {
            "mode": "rerun_model_blocked_only",
            "total_files": len(files),
            "target_files": len(targets),
            "patina_doctor": patina_doctor(),
        },
    )

    for index, path in enumerate(targets, 1):
        rel = str(path.relative_to(ROOT))
        print(f"[retry {index}/{len(targets)}] {rel}", flush=True)
        copy_report, patina_report, quality = run_gate_for_file(
            path,
            prefer_existing_final=True,
        )
        copy_by_file[rel] = copy_report
        patina_by_file[rel] = patina_report
        quality_by_file[rel] = quality
        ordered_copy_reports = [copy_by_file[str(item.relative_to(ROOT))] for item in files if str(item.relative_to(ROOT)) in copy_by_file]
        ordered_patina_reports = [patina_by_file[str(item.relative_to(ROOT))] for item in files if str(item.relative_to(ROOT)) in patina_by_file]
        ordered_quality_rows = [quality_by_file[str(item.relative_to(ROOT))] for item in files if str(item.relative_to(ROOT)) in quality_by_file]
        save_json(COPYEDITOR_REPORT, ordered_copy_reports)
        save_json(PATINA_REPORT, ordered_patina_reports)
        save_json(QUALITY_REPORT, ordered_quality_rows)
        save_json(
            STATE_JSON,
            {
                "mode": "rerun_model_blocked_only",
                "total_files": len(files),
                "target_files": len(targets),
                "completed_retry_files": index,
                "last_file": rel,
                "patina_doctor": patina_doctor(),
            },
        )

    ordered_copy_reports = [copy_by_file[str(item.relative_to(ROOT))] for item in files if str(item.relative_to(ROOT)) in copy_by_file]
    ordered_patina_reports = [patina_by_file[str(item.relative_to(ROOT))] for item in files if str(item.relative_to(ROOT)) in patina_by_file]
    ordered_quality_rows = [quality_by_file[str(item.relative_to(ROOT))] for item in files if str(item.relative_to(ROOT)) in quality_by_file]
    write_combined_docx(ordered_quality_rows)
    write_summary(ordered_copy_reports, ordered_patina_reports, ordered_quality_rows)

    total_questions = sum(row["question_count"] for row in ordered_quality_rows)
    copy_backend_ok = sum(
        1
        for report in ordered_copy_reports
        for item in report["items"]
        if item["status"] != "fallback_backend_error"
    )
    patina_scored = sum(
        1
        for report in ordered_patina_reports
        for item in report["items"]
        if item["selected_ai_score"] is not None
    )
    pass_count = sum(1 for row in ordered_quality_rows if row["recommendation"] == "제출권장")
    result = {
        "processed": len(ordered_quality_rows),
        "submission_recommended": pass_count,
        "copyeditor_backend_ok": copy_backend_ok,
        "patina_scored": patina_scored,
        "total_questions": total_questions,
    }
    print(json.dumps(result, ensure_ascii=False))
    return 0 if (
        pass_count == len(ordered_quality_rows)
        and copy_backend_ok == total_questions
        and patina_scored == total_questions
    ) else 1


if __name__ == "__main__":
    if "--rebuild-quality-only" in sys.argv:
        raise SystemExit(rebuild_quality_only())
    if "--rerun-copyeditor-only" in sys.argv:
        raise SystemExit(rerun_copyeditor_only())
    if "--rerun-model-blocked-only" in sys.argv:
        raise SystemExit(rerun_model_blocked_only())
    raise SystemExit(main())
