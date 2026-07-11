from pathlib import Path
import csv
import json
import zipfile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

root = Path.cwd()
outdir = root / "jasoseo_all_review_20260705"
records = json.loads((outdir / "evaluation_records.json").read_text(encoding="utf-8"))

blank_files = {
    "26-01-16_NH농협은행 6급초급_일반.docx",
    "26-03-04_국민건강보험공단(h·well) [일반]강원_행정 6급가.docx",
    "26-03-17_국민연금공단(NPS) [일반전형-권역] 6급 사무직.docx",
    "26-05-27_우리은행 지역인재 (강원).docx",
    "26-06-21_주택도시보증공사(HUG) 일반전형_금융·기금(강원).docx",
    "새로운 자기소개서\\우리인턴.docx",
    "새로운 자기소개서\\해양박물관.docx",
    "서울교통공사 자소서.docx",
}

for r in records:
    if r["file"] in blank_files:
        r["status"] = "초안 작성 필요"
        r["score"] = 25 if "서울교통공사 자소서" not in r["file"] else 30
        r["note"] = "답변 공란 또는 문항 템플릿 중심"
        r["issues"] = ["답변 본문이 없거나 문항 안내 중심이라 평가 가능한 서술이 부족함"]
        r["actions"] = ["기존 검증 경험 4종(데이터 정리, 숙박비 검증, 도서관 개선, 고령 고객 응대)을 문항별로 배치해 초안부터 작성"]
        r["outline"] = (
            f"{r['group']} 문항은 아직 답변이 비어 있으므로, 데이터 정리 경험은 업무효율/직무역량, "
            "숙박비 검증 경험은 원칙/윤리, 도서관 개선 경험은 협업/개선, 고령 고객 응대는 고객보호/소통 문항에 배치해 초안을 작성합니다."
        )

best_labels = [
    "NH농협은행 6급 제출 기준본",
    "지역농협 제출 기준본",
    "국민연금공단 최신 교열본",
    "HUG 강원 최신 상위권본",
    "사회보장정보원 제출 기준본",
    "IBK기업은행 제출 기준본",
    "하나저축은행 최신 수정안",
    "서울교통공사 고객안전 최신본",
    "한국도로공사서비스 통합본",
]

with (outdir / "evaluation_table.csv").open("w", encoding="utf-8-sig", newline="") as f:
    writer = csv.DictWriter(
        f,
        fieldnames=["no", "group", "status", "score", "file", "note", "issues", "actions", "outline", "chars_no_space"],
    )
    writer.writeheader()
    for row in records:
        out = row.copy()
        out["issues"] = " / ".join(row["issues"])
        out["actions"] = " / ".join(row["actions"])
        writer.writerow(out)

md = []
md.append("# 취업 폴더 자기소개서 전체 평가 및 개선안")
md.append("")
md.append("- 생성일: 2026-07-05")
md.append(f"- 처리 대상: 제출용/초안 성격 자기소개서 {len(records)}개")
md.append("- 적용 기준: 문항 대응도, 기관·직무 적합도, 경험 근거, 사실성/면접 방어, 문장 가독성, 블라인드/분량 리스크")
md.append("- 제한: 이번 일괄 평가는 로컬 문서 기준입니다. 최신 공식 공고·기관 수치·경영진 인용은 별도 외부 검증이 필요합니다.")
md.append("")
md.append("## 제출 기준 우선순위")
best_rows = [r for r in records if r["status"] == "제출 기준본"]
for label, row in zip(best_labels, best_rows[: len(best_labels)]):
    md.append(f"- {label}: {row['file']} ({row['score']}점, {row['status']})")
md.append("")
md.append("## 전체 평가표")
md.append("|No|기관/그룹|상태|점수|파일|핵심 조치|")
md.append("|---:|---|---|---:|---|---|")
for r in records:
    md.append(f"|{r['no']}|{r['group']}|{r['status']}|{r['score']}|{r['file']}|{'; '.join(r['actions'])}|")
md.append("")
md.append("## 문서별 개선안")
for r in records:
    md.append(f"### {r['no']}. {r['file']}")
    md.append(f"- 평가: {r['score']}점 / {r['status']} / {r['note']}")
    md.append(f"- 문제: {'; '.join(r['issues'])}")
    md.append(f"- 개선안: {'; '.join(r['actions'])}")
    md.append(f"- 작성 방향: {r['outline']}")
    md.append("")
(outdir / "all_self_intro_evaluation_and_revisions.md").write_text("\n".join(md), encoding="utf-8")

doc = Document()
styles = doc.styles
styles["Normal"].font.name = "Malgun Gothic"
styles["Normal"].font.size = Pt(10)
for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
    styles[style_name].font.name = "Malgun Gothic"

title = doc.add_heading("취업 폴더 자기소개서 전체 평가 및 개선안", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(
    "생성일: 2026-07-05\n"
    f"처리 대상: 제출용/초안 성격 자기소개서 {len(records)}개\n"
    "적용 기준: 문항 대응도, 기관·직무 적합도, 경험 근거, 사실성/면접 방어, 문장 가독성, 블라인드/분량 리스크\n"
    "제한: 최신 공식 공고·기관 수치·경영진 인용은 이번 일괄 평가에서 외부 검증하지 않았습니다."
)

doc.add_heading("1. 제출 기준 우선순위", level=2)
for label, row in zip(best_labels, best_rows[: len(best_labels)]):
    doc.add_paragraph(f"{label}: {row['file']} ({row['score']}점, {row['status']})")

doc.add_heading("2. 전체 평가표", level=2)
table = doc.add_table(rows=1, cols=6)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = "Table Grid"
for idx, header in enumerate(["No", "기관/그룹", "상태", "점수", "파일", "핵심 조치"]):
    table.rows[0].cells[idx].text = header
for r in records:
    cells = table.add_row().cells
    values = [str(r["no"]), r["group"], r["status"], str(r["score"]), r["file"], "; ".join(r["actions"])]
    for cell, value in zip(cells, values):
        cell.text = value
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

doc.add_heading("3. 문서별 개선안", level=2)
for r in records:
    doc.add_heading(f"{r['no']}. {r['file']}", level=3)
    doc.add_paragraph(f"평가: {r['score']}점 / {r['status']} / {r['note']}")
    doc.add_paragraph(f"문제: {'; '.join(r['issues'])}")
    doc.add_paragraph(f"개선안: {'; '.join(r['actions'])}")
    doc.add_paragraph(f"작성 방향: {r['outline']}")

report_path = outdir / "all_self_intro_evaluation_and_revisions.docx"
doc.save(report_path)
(outdir / "evaluation_records.json").write_text(json.dumps(records, ensure_ascii=False, indent=2), encoding="utf-8")

with zipfile.ZipFile(report_path) as z:
    assert z.testzip() is None
    assert "word/document.xml" in z.namelist()
Document(str(report_path))
print("rewritten:", report_path)
print("records:", len(records))
print("blank_count:", sum(1 for r in records if r["status"] == "초안 작성 필요"))
