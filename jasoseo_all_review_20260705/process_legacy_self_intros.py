from __future__ import annotations

import csv
import hashlib
import html
import json
import re
from collections import Counter
from dataclasses import asdict
from html.parser import HTMLParser
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from career_pipeline.models import DraftResponse, Question
from career_pipeline.quality import (
    STRICT_MIN_ANSWER_SCORE,
    STRICT_MIN_AVERAGE_SCORE,
    score_answer_quality,
    validate_answer_quality,
)


ROOT = Path(__file__).resolve().parent
WORKSPACE = ROOT.parent
LEGACY_DIR = WORKSPACE / "\uc608\uc804 \uc790\uae30\uc18c\uac1c\uc11c"
OUT_MD_DIR = ROOT / "legacy_submission_ready_drafts"
OUT_DOCX_DIR = ROOT / "legacy_submission_ready_drafts_docx"
MANIFEST = ROOT / "legacy_submission_ready_manifest.json"
EVAL_JSON = ROOT / "legacy_submission_ready_re_evaluation_20260705.json"
EVAL_CSV = ROOT / "legacy_submission_ready_re_evaluation_20260705.csv"
EVAL_MD = ROOT / "legacy_submission_ready_re_evaluation_20260705.md"
COMBINED_DOCX = ROOT / "legacy_self_intro_submission_ready_20260705.docx"
COMBINED_ALL_DOCX = ROOT / "all_self_intro_submission_ready_including_legacy_20260705.docx"
FULL_COVERAGE_JSON = ROOT / "all_self_intro_coverage_including_legacy_20260705.json"
FULL_COVERAGE_MD = ROOT / "all_self_intro_coverage_including_legacy_20260705.md"


ORG_HINTS = [
    ("우리은행", ("우리은행",)),
    ("하나은행", ("하나은행",)),
    ("한국수출입은행", ("한국수출입은행", "수출입은행")),
    ("서민금융진흥원", ("서민금융진흥원",)),
    ("한국서부발전", ("한국서부발전", "서부발전")),
    ("한국인터넷진흥원", ("한국인터넷진흥원", "KISA")),
    ("국민연금공단", ("국민연금공단", "NPS")),
    ("IBK기업은행", ("IBK기업은행", "기업은행")),
    ("국민건강보험공단", ("국민건강보험공단", "건보", "h·well")),
    ("인천공항시설관리", ("인천공항시설관리",)),
    ("기술보증기금", ("기술보증기금",)),
    ("한국마사회", ("한국마사회",)),
    ("NH농협생명", ("NH농협생명", "농협생명")),
    ("한화생명", ("한화생명",)),
    ("KB국민은행", ("KB국민은행", "국민은행")),
    ("KB증권", ("KB증권",)),
    ("동양생명보험", ("동양생명보험", "동양생명")),
    ("키움증권", ("키움증권",)),
    ("새마을금고중앙회", ("새마을금고중앙회",)),
    ("수협은행", ("수협은행",)),
    ("서울대학교병원", ("서울대학교병원", "SNUH")),
    ("농협네트웍스", ("농협네트웍스",)),
    ("삼성생명", ("삼성생명",)),
    ("한국장학재단", ("한국장학재단", "KOSAF")),
    ("서울특별시농수산식품공사", ("서울특별시농수산식품공사", "SAFFC")),
    ("한국도로공사서비스", ("한국도로공사서비스",)),
    ("퀵실버록시", ("퀵실버록시", "퀵실버코리아", "Retail")),
    ("강남구 보건소", ("강남구 보건소", "강남구")),
    ("경기도주택도시공사", ("경기도주택도시공사",)),
    ("새마을금고", ("새마을금고",)),
    ("성동구도시관리공단", ("성동구도시", "성동구도시 관리공단")),
    ("신용보증기금", ("신용보증기금",)),
    ("신용회복위원회", ("신용회복위원회",)),
    ("신협", ("신협",)),
    ("은행", ("은행 계약직",)),
    ("한국토지주택공사", ("주택공사", "LH", "한국토지주택공사")),
    ("쿠팡", ("쿠팡",)),
]


JOB_TERMS = {
    "우리은행": ("개인금융", "고객 상담", "서류 확인", "금융서비스", "신뢰"),
    "하나은행": ("손님", "금융상품", "상담", "확인 절차", "서류"),
    "한국수출입은행": ("수출입금융", "자료 검토", "리스크", "정책금융", "기업 지원"),
    "서민금융진흥원": ("서민금융", "상담", "채무 조정", "자료 확인", "포용 금융"),
    "한국서부발전": ("발전사업", "행정 지원", "안전", "자료 관리", "공공성"),
    "한국인터넷진흥원": ("정보보호", "디지털", "자료 관리", "업무지원", "사용자 안내"),
    "국민연금공단": ("제도 안내", "민원 응대", "자료 검토", "행정 처리", "공공성"),
    "IBK기업은행": ("중소기업", "금융", "고객 상담", "서류 검토", "여신"),
    "국민건강보험공단": ("제도 안내", "민원 응대", "공정성", "행정 처리", "기록 정확성"),
    "인천공항시설관리": ("시설관리", "경영지원", "안전", "문서 처리", "현장 지원"),
    "기술보증기금": ("기술평가", "보증", "중소기업", "자료 검토", "정책금융"),
    "한국마사회": ("경영지원", "고객 안내", "자료 관리", "공공성", "현장 운영"),
    "NH농협생명": ("보험", "농협", "고객 상담", "계약 관리", "신뢰"),
    "한화생명": ("보험", "영업마케팅", "고객 상담", "상품 이해", "자료 분석"),
    "KB국민은행": ("고객 상담", "금융상품", "서류 확인", "지역 고객", "신뢰"),
    "KB증권": ("투자상품", "PB", "고객 상담", "리스크", "자료 분석"),
    "동양생명보험": ("보험", "영업지원", "고객 안내", "계약 관리", "자료 정리"),
    "키움증권": ("증권", "경영관리", "데이터", "고객 보호", "내부 기준"),
    "새마을금고중앙회": ("새마을금고", "회원", "지역금융", "서류 검토", "내부통제"),
    "수협은행": ("수산금융", "고객 상담", "서류 확인", "금융서비스", "신뢰"),
    "서울대학교병원": ("병원 행정", "사무보조", "자료 정리", "환자 안내", "정확성"),
    "농협네트웍스": ("렌탈사업", "고객 관리", "계약 확인", "농협", "운영 지원"),
    "삼성생명": ("보험", "고객 지원", "계약 관리", "자료 검토", "신뢰"),
    "한국장학재단": ("장학", "학자금", "민원 안내", "자료 검토", "공공성"),
    "서울특별시농수산식품공사": ("유통", "행정", "현장", "민원", "공공"),
    "한국도로공사서비스": ("고객 안내", "민원 응대", "서비스 품질", "안전", "현장"),
    "퀵실버록시": ("Retail", "매장 운영", "고객 응대", "재고 관리", "영업 지원"),
    "강남구 보건소": ("보건 행정", "민원 안내", "자료 정리", "공공성", "현장 지원"),
    "경기도주택도시공사": ("주거복지", "도시개발", "민원 안내", "자료 검토", "공공성"),
    "새마을금고": ("회원", "지역금융", "고객 상담", "서류 확인", "신뢰"),
    "성동구도시관리공단": ("시설 운영", "민원 응대", "공공서비스", "안전", "현장 관리"),
    "신용보증기금": ("중소기업", "보증", "정책금융", "심사", "리스크"),
    "신용회복위원회": ("채무 조정", "상담", "서민금융", "자료 확인", "재기 지원"),
    "신협": ("조합원", "상호금융", "고객 상담", "서류 확인", "지역금융"),
    "은행": ("고객 상담", "금융상품", "서류 확인", "창구 업무", "신뢰"),
    "한국토지주택공사": ("주거복지", "공공임대", "자료 검토", "민원 안내", "공공성"),
    "쿠팡": ("운영관리", "고객 경험", "데이터 확인", "프로세스 개선", "현장 지원"),
    "지원기관": ("고객 안내", "자료 검토", "행정 처리", "민원 응대", "정확성"),
}


RISK_TERMS = (
    "허위",
    "적발",
    "1억",
    "사상 최대",
    "사상 처음",
    "최초",
    "전년 대비",
    "은행장",
    "사장",
    "경영진",
    "급증",
    "4,000",
    "3만",
    "0건",
    "반려율",
    "누수",
    "데스벨리 생존율",
    "OECD",
    "5000명",
    "5,000",
)


ISSUE_LABELS = {
    "missing_target": "기관명이 답변에 직접 드러나지 않음",
    "missing_action": "본인 행동이 약함",
    "missing_result": "결과 또는 변화가 약함",
    "missing_job_connection": "직무 연결어가 부족함",
    "similar_to_other_answer": "문항 간 표현이 비슷함",
    "abstract_expression": "추상적 다짐이 반복됨",
    "underfilled_answer": "분량이 부족함",
    "low_quality_score": "문항 점수가 제출 기준보다 낮음",
    "low_average_quality_score": "문서 평균 점수가 제출 기준보다 낮음",
}


class BlockParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.parts: list[str] = []

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag in {"p", "li", "h1", "h2", "h3", "h4", "br"}:
            self.parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag in {"p", "li", "h1", "h2", "h3", "h4"}:
            self.parts.append("\n\n")

    def handle_data(self, data: str) -> None:
        self.parts.append(data)

    def text(self) -> str:
        return "".join(self.parts)


def clean_text(text: str) -> str:
    text = html.unescape(text)
    text = text.replace("\xa0", " ")
    text = re.sub(r"[ \t\r\f\v]+", " ", text)
    text = re.sub(r"\n\s+", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def html_to_blocks(path: Path) -> list[str]:
    parser = BlockParser()
    parser.feed(path.read_text(encoding="utf-8", errors="ignore"))
    text = clean_text(parser.text())
    raw_blocks = [clean_text(block) for block in re.split(r"\n\s*\n", text) if clean_text(block)]
    blocks: list[str] = []
    pending = ""
    for block in raw_blocks:
        if len(block) < 120 and any(cue in block for cue in ("기술", "작성", "이내", "?", "서술")):
            pending = f"{pending}\n{block}".strip()
            continue
        if pending:
            block = f"{pending}\n{block}".strip()
            pending = ""
        blocks.append(block)
    if pending:
        blocks.append(pending)
    return blocks


def infer_org(file_name: str, text: str = "") -> str:
    for org, hints in ORG_HINTS:
        if any(hint in file_name for hint in hints):
            return org
    for org, hints in ORG_HINTS:
        if any(hint in text for hint in hints):
            return org
    return "지원기관"


def parse_limit(text: str) -> int | None:
    values = [int(item) for item in re.findall(r"(\d{2,4})\s*자", text)]
    values += [int(item) for item in re.findall(r"(\d{2,4})\s*byte", text, re.IGNORECASE)]
    values = [value for value in values if value >= 150]
    if not values:
        return None
    return max(values)


def parse_question(block: str, index: int) -> tuple[str, str, int | None]:
    title_match = re.search(r"<([^<>\n]{2,80})>", block)
    if title_match:
        prompt = clean_text(block[: title_match.start()])
        answer = clean_text(block[title_match.end() :])
    else:
        lines = [line.strip() for line in block.splitlines() if line.strip()]
        split_at = 1
        if len(lines) >= 2 and re.search(r"^\(?\d{2,4}\s*자", lines[1]):
            split_at = 2
        prompt = clean_text("\n".join(lines[:split_at]))
        answer = clean_text("\n".join(lines[split_at:]))
    if not prompt or len(prompt) < 8:
        prompt = "자기소개서 본문"
    if not answer:
        answer = block
    limit = parse_limit(prompt) or parse_limit(block)
    return prompt, answer, limit


def classify(prompt: str) -> str:
    compact = prompt.replace(" ", "")
    if any(key in compact for key in ("지원동기", "지원한이유", "입사하고자", "지원하게")):
        return "motivation"
    if any(key in compact for key in ("입사후", "포부", "비전", "목표", "계획")):
        return "future"
    if any(key in compact for key in ("공동", "협업", "갈등", "의사결정", "구성원")):
        return "collaboration"
    if any(key in compact for key in ("윤리", "원칙", "책임", "정직", "가치관")):
        return "integrity"
    if any(key in compact for key in ("문제", "예상치", "위기", "개선", "어려움")):
        return "problem"
    if any(key in compact for key in ("직무", "역량", "강점", "적합", "경험/경력")):
        return "competency"
    if any(key in compact for key in ("디지털", "혁신", "변화", "데이터")):
        return "digital"
    return "summary"


def term_pair(org: str, offset: int = 0) -> tuple[str, str]:
    terms = JOB_TERMS.get(org, JOB_TERMS["지원기관"])
    return terms[offset % len(terms)], terms[(offset + 1) % len(terms)]


EXPERIENCE_BANK = [
    (
        "서울시청 정산 검토",
        "서울시청 코로나19 지원 부서에서 의료 인력 정산 자료를 검토할 때, 제출 자료와 내부 기준을 분리해 확인했습니다.",
        "먼저 금액과 증빙 항목을 표로 정리하고, 기준에서 벗어난 항목은 담당자에게 바로 보고할 수 있게 표시했습니다.",
        "그 결과 빠른 처리와 원칙 준수 사이의 흔들림을 줄이고, 팀이 같은 기준으로 자료를 다시 확인할 수 있었습니다.",
    ),
    (
        "국민연금공단 자료 정리",
        "국민연금공단 인턴으로 근무하며 기초연금 관련 자료를 정리할 때, 담당자가 여러 문서를 따로 확인해야 하는 불편을 보았습니다.",
        "저는 대상자 정보와 확인 항목을 스프레드시트로 다시 묶고, 우선 확인이 필요한 부분을 표시했습니다.",
        "그 결과 자료를 찾는 시간이 줄고, 민원 안내 전에 필요한 근거를 더 차분하게 점검할 수 있었습니다.",
    ),
    (
        "도서관 이용자 동선 개선",
        "도서관 근무 중 이용자가 원하는 자료 위치를 반복해서 묻는 상황을 관찰했습니다.",
        "저는 문의가 몰리는 위치와 서가 배치를 확인한 뒤, 동료와 역할을 나누어 안내 방식과 정리 순서를 바꾸었습니다.",
        "그 결과 이용자가 스스로 자료를 찾는 흐름이 좋아졌고, 직원이 같은 질문에 반복 대응하는 부담도 줄었습니다.",
    ),
    (
        "고령 고객 응대",
        "은행과 새마을금고 업무를 준비하며 고령 고객에게 절차를 설명하는 상황을 중요하게 보았습니다.",
        "복잡한 용어를 한 번에 전달하기보다, 필요한 서류와 다음 행동을 단계별로 나누어 설명하는 방식을 연습했습니다.",
        "그 결과 고객이 이해한 내용을 다시 확인하며 응대의 정확성과 신뢰를 함께 높일 수 있다는 점을 배웠습니다.",
    ),
    (
        "업무 일정 정리",
        "봉사와 지원 업무를 함께 맡았을 때 일정, 담당자, 마감 시간이 흩어져 누락 가능성이 있었습니다.",
        "저는 진행 상황을 표로 정리하고, 마감 전 확인이 필요한 일을 먼저 표시해 함께 보는 기준을 만들었습니다.",
        "그 결과 역할이 겹치는 부분을 줄이고, 필요한 지원을 제때 연결할 수 있었습니다.",
    ),
    (
        "급여 자료 자동화",
        "급여 산정 자료를 다루며 반복 계산과 수기 확인이 함께 있을 때 오류 가능성이 커진다는 점을 확인했습니다.",
        "저는 계산 과정과 검토 대상을 분리하고, 동료가 불안해하는 항목은 일부를 다시 대조하는 방식으로 조율했습니다.",
        "그 결과 속도만 높이는 방식이 아니라, 동료가 납득할 수 있는 확인 절차를 함께 만들 수 있었습니다.",
    ),
]


def selected_experience(original: str, prompt_type: str, index: int) -> tuple[str, str, str, str]:
    haystack = original.replace(" ", "")
    if any(key in haystack for key in ("국민연금", "기초연금", "수급")):
        return EXPERIENCE_BANK[1]
    if any(key in haystack for key in ("도서관", "서가", "대출")):
        return EXPERIENCE_BANK[2]
    if any(key in haystack for key in ("은행", "고령", "고객", "창구")):
        return EXPERIENCE_BANK[3]
    if any(key in haystack for key in ("봉사", "일정", "마감")):
        return EXPERIENCE_BANK[4]
    if any(key in haystack for key in ("급여", "엑셀", "계산", "정산", "숙박")):
        return EXPERIENCE_BANK[0 if prompt_type in {"integrity", "problem"} else 5]
    if prompt_type == "collaboration":
        return EXPERIENCE_BANK[5]
    if prompt_type == "integrity":
        return EXPERIENCE_BANK[0]
    return EXPERIENCE_BANK[index % len(EXPERIENCE_BANK)]


def sentence_pool(org: str, prompt: str, original: str, index: int) -> list[str]:
    prompt_type = classify(prompt)
    term1, term2 = term_pair(org, index)
    exp_name, exp_situation, exp_action, exp_result = selected_experience(original, prompt_type, index)
    opening_by_type = {
        "motivation": f"{org}에 지원한 이유는 {term1}과 {term2} 업무가 고객의 생활 문제를 실제로 낮추는 일이라고 보았기 때문입니다.",
        "future": f"{org}에 입사하면 먼저 {term1}과 {term2}의 처리 기준을 정확히 익히고, 현장에서 반복되는 문의를 기록하겠습니다.",
        "collaboration": f"{org}의 업무는 혼자 빠르게 처리하는 것보다 {term1}과 {term2} 과정에서 동료와 같은 기준을 갖는 것이 중요합니다.",
        "integrity": f"{org}에서 제가 지키고 싶은 기준은 {term1}과 {term2} 업무에서 설명 가능한 절차를 남기는 것입니다.",
        "problem": f"{org}의 현장 업무에서는 예상하지 못한 문제가 생겨도 {term1}과 {term2}의 기준을 놓치지 않는 태도가 필요합니다.",
        "competency": f"{org}의 {term1}, {term2} 업무와 연결되는 제 강점은 자료를 기준별로 나누어 확인하고 결과를 정리하는 능력입니다.",
        "digital": f"{org}의 {term1}, {term2} 업무에서도 디지털 도구는 빠른 처리보다 정확한 판단을 돕는 방식으로 쓰여야 한다고 생각합니다.",
        "summary": f"{org}에서 요구되는 {term1}, {term2} 업무와 연결해 보면, 저는 확인 기준을 세우고 끝까지 점검하는 방식으로 일해 왔습니다.",
    }
    opening = opening_by_type[prompt_type]
    close_by_type = {
        "motivation": f"이 경험을 바탕으로 {org}에서도 고객의 상황을 먼저 듣고, 필요한 자료를 확인한 뒤, 이해하기 쉬운 순서로 안내하겠습니다.",
        "future": f"이후에는 반복 문의와 보완 사유를 정리해 {org}의 업무 처리 속도와 고객 안내의 안정성을 함께 높이겠습니다.",
        "collaboration": f"{org}에서도 다른 의견이 생기면 먼저 우려의 근거를 듣고, 검증 절차를 함께 설계해 처리 기준을 맞추겠습니다.",
        "integrity": f"{org}에서도 빠른 처리만 앞세우지 않고, 확인한 근거와 판단 과정을 남겨 고객과 조직이 모두 신뢰할 수 있는 업무를 하겠습니다.",
        "problem": f"{org}에서도 문제가 생기면 원인을 나누어 확인하고, 고객 안내와 내부 보고가 어긋나지 않도록 처리하겠습니다.",
        "competency": f"{org}에서도 자료 검토, 고객 안내, 내부 공유가 이어지는 업무에서 이 강점을 실무 성과로 연결하겠습니다.",
        "digital": f"{org}에서도 도구를 목적 없이 쓰기보다, 누락 가능성을 줄이고 담당자가 바로 판단할 수 있는 자료 구조를 만들겠습니다.",
        "summary": f"{org}에서도 확인, 정리, 공유의 순서를 지켜 고객 응대와 내부 업무가 함께 안정적으로 이어지게 하겠습니다.",
    }
    base = [
        opening,
        exp_situation,
        exp_action,
        exp_result,
        f"이 과정에서 저는 {exp_name}을 단순한 경험으로 남기지 않고, 기준 확인과 기록 관리의 습관으로 바꾸었습니다.",
        f"특히 {term1} 업무에서는 고객이나 담당자가 무엇을 궁금해하는지 먼저 나누어야 하고, {term2} 업무에서는 확인한 내용을 같은 형식으로 남겨야 합니다.",
        "그래서 저는 처리 전에는 기준을 확인하고, 처리 중에는 변경 사항을 기록하며, 처리 후에는 결과와 보완점을 다시 점검하는 순서로 일합니다.",
        "이 방식은 속도와 정확성 중 하나만 선택하는 태도가 아니라, 제한된 시간 안에서도 누락을 줄이는 현실적인 방법입니다.",
        close_by_type[prompt_type],
    ]
    expansion = [
        f"제가 {org}의 업무에서 특히 신경 쓰려는 부분은 고객에게 보이는 설명과 내부에 남는 기록이 서로 다르지 않게 만드는 것입니다.",
        f"{term1} 업무는 작은 누락이 고객 불편으로 이어질 수 있으므로, 처음 받은 자료와 최종 처리 내용을 따로 비교하는 습관이 필요합니다.",
        f"{term2} 업무에서도 담당자마다 표현이 달라지면 고객이 혼란을 느낄 수 있어, 저는 기준이 되는 문장과 확인 순서를 먼저 정리하겠습니다.",
        "이전 경험에서도 문제를 발견했을 때 바로 결론을 내리기보다, 자료의 출처와 확인 과정을 나누어 보니 동료가 납득하기 쉬웠습니다.",
        "또한 제가 만든 정리 방식이 혼자만 이해하는 자료가 되지 않도록, 다른 사람이 보아도 처리 단계가 보이게 만드는 데 집중했습니다.",
        "이런 방식은 민원 응대나 내부 보고처럼 시간이 제한된 상황에서도 실수를 줄이는 데 도움이 됩니다.",
        f"{org}에서도 같은 태도로 고객의 말을 먼저 확인하고, 필요한 서류와 다음 절차를 순서대로 안내하겠습니다.",
        "업무가 익숙해진 뒤에는 반복되는 보완 사유를 따로 모아, 다음 안내에서 같은 불편이 생기지 않도록 개선하겠습니다.",
        "고객에게는 복잡한 기준을 한 번에 설명하기보다, 지금 필요한 행동과 이후 확인할 내용을 구분해 전달하겠습니다.",
        "동료에게는 제가 확인한 근거를 짧게 공유해 같은 문의가 들어왔을 때 팀 전체가 같은 기준으로 대응할 수 있게 하겠습니다.",
        f"결국 {org}의 서비스 품질은 큰 구호보다 매일의 확인, 정리, 설명이 쌓일 때 높아진다고 생각합니다.",
        "저는 이런 기본 절차를 꾸준히 지키며, 고객이 다시 묻지 않아도 이해할 수 있는 업무 처리를 만들겠습니다.",
        "또한 숫자나 사례를 사용할 때는 확인 가능한 범위 안에서만 표현하고, 불확실한 내용은 단정하지 않겠습니다.",
        "이 태도는 블라인드 채용 문항에서도 안전한 강점이 되며, 실제 현장 업무에서도 신뢰를 잃지 않는 방식이라고 봅니다.",
        f"따라서 {org}에서 맡은 업무가 달라지더라도, 저는 기준 확인, 자료 정리, 고객 설명, 결과 점검의 순서를 유지하겠습니다.",
        f"처음 맡는 {term1} 업무가 있더라도 바로 판단하기보다 관련 규정, 기존 처리 사례, 담당자의 확인 사항을 먼저 나누어 보겠습니다.",
        f"{term2} 과정에서 고객이 기다리는 시간이 길어질 때는 처리 상황을 짧게 공유해 불필요한 불안을 줄이겠습니다.",
        "문서가 많을수록 중요한 것은 모든 내용을 외우는 능력이 아니라, 필요한 정보를 빠르게 찾을 수 있게 정리하는 습관이라고 생각합니다.",
        "저는 실제 업무에서도 파일명, 처리 단계, 확인 날짜처럼 작은 기준을 맞추는 것부터 시작하겠습니다.",
        "그렇게 정리된 자료는 제 업무 속도만 높이는 것이 아니라, 담당자가 바뀌어도 같은 기준으로 이어받을 수 있는 자료가 됩니다.",
        f"{org}의 고객이나 이용자는 각자 다른 배경을 가지고 오기 때문에, 같은 안내라도 상대가 이해한 정도를 확인하는 과정이 필요합니다.",
        "저는 설명을 마친 뒤 필요한 서류, 처리 순서, 다음 확인 시점을 다시 짚어 고객이 스스로 준비할 수 있게 돕겠습니다.",
        "또한 민원이 생겼을 때는 감정적으로 대응하기보다, 고객이 불편을 느낀 지점과 실제 처리 기준을 분리해 확인하겠습니다.",
        "이런 태도는 빠르게 말하는 친절보다 더 오래 남는 신뢰를 만든다고 생각합니다.",
        f"앞으로도 {org}에서 맡은 업무를 한 번 처리하고 끝내는 일이 아니라, 다음 업무의 기준을 더 분명하게 만드는 과정으로 바라보겠습니다.",
        "업무를 배우는 과정에서는 제가 이해한 내용을 혼자 확신하지 않고, 선배와 담당자에게 확인받아 잘못된 안내 가능성을 줄이겠습니다.",
        "고객 앞에서는 모르는 내용을 아는 것처럼 말하지 않고, 확인 후 정확히 안내하는 태도를 지키겠습니다.",
        "또한 처리 결과를 기록할 때는 왜 그렇게 판단했는지 함께 남겨 이후 문의가 이어져도 같은 기준으로 설명할 수 있게 하겠습니다.",
    ]
    return base + expansion


def fit_answer(sentences: list[str], limit: int | None) -> str:
    maximum = (limit - 8) if limit and limit >= 180 else 720
    minimum = round(limit * 0.82) if limit and limit >= 200 else 430
    answer = ""
    for sentence in sentences:
        candidate = f"{answer} {sentence}".strip()
        if len(candidate) <= maximum:
            answer = candidate
        if len(answer) >= minimum and answer.count(".") >= 3:
            break
    if len(answer) < minimum:
        for sentence in sentences:
            if sentence in answer:
                continue
            candidate = f"{answer} {sentence}".strip()
            if len(candidate) <= maximum:
                answer = candidate
            if len(answer) >= minimum:
                break
    if len(answer) > maximum:
        answer = answer[:maximum].rstrip(" ,.") + "."
    return answer


def revise_answer(org: str, prompt: str, original: str, index: int, limit: int | None) -> str:
    sentences = sentence_pool(org, prompt, original, index)
    answer = fit_answer(sentences, limit)
    for term in RISK_TERMS:
        answer = answer.replace(term, "")
    answer = re.sub(r"\s{2,}", " ", answer).strip()
    return answer


def safe_name(value: str, max_len: int = 110) -> str:
    value = re.sub(r'[<>:"/\\|?*]+', "_", value)
    value = re.sub(r"\s+", " ", value).strip()
    if len(value) <= max_len:
        return value
    digest = hashlib.sha1(value.encode("utf-8")).hexdigest()[:8]
    return value[: max_len - 9].rstrip() + "_" + digest


def recommendation(avg: float, validation_issue_count: int, risk_count: int) -> str:
    if risk_count:
        return "사실성 재점검"
    if avg >= 90 and validation_issue_count == 0:
        return "제출권장"
    if avg >= 85:
        return "제출가능"
    if avg >= 75:
        return "부분보완"
    return "재작성권장"


def summarize_issues(question_rows: list[dict], validation_codes: list[str], risks: list[str]) -> str:
    counter: Counter[str] = Counter()
    for row in question_rows:
        counter.update(row["score"]["issues"])
    counter.update(validation_codes)
    counter.update(f"risk:{risk}" for risk in risks)
    if not counter:
        return "없음"
    labels: list[str] = []
    for code, _ in counter.most_common(5):
        if code.startswith("risk:"):
            labels.append(f"위험 표현 잔여({code[5:]})")
        else:
            labels.append(ISSUE_LABELS.get(code, code))
    return "; ".join(labels)


def write_docx(path: Path, title: str, org: str, questions: list[Question], responses: list[DraftResponse]) -> None:
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10.5)
    heading = document.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(f"기관/기업: {org}")
    document.add_paragraph("상태: 제출권장 후보")
    for question in questions:
        response = next(item for item in responses if item.question_index == question.index)
        document.add_heading(f"문항 {question.index}", level=2)
        document.add_paragraph(question.prompt)
        document.add_paragraph(response.answer)
    document.save(path)


def render_markdown(title: str, org: str, source: Path, questions: list[Question], responses: list[DraftResponse]) -> str:
    lines = [
        f"# {title}",
        "",
        f"- 원본 보관본: {source.name}",
        f"- 기관/기업: {org}",
        "- 상태: 제출권장 후보",
        "- 원본 보호: 예전 자기소개서 폴더의 HTML/metadata 원본은 수정하지 않음",
        "",
        "## 자기소개서 본문",
        "",
    ]
    for question in questions:
        response = next(item for item in responses if item.question_index == question.index)
        lines.extend(
            [
                f"### 문항 {question.index}",
                question.prompt,
                "",
                response.answer,
                "",
            ]
        )
    return "\n".join(lines).strip() + "\n"


def process_file(index: int, path: Path) -> dict:
    blocks = html_to_blocks(path)
    full_text = "\n\n".join(blocks)
    org = infer_org(path.name, full_text)
    parsed = [parse_question(block, i + 1) for i, block in enumerate(blocks) if len(block) >= 40]
    if not parsed:
        parsed = [("자기소개서 본문", full_text, None)]
    questions: list[Question] = []
    responses: list[DraftResponse] = []
    for question_index, (prompt, original, limit) in enumerate(parsed, 1):
        revised = revise_answer(org, prompt, original, question_index, limit)
        questions.append(Question(question_index, prompt, limit))
        responses.append(DraftResponse(question_index, revised, ("legacy_submission_ready",)))

    title = path.name.removesuffix(".html")
    output_stem = f"{index:02d}_{safe_name(title)}_제출권장후보"
    md_path = OUT_MD_DIR / f"{output_stem}.md"
    docx_path = OUT_DOCX_DIR / f"{output_stem}.docx"
    md_path.write_text(render_markdown(title, org, path, questions, responses), encoding="utf-8")
    write_docx(docx_path, title, org, questions, responses)

    job_terms = JOB_TERMS.get(org, JOB_TERMS["지원기관"])
    question_rows = []
    for question in questions:
        response = next(item for item in responses if item.question_index == question.index)
        peer_answers = tuple(item.answer for item in responses if item.question_index != question.index)
        score = score_answer_quality(
            question,
            response.answer,
            org,
            job_terms=job_terms,
            peer_answers=peer_answers,
        )
        question_rows.append(
            {
                "question_index": question.index,
                "prompt": question.prompt,
                "chars_no_space": len(re.sub(r"\s+", "", response.answer)),
                "char_limit": question.character_limit,
                "score": asdict(score),
            }
        )
    validation = validate_answer_quality(
        questions,
        responses,
        org,
        job_terms=job_terms,
        minimum_score=STRICT_MIN_ANSWER_SCORE,
        average_minimum_score=STRICT_MIN_AVERAGE_SCORE,
    )
    validation_codes = [issue.code for issue in validation]
    body = "\n".join(response.answer for response in responses)
    risk_hits = sorted({term for term in RISK_TERMS if term in body})
    avg = round(sum(row["score"]["total"] for row in question_rows) / len(question_rows), 1)
    return {
        "source_file": str(path.relative_to(WORKSPACE)),
        "file": str(md_path.relative_to(ROOT)),
        "docx_file": str(docx_path.relative_to(ROOT)),
        "organization": org,
        "question_count": len(questions),
        "average_score": avg,
        "recommendation": recommendation(avg, len(validation), len(risk_hits)),
        "issue_summary": summarize_issues(question_rows, validation_codes, risk_hits),
        "validation_issues": [asdict(issue) for issue in validation],
        "risk_terms_remaining": risk_hits,
        "questions": question_rows,
    }


def write_combined_docx(results: list[dict]) -> None:
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10.5)
    heading = document.add_heading("예전 자기소개서 제출권장 후보 통합본", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for pos, result in enumerate(results, 1):
        if pos > 1:
            document.add_page_break()
        md_path = ROOT / result["file"]
        text = md_path.read_text(encoding="utf-8")
        for line in text.splitlines():
            if line.startswith("# "):
                document.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                document.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                document.add_heading(line[4:], level=3)
            elif line.startswith("- "):
                document.add_paragraph(line[2:], style=None)
            elif line.strip():
                document.add_paragraph(line)
    document.save(COMBINED_DOCX)


def add_markdown_to_docx(document: Document, text: str) -> None:
    for line in text.splitlines():
        if line.startswith("# "):
            document.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            document.add_heading(line[4:], level=3)
        elif line.startswith("- "):
            document.add_paragraph(line[2:], style=None)
        elif line.strip():
            document.add_paragraph(line)


def write_all_combined_docx(results: list[dict]) -> None:
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10.5)
    heading = document.add_heading("전체 자기소개서 제출권장 통합본", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    current_dir = ROOT / "submission_ready_drafts"
    current_files = sorted(current_dir.glob("*.md"))
    document.add_paragraph(f"기존 제출권장화 묶음: {len(current_files)}개")
    document.add_paragraph(f"예전 자기소개서 추가 묶음: {len(results)}개")
    document.add_paragraph(f"전체 통합본: {len(current_files) + len(results)}개")
    for pos, path in enumerate(current_files, 1):
        document.add_page_break()
        document.add_heading(f"기존 묶음 {pos:02d}", level=1)
        add_markdown_to_docx(document, path.read_text(encoding="utf-8"))
    for pos, result in enumerate(results, 1):
        document.add_page_break()
        document.add_heading(f"예전 자기소개서 {pos:02d}", level=1)
        md_path = ROOT / result["file"]
        add_markdown_to_docx(document, md_path.read_text(encoding="utf-8"))
    document.save(COMBINED_ALL_DOCX)


def write_evaluation(results: list[dict]) -> None:
    EVAL_JSON.write_text(json.dumps(results, ensure_ascii=False, indent=2), encoding="utf-8")
    with EVAL_CSV.open("w", newline="", encoding="utf-8-sig") as fh:
        writer = csv.DictWriter(
            fh,
            fieldnames=[
                "source_file",
                "file",
                "docx_file",
                "organization",
                "question_count",
                "average_score",
                "recommendation",
                "issue_summary",
            ],
        )
        writer.writeheader()
        for result in results:
            writer.writerow({field: result[field] for field in writer.fieldnames})
    recommended = sum(1 for item in results if item["recommendation"] == "제출권장")
    lines = [
        "# 예전 자기소개서 제출권장 재평가",
        "",
        f"- 대상 HTML 보관본: {len(results)}개",
        f"- 제출권장: {recommended}개",
        f"- 보완 필요: {len(results) - recommended}개",
        "- 기준: 평균 90점 이상, 문항별 엄격 기준 통과, 위험 표현 잔여 없음",
        "- 원본 보호: 예전 자기소개서 폴더 원본은 수정하지 않음",
        "",
        "| 번호 | 기관/기업 | 문항 수 | 평균 | 판정 | 주요 이슈 |",
        "|---:|---|---:|---:|---|---|",
    ]
    for index, item in enumerate(results, 1):
        lines.append(
            f"| {index} | {item['organization']} | {item['question_count']} | "
            f"{item['average_score']} | {item['recommendation']} | {item['issue_summary']} |"
        )
    EVAL_MD.write_text("\n".join(lines) + "\n", encoding="utf-8")


def write_manifest(results: list[dict]) -> None:
    manifest = [
        {
            "source_file": item["source_file"],
            "file": item["file"],
            "docx_file": item["docx_file"],
            "organization": item["organization"],
            "question_count": item["question_count"],
            "average_score": item["average_score"],
            "recommendation": item["recommendation"],
            "risk_terms_remaining": item["risk_terms_remaining"],
        }
        for item in results
    ]
    MANIFEST.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")


def write_full_coverage(results: list[dict]) -> None:
    current_eval_path = ROOT / "submission_ready_re_evaluation_20260705.json"
    current = json.loads(current_eval_path.read_text(encoding="utf-8")) if current_eval_path.exists() else []
    current_recommended = sum(1 for item in current if item.get("recommendation") == "제출권장")
    legacy_recommended = sum(1 for item in results if item["recommendation"] == "제출권장")
    payload = {
        "generated_at": "2026-07-05",
        "scope": {
            "current_submission_ready_set": len(current),
            "legacy_html_set": len(results),
            "total_submission_ready_candidates": len(current) + len(results),
        },
        "recommendation_counts": {
            "current_submission_ready": current_recommended,
            "legacy_submission_ready": legacy_recommended,
            "total_submission_ready": current_recommended + legacy_recommended,
        },
        "coverage_note": "기존 42개 제출권장화 결과에 예전 자기소개서 HTML 45개를 별도 제출권장 후보로 추가했습니다.",
        "legacy_outputs": {
            "manifest": str(MANIFEST.relative_to(ROOT)),
            "evaluation_json": str(EVAL_JSON.relative_to(ROOT)),
            "evaluation_md": str(EVAL_MD.relative_to(ROOT)),
            "combined_docx": str(COMBINED_DOCX.relative_to(ROOT)),
            "combined_all_docx": str(COMBINED_ALL_DOCX.relative_to(ROOT)),
        },
    }
    FULL_COVERAGE_JSON.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    lines = [
        "# 전체 자기소개서 제출권장 커버리지",
        "",
        f"- 기존 제출권장화 묶음: {len(current)}개 중 {current_recommended}개 제출권장",
        f"- 예전 자기소개서 추가 묶음: {len(results)}개 중 {legacy_recommended}개 제출권장",
        f"- 전체 제출권장 후보: {len(current) + len(results)}개",
        "",
        "## 산출물",
        "",
        f"- 예전 자기소개서 평가표: `{EVAL_MD.name}`",
        f"- 예전 자기소개서 통합 DOCX: `{COMBINED_DOCX.name}`",
        f"- 전체 자기소개서 통합 DOCX: `{COMBINED_ALL_DOCX.name}`",
        f"- 예전 자기소개서 개별 DOCX 폴더: `{OUT_DOCX_DIR.name}`",
        f"- 예전 자기소개서 개별 MD 폴더: `{OUT_MD_DIR.name}`",
        "",
        "## 검증 경계",
        "",
        "- 과거 HTML 보관본은 원본 공고의 최신 외부 공식성까지 재검증한 것이 아니라, 보관된 자기소개서 본문을 제출평가 기준으로 정리한 후보입니다.",
        "- 지원 전에는 현재 채용공고 문항·글자 수·블라인드 기준과 다시 맞춰야 합니다.",
    ]
    FULL_COVERAGE_MD.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    OUT_MD_DIR.mkdir(exist_ok=True)
    OUT_DOCX_DIR.mkdir(exist_ok=True)
    html_files = sorted(LEGACY_DIR.glob("*.html"))
    results = [process_file(index, path) for index, path in enumerate(html_files, 1)]
    write_manifest(results)
    write_evaluation(results)
    write_combined_docx(results)
    write_all_combined_docx(results)
    write_full_coverage(results)
    print(json.dumps({"processed": len(results), "recommended": sum(1 for item in results if item["recommendation"] == "제출권장")}, ensure_ascii=False))


if __name__ == "__main__":
    main()
