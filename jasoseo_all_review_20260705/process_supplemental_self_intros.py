from __future__ import annotations

import csv
import json
import re
from dataclasses import asdict
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from career_pipeline.models import DraftResponse, Question
from career_pipeline.quality import (
    STRICT_MIN_ANSWER_SCORE,
    STRICT_MIN_AVERAGE_SCORE,
    score_answer_quality,
    validate_answer_quality,
)

from process_legacy_self_intros import (
    JOB_TERMS,
    RISK_TERMS,
    add_markdown_to_docx,
    infer_org,
    parse_limit,
    recommendation,
    render_markdown,
    revise_answer,
    safe_name,
    summarize_issues,
    write_docx,
)


ROOT = Path(__file__).resolve().parent
WORKSPACE = ROOT.parent
OUT_MD_DIR = ROOT / "supplemental_submission_ready_drafts"
OUT_DOCX_DIR = ROOT / "supplemental_submission_ready_drafts_docx"
MANIFEST = ROOT / "supplemental_submission_ready_manifest.json"
EVAL_JSON = ROOT / "supplemental_submission_ready_re_evaluation_20260705.json"
EVAL_CSV = ROOT / "supplemental_submission_ready_re_evaluation_20260705.csv"
EVAL_MD = ROOT / "supplemental_submission_ready_re_evaluation_20260705.md"
COMBINED_DOCX = ROOT / "supplemental_self_intro_submission_ready_20260705.docx"
COMPLETE_DOCX = ROOT / "all_self_intro_submission_ready_complete_20260705.docx"
COMPLETE_JSON = ROOT / "all_self_intro_coverage_complete_20260705.json"
COMPLETE_MD = ROOT / "all_self_intro_coverage_complete_20260705.md"


def read_docx(path: Path) -> list[str]:
    doc = Document(str(path))
    return [para.text.strip().replace("\xa0", " ") for para in doc.paragraphs if para.text.strip()]


def read_txt(path: Path) -> list[str]:
    text = path.read_text(encoding="utf-8", errors="ignore").replace("\xa0", " ")
    return [line.strip() for line in text.splitlines() if line.strip()]


def is_prompt(line: str) -> bool:
    compact = line.replace(" ", "")
    if re.search(r"^\d+[\.\)]", compact):
        return True
    if re.search(r"\d{2,4}\s*(자|byte)", line, re.IGNORECASE):
        return True
    if len(line) <= 260 and any(
        cue in compact
        for cue in (
            "지원동기",
            "기술해",
            "서술하",
            "말씀해",
            "작성하",
            "본인의",
            "경험을",
            "역량",
            "입사후",
            "포부",
        )
    ):
        return True
    return False


def group_questions(lines: list[str]) -> list[tuple[str, str, int | None]]:
    full_text = "\n".join(lines)
    limit = parse_limit(full_text)
    if len(lines) >= 1:
        return [("자기소개서 보완 통합본", full_text, limit)]

    groups: list[tuple[str, list[str]]] = []
    current_prompt = ""
    current_answer: list[str] = []
    header_lines: list[str] = []
    for line in lines:
        if line.startswith("마감일") or line.startswith("지원 기업") or line.startswith("직무:"):
            header_lines.append(line)
            continue
        if is_prompt(line):
            if current_prompt:
                groups.append((current_prompt, current_answer))
            current_prompt = line
            current_answer = []
            continue
        if current_prompt:
            current_answer.append(line)
        elif len(line) > 40:
            header_lines.append(line)
    if current_prompt:
        groups.append((current_prompt, current_answer))
    if not groups:
        text = "\n".join(lines)
        return [("자기소개서 보완 통합본", text, parse_limit(text))]
    parsed: list[tuple[str, str, int | None]] = []
    for prompt, answer_lines in groups:
        answer = "\n".join(
            line for line in answer_lines if not re.match(r"^\[[^\]]{2,80}\]$", line.strip())
        ).strip()
        if not answer:
            answer = "\n".join(header_lines).strip() or prompt
        parsed.append((prompt, answer, parse_limit(prompt)))
    return parsed


def selected_relpaths() -> set[str]:
    selected_path = ROOT / "selected_self_intros.json"
    if not selected_path.exists():
        return set()
    data = json.loads(selected_path.read_text(encoding="utf-8"))
    return {item["relpath"] for item in data}


def supplemental_sources() -> list[Path]:
    selected = selected_relpaths()
    sources: list[Path] = []
    folder = WORKSPACE / "\uc0c8\ub85c\uc6b4 \uc790\uae30\uc18c\uac1c\uc11c"
    for path in sorted(folder.glob("*")):
        if path.suffix.lower() not in {".docx", ".txt"}:
            continue
        rel = str(path.relative_to(WORKSPACE))
        if rel not in selected:
            sources.append(path)
    feedback = WORKSPACE / "250616 2025\ub144 \uc0c1\ubc18\uae30 \uad6d\ubbfc\uc5f0\uae08\uacf5\ub2e8 \uc790\uae30\uc18c\uac1c\uc11c \ud53c\ub4dc\ubc31.docx"
    if feedback.exists() and str(feedback.relative_to(WORKSPACE)) not in selected:
        sources.append(feedback)
    return sources


def extract_lines(path: Path) -> list[str]:
    if path.suffix.lower() == ".docx":
        return read_docx(path)
    return read_txt(path)


def process_file(index: int, path: Path) -> dict:
    lines = extract_lines(path)
    full_text = "\n".join(lines)
    org = infer_org(path.name, full_text)
    parsed = group_questions(lines)
    questions: list[Question] = []
    responses: list[DraftResponse] = []
    for question_index, (prompt, original, limit) in enumerate(parsed, 1):
        revised = revise_answer(org, prompt, original, question_index, limit)
        questions.append(Question(question_index, prompt, limit))
        responses.append(DraftResponse(question_index, revised, ("supplemental_submission_ready",)))
    title = path.name.removesuffix(path.suffix)
    output_stem = f"{index:02d}_{safe_name(title)}_제출권장후보"
    md_path = OUT_MD_DIR / f"{output_stem}.md"
    docx_path = OUT_DOCX_DIR / f"{output_stem}.docx"
    md_path.write_text(render_markdown(title, org, path, questions, responses), encoding="utf-8")
    write_docx(docx_path, title, org, questions, responses)

    job_terms = JOB_TERMS.get(org, JOB_TERMS["지원기관"])
    question_rows = []
    for question in questions:
        response = next(item for item in responses if item.question_index == question.index)
        peer_answers = tuple(item.answer for item in responses if item.question_index != question.index)
        score = score_answer_quality(
            question,
            response.answer,
            org,
            job_terms=job_terms,
            peer_answers=peer_answers,
        )
        question_rows.append(
            {
                "question_index": question.index,
                "prompt": question.prompt,
                "chars_no_space": len(re.sub(r"\s+", "", response.answer)),
                "char_limit": question.character_limit,
                "score": asdict(score),
            }
        )
    validation = validate_answer_quality(
        questions,
        responses,
        org,
        job_terms=job_terms,
        minimum_score=STRICT_MIN_ANSWER_SCORE,
        average_minimum_score=STRICT_MIN_AVERAGE_SCORE,
    )
    validation_codes = [issue.code for issue in validation]
    body = "\n".join(response.answer for response in responses)
    risk_hits = sorted({term for term in RISK_TERMS if term in body})
    avg = round(sum(row["score"]["total"] for row in question_rows) / len(question_rows), 1)
    return {
        "source_file": str(path.relative_to(WORKSPACE)),
        "file": str(md_path.relative_to(ROOT)),
        "docx_file": str(docx_path.relative_to(ROOT)),
        "organization": org,
        "question_count": len(questions),
        "average_score": avg,
        "recommendation": recommendation(avg, len(validation), len(risk_hits)),
        "issue_summary": summarize_issues(question_rows, validation_codes, risk_hits),
        "validation_issues": [asdict(issue) for issue in validation],
        "risk_terms_remaining": risk_hits,
        "questions": question_rows,
    }


def write_combined_docx(results: list[dict]) -> None:
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10.5)
    heading = document.add_heading("보충 자기소개서 제출권장 후보 통합본", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for pos, result in enumerate(results, 1):
        if pos > 1:
            document.add_page_break()
        add_markdown_to_docx(document, (ROOT / result["file"]).read_text(encoding="utf-8"))
    document.save(COMBINED_DOCX)


def write_evaluation(results: list[dict]) -> None:
    EVAL_JSON.write_text(json.dumps(results, ensure_ascii=False, indent=2), encoding="utf-8")
    MANIFEST.write_text(
        json.dumps(
            [
                {
                    "source_file": item["source_file"],
                    "file": item["file"],
                    "docx_file": item["docx_file"],
                    "organization": item["organization"],
                    "question_count": item["question_count"],
                    "average_score": item["average_score"],
                    "recommendation": item["recommendation"],
                    "risk_terms_remaining": item["risk_terms_remaining"],
                }
                for item in results
            ],
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )
    with EVAL_CSV.open("w", newline="", encoding="utf-8-sig") as fh:
        writer = csv.DictWriter(
            fh,
            fieldnames=[
                "source_file",
                "file",
                "docx_file",
                "organization",
                "question_count",
                "average_score",
                "recommendation",
                "issue_summary",
            ],
        )
        writer.writeheader()
        for result in results:
            writer.writerow({field: result[field] for field in writer.fieldnames})
    recommended = sum(1 for item in results if item["recommendation"] == "제출권장")
    lines = [
        "# 보충 자기소개서 제출권장 재평가",
        "",
        f"- 대상 후보: {len(results)}개",
        f"- 제출권장: {recommended}개",
        f"- 보완 필요: {len(results) - recommended}개",
        "- 기준: 평균 90점 이상, 문항별 엄격 기준 통과, 위험 표현 잔여 없음",
        "",
        "| 번호 | 원본 | 기관/기업 | 문항 수 | 평균 | 판정 | 주요 이슈 |",
        "|---:|---|---|---:|---:|---|---|",
    ]
    for index, item in enumerate(results, 1):
        lines.append(
            f"| {index} | {item['source_file']} | {item['organization']} | {item['question_count']} | "
            f"{item['average_score']} | {item['recommendation']} | {item['issue_summary']} |"
        )
    EVAL_MD.write_text("\n".join(lines) + "\n", encoding="utf-8")


def write_complete_docx(results: list[dict]) -> None:
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10.5)
    heading = document.add_heading("전체 자기소개서 제출권장 완전 통합본", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    groups = [
        ("기존 제출권장화 묶음", ROOT / "submission_ready_drafts"),
        ("예전 자기소개서 묶음", ROOT / "legacy_submission_ready_drafts"),
        ("보충 후보 묶음", ROOT / "supplemental_submission_ready_drafts"),
    ]
    total = sum(len(list(path.glob("*.md"))) for _, path in groups)
    document.add_paragraph(f"전체 통합본: {total}개")
    for label, folder in groups:
        files = sorted(folder.glob("*.md"))
        for pos, path in enumerate(files, 1):
            document.add_page_break()
            document.add_heading(f"{label} {pos:02d}", level=1)
            add_markdown_to_docx(document, path.read_text(encoding="utf-8"))
    document.save(COMPLETE_DOCX)


def write_complete_coverage(results: list[dict]) -> None:
    current = json.loads((ROOT / "submission_ready_re_evaluation_20260705.json").read_text(encoding="utf-8"))
    legacy = json.loads((ROOT / "legacy_submission_ready_re_evaluation_20260705.json").read_text(encoding="utf-8"))
    current_recommended = sum(1 for item in current if item["recommendation"] == "제출권장")
    legacy_recommended = sum(1 for item in legacy if item["recommendation"] == "제출권장")
    supplemental_recommended = sum(1 for item in results if item["recommendation"] == "제출권장")
    total = len(current) + len(legacy) + len(results)
    payload = {
        "generated_at": "2026-07-05",
        "scope": {
            "current_submission_ready_set": len(current),
            "legacy_html_set": len(legacy),
            "supplemental_candidate_set": len(results),
            "total_submission_ready_candidates": total,
        },
        "recommendation_counts": {
            "current_submission_ready": current_recommended,
            "legacy_submission_ready": legacy_recommended,
            "supplemental_submission_ready": supplemental_recommended,
            "total_submission_ready": current_recommended + legacy_recommended + supplemental_recommended,
        },
        "outputs": {
            "complete_docx": str(COMPLETE_DOCX.relative_to(ROOT)),
            "supplemental_evaluation_md": str(EVAL_MD.relative_to(ROOT)),
            "supplemental_combined_docx": str(COMBINED_DOCX.relative_to(ROOT)),
        },
        "verification_boundary": "지원 전에는 실제 최신 공고 문항, 글자 수, 블라인드 기준을 다시 대조해야 합니다.",
    }
    COMPLETE_JSON.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    lines = [
        "# 전체 자기소개서 제출권장 완전 커버리지",
        "",
        f"- 기존 제출권장화 묶음: {len(current)}개 중 {current_recommended}개 제출권장",
        f"- 예전 자기소개서 묶음: {len(legacy)}개 중 {legacy_recommended}개 제출권장",
        f"- 보충 후보 묶음: {len(results)}개 중 {supplemental_recommended}개 제출권장",
        f"- 전체 제출권장 후보: {total}개 중 {current_recommended + legacy_recommended + supplemental_recommended}개 제출권장",
        "",
        "## 산출물",
        "",
        f"- 전체 완전 통합 DOCX: `{COMPLETE_DOCX.name}`",
        f"- 보충 후보 평가표: `{EVAL_MD.name}`",
        f"- 보충 후보 통합 DOCX: `{COMBINED_DOCX.name}`",
        "",
        "## 검증 경계",
        "",
        "- 원본 문서는 수정하지 않았고, 제출권장 후보본을 별도 산출했습니다.",
        "- 실제 지원 전에는 최신 공고의 문항, 글자 수, 블라인드 기준과 다시 대조해야 합니다.",
    ]
    COMPLETE_MD.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    OUT_MD_DIR.mkdir(exist_ok=True)
    OUT_DOCX_DIR.mkdir(exist_ok=True)
    sources = supplemental_sources()
    results = [process_file(index, path) for index, path in enumerate(sources, 1)]
    write_evaluation(results)
    write_combined_docx(results)
    write_complete_docx(results)
    write_complete_coverage(results)
    print(
        json.dumps(
            {
                "processed": len(results),
                "recommended": sum(1 for item in results if item["recommendation"] == "제출권장"),
            },
            ensure_ascii=False,
        )
    )


if __name__ == "__main__":
    main()
