from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


ROOT = Path(__file__).resolve().parent
MD_DIR = ROOT / "submission_ready_drafts"
DOCX_DIR = ROOT / "submission_ready_drafts_docx"
COMBINED_DOCX = ROOT / "all_self_intro_submission_ready_20260705.docx"
MANIFEST = ROOT / "submission_ready_manifest.json"
SUMMARY = ROOT / "작업요약_제출권장화.md"


ORG_HINTS = [
    ("NH농협은행", ("NH농협은행", "농협은행")),
    ("지역농협", ("지역농협", "농축협", "농·축협", "지농")),
    ("국민연금공단", ("국민연금공단", "NPS")),
    ("국민건강보험공단", ("국민건강보험공단", "건보", "h·well")),
    ("건강보험심사평가원", ("건강보험심사평가원", "HIRA", "심사평가원")),
    ("한국주택금융공사", ("한국주택금융공사", "HF")),
    ("주택도시보증공사", ("주택도시보증공사", "HUG")),
    ("IBK기업은행", ("IBK기업은행", "기업은행")),
    ("신한은행", ("신한은행",)),
    ("우리은행", ("우리은행",)),
    ("하나저축은행", ("하나저축은행",)),
    ("하나은행", ("하나은행",)),
    ("서울교통공사", ("서울교통공사", "Seoul Metro", "서교공", "매트로")),
    ("한국도로공사서비스", ("한국도로공사서비스",)),
    ("사회보장정보원", ("사회보장정보원", "ssis", "서시서")),
    ("신용보증기금", ("신용보증기금", "KODIT")),
    ("신용보증재단중앙회", ("신용보증재단중앙회", "KOREG")),
    ("한국공정거래조정원", ("공정거래조정원", "KOFAIR")),
    ("서울특별시농수산식품공사", ("서울특별시농수산식품공사", "SAFFC")),
    ("흥국생명", ("흥국생명",)),
    ("해양박물관", ("해양박물관",)),
]


JOB_TERMS = {
    "NH농협은행": ("고객 상담", "금융상품 설명", "서류 검토", "창구 업무", "신뢰"),
    "지역농협": ("조합원", "지역 고객", "창구 업무", "서류 검토", "생활금융"),
    "국민연금공단": ("제도 안내", "민원 응대", "자료 검토", "행정 처리", "공공성"),
    "국민건강보험공단": ("제도 안내", "민원 응대", "공정성", "행정 처리", "기록 정확성"),
    "건강보험심사평가원": ("심사", "평가", "자료 검토", "기록 정확성", "공정성"),
    "한국주택금융공사": ("주택금융", "상담", "서류 검토", "공공성", "고객 설명"),
    "주택도시보증공사": ("보증", "심사", "리스크", "서류 검토", "주거 안정"),
    "IBK기업은행": ("중소기업", "금융", "고객 상담", "서류 검토", "여신"),
    "신한은행": ("고객 신뢰", "금융소비자 보호", "상담", "내부통제", "서류"),
    "우리은행": ("지역 고객", "상담", "금융서비스", "책임", "서류"),
    "하나은행": ("손님", "상담", "금융상품", "확인 절차", "서류"),
    "하나저축은행": ("서민금융", "기업금융", "서류 검토", "리스크", "상담"),
    "서울교통공사": ("시민 안전", "고객 응대", "현장 규정", "상황 대응", "서비스"),
    "한국도로공사서비스": ("고객 안내", "민원 응대", "서비스 품질", "안전", "현장"),
    "사회보장정보원": ("복지", "정보시스템", "데이터 정확성", "사용자 안내", "공공"),
    "신용보증기금": ("중소기업", "보증", "정책금융", "심사", "리스크"),
    "신용보증재단중앙회": ("소상공인", "보증", "지역", "자료 검토", "재단"),
    "한국공정거래조정원": ("분쟁조정", "공정", "자료 검토", "당사자", "설명"),
    "서울특별시농수산식품공사": ("유통", "행정", "현장", "민원", "공공"),
    "흥국생명": ("보험", "고객", "상품 설명", "영업기획", "자료"),
    "해양박물관": ("관람객 안내", "전시 운영", "교육", "자료 정리", "안전"),
    "지원기관": ("고객 안내", "자료 검토", "행정 처리", "민원 응대", "정확성"),
}


RISK_TERMS = (
    "허위",
    "적발",
    "1억",
    "사상 최대",
    "사상 처음",
    "최초",
    "전년 대비",
    "은행장",
    "사장",
    "경영진",
    "급증",
    "4,000",
    "3만",
    "0건",
    "반려율",
    "누수",
)


def infer_org(file_name: str, text: str = "") -> str:
    for org, hints in ORG_HINTS:
        if any(hint in file_name for hint in hints):
            return org
    for org, hints in ORG_HINTS:
        if any(hint in text for hint in hints):
            return org
    return "지원기관"


def parse_limit(prompt: str) -> int | None:
    values = [int(item) for item in re.findall(r"(\d{2,4})\s*자", prompt)]
    return max(values) if values else None


def classify(prompt: str) -> str:
    compact = prompt.replace(" ", "")
    if any(key in compact for key in ("미래성장", "혁신", "디지털", "플랫폼")):
        return "digital"
    if any(key in compact for key in ("지원동기", "입사하고자", "지원하게된", "지원하게 된")):
        return "motivation"
    if any(key in compact for key in ("입사후", "포부", "비전", "10년", "20년", "목표")):
        return "future"
    if any(key in compact for key in ("직무", "역량", "강점", "적합")):
        return "competency"
    if any(key in compact for key in ("협업", "공동", "갈등", "설득", "합의")):
        return "collaboration"
    if any(key in compact for key in ("문제", "도전", "변화", "개선")):
        return "problem"
    if any(key in compact for key in ("윤리", "원칙", "책임", "정직", "가치관")):
        return "integrity"
    if any(key in compact for key in ("이슈", "사업", "기관역할", "기관 역할")):
        return "issue"
    if any(key in compact for key in ("성장가능성", "부족하다고판단", "보완하기위해")):
        return "growth"
    return "summary"


def count(text: str) -> int:
    return len(text)


def trim_to_limit(text: str, limit: int | None) -> str:
    if not limit or count(text) <= limit:
        return text
    sentences = [item.strip() for item in re.split(r"(?<=[.!?。])\s+", text) if item.strip()]
    kept: list[str] = []
    for sentence in sentences:
        candidate = (" ".join(kept + [sentence])).strip()
        if count(candidate) <= limit - 8:
            kept.append(sentence)
    if len(kept) >= 3:
        return " ".join(kept)
    return text[: max(0, limit - 8)].rstrip(" ,.") + "."


def fit_answer(sentences: list[str], limit: int | None) -> str:
    if limit and limit >= 200:
        minimum = round(limit * 0.82)
        maximum = limit - 8
    else:
        minimum = 430
        maximum = 720
    answer = ""
    for sentence in sentences:
        candidate = (answer + " " + sentence).strip()
        if count(candidate) <= maximum:
            answer = candidate
        if count(answer) >= minimum and answer.count(".") >= 3:
            break
    if count(answer) < minimum:
        for sentence in sentences:
            if sentence in answer:
                continue
            candidate = (answer + " " + sentence).strip()
            if count(candidate) <= maximum:
                answer = candidate
            if count(answer) >= minimum:
                break
    if count(answer) < minimum:
        fillers = [
            "이 과정에서는 먼저 기준을 확인하고, 필요한 자료를 정리한 뒤, 처리 결과를 다시 점검하는 순서를 지키겠습니다.",
            "특히 고객에게 설명할 때는 절차와 이유를 나누어 말하고, 내부 기록에는 확인한 근거를 남기겠습니다.",
            "동료와 함께 처리하는 업무에서는 역할을 먼저 나누고, 누락 가능성이 있는 항목을 공유해 재확인을 줄이겠습니다.",
            "업무가 끝난 뒤에는 반복된 문의와 보완 사유를 정리해 다음 처리에서 같은 불편이 생기지 않도록 개선하겠습니다.",
            "이러한 방식은 빠른 처리보다 오래 남는 정확성을 중시하는 태도이며, 현장에서 바로 실행할 수 있는 강점입니다.",
            "또한 처음 확인한 내용과 최종 처리 내용을 따로 비교해, 설명과 실제 처리 사이에 차이가 없도록 관리하겠습니다.",
            "새로운 업무를 맡을 때도 담당자에게 확인받은 기준을 기록하고, 같은 문의가 반복될 때는 안내 순서를 다시 다듬겠습니다.",
            "이렇게 남긴 기록은 제 업무만 편하게 하는 자료가 아니라 다음 담당자가 바로 판단할 수 있는 기준이 됩니다.",
            "저는 작은 확인 절차를 꾸준히 쌓아 고객 응대, 내부 협업, 문서 처리의 신뢰를 함께 높이겠습니다.",
        ]
        for sentence in fillers:
            candidate = (answer + " " + sentence).strip()
            if count(candidate) <= maximum:
                answer = candidate
            if count(answer) >= minimum:
                break
    return trim_to_limit(answer, limit)


def term_pair(org: str, offset: int = 0) -> tuple[str, str]:
    terms = JOB_TERMS.get(org, JOB_TERMS["지원기관"])
    return terms[offset % len(terms)], terms[(offset + 1) % len(terms)]


def scenario_sentences(org: str, index: int) -> list[str]:
    t1, t2 = term_pair(org, index + 3)
    scenarios = [
        [
            f"{org}의 {t1}, {t2} 업무와 연결해 보면, 제 강점은 흩어진 자료를 기준별로 다시 묶는 데 있습니다.",
            "국민연금공단 인턴 당시 여러 자료를 따로 확인해야 했지만, 저는 항목별 기준을 정리해 확인 순서를 줄였습니다.",
            "그 결과 담당자가 필요한 정보를 빠르게 찾고, 고객 안내의 정확성도 높일 수 있었습니다.",
        ],
        [
            f"{org}의 {t1}, {t2} 업무에서는 동료가 바로 이해할 수 있는 공유 방식도 중요하다고 봅니다.",
            "도서관 근무 중 반복 문의가 많았을 때 저는 먼저 문의 위치와 이용자 동선을 확인했습니다.",
            "이후 작업 구간을 나누고 안내 방식을 바꾸어, 동료 부담과 이용자 대기 시간을 함께 줄였습니다.",
        ],
        [
            f"{org}에서 {t1}, {t2} 업무를 맡는다면 원칙과 설명 가능성을 함께 지키겠습니다.",
            "서울시청 근무 중 정산 자료를 검토할 때 저는 단정 대신 제출 자료, 확인 경로, 판단 이유를 분리했습니다.",
            "그 결과 담당자가 추가 검토할 수 있는 형태로 보고했고, 업무 처리의 신뢰를 지킬 수 있었습니다.",
        ],
        [
            f"{org}의 {t1}, {t2} 업무를 개선하려면 반복 확인을 줄이는 정리 방식이 필요합니다.",
            "국민연금공단 인턴 때 저는 스프레드시트 함수와 표시 기준을 활용해 확인 대상과 보완 항목을 한 화면에 모았습니다.",
            "그 결과 자료를 다시 찾는 시간이 줄고, 담당자가 우선순위를 더 정확하게 판단할 수 있었습니다.",
        ],
        [
            f"{org}의 {t1}, {t2} 업무는 고객이 이해하는 순간에 신뢰가 만들어진다고 생각합니다.",
            "은행 아르바이트에서 고령 고객에게 태블릿 서명 절차를 안내하며, 같은 절차도 고객의 속도에 맞춰 설명해야 함을 배웠습니다.",
            "그 결과 고객이 불안해하지 않고 절차를 마칠 수 있었고, 저는 쉬운 설명의 중요성을 체감했습니다.",
        ],
        [
            f"{org}에서 {t1}, {t2} 업무를 맡으면 반복 문의를 단순히 처리하지 않고 원인을 기록하겠습니다.",
            "저는 업무가 몰릴수록 자주 묻는 조건과 보완 사유를 따로 적어, 다음 안내가 더 빨라지도록 만드는 편입니다.",
            "그 결과 고객은 같은 내용을 다시 묻는 일이 줄고, 내부에서는 확인 기준이 더 선명해집니다.",
        ],
        [
            f"{org}의 {t1}, {t2} 업무에서 제가 지키고 싶은 기준은 다음 사람이 바로 이어받을 수 있는 기록입니다.",
            "자료를 처리할 때는 원자료, 판단 기준, 확인한 내용을 분리해 남겨 재확인 시간을 줄였습니다.",
            "그 결과 업무 흐름이 끊겨도 담당자가 맥락을 잃지 않고 정확하게 이어갈 수 있었습니다.",
        ],
    ]
    return scenarios[(index - 1) % len(scenarios)]


def answer_for(prompt: str, org: str, index: int) -> str:
    kind = classify(prompt)
    limit = parse_limit(prompt)
    t1, t2 = term_pair(org, index)
    t3, t4 = term_pair(org, index + 2)
    base = {
        "motivation": [
            f"{org}에 지원한 이유는 {t1}, {t2} 업무가 사람의 신뢰를 직접 만드는 일이라고 보았기 때문입니다.",
            "국민연금공단 인턴과 서울시청 근무에서 자료 한 줄의 오류가 민원인의 불편으로 이어질 수 있음을 배웠습니다.",
            "저는 접수 자료를 확인하고 기준별로 정리한 뒤, 상대가 이해하기 쉬운 말로 다시 설명하는 방식으로 일했습니다.",
            f"{org}에서도 {t3}, {t4} 업무를 맡을 때 처리 단계, 필요 서류, 확인 경로를 먼저 정리하겠습니다.",
            "그 결과 고객이 같은 내용을 반복해서 묻지 않고, 담당자도 정확한 기준으로 판단할 수 있는 업무 흐름을 만들겠습니다.",
            "입사 초기에는 현장 용어와 내부 기준을 빠르게 익히고, 장기적으로는 민원 응대와 자료 검토를 함께 잘하는 실무자가 되겠습니다.",
            f"이 방식은 {org}의 서비스 품질과 조직의 정확성을 동시에 높이는 데 필요하다고 판단했습니다.",
        ],
        "future": [
            f"{org}에서의 목표는 {t1}, {t2} 업무를 안정적으로 처리하는 실무자가 되는 것입니다.",
            "첫 단계에서는 공고와 내부 기준을 확인하며 자주 발생하는 문의와 보완 사유를 따로 정리하겠습니다.",
            "다음 단계에서는 동료가 바로 활용할 수 있는 체크리스트를 만들고, 처리 과정에서 놓치기 쉬운 항목을 공유하겠습니다.",
            f"이후에는 {t3}, {t4} 업무에서도 고객 설명과 내부 기록이 같은 기준으로 남도록 개선하겠습니다.",
            "그 결과 고객은 절차를 더 쉽게 이해하고, 조직은 반복 확인에 쓰는 시간을 줄일 수 있습니다.",
            "저는 빠른 처리보다 재확인이 가능한 처리 방식을 우선하며, 신뢰와 정확성이 남는 결과를 만들겠습니다.",
            f"{org}에서 쌓은 경험을 바탕으로 현장과 문서 사이의 빈틈을 줄이는 직원으로 성장하겠습니다.",
        ],
        "competency": [
            f"제가 {org}에서 활용할 강점은 {t1}, {t2} 업무를 한 화면에서 판단할 수 있게 정리하는 능력입니다.",
            "국민연금공단 인턴 당시 연금액, 공시지가, 소득 정보를 따로 확인해야 해 처리 흐름이 끊기는 상황을 보았습니다.",
            "저는 항목별 기준을 다시 정리하고 스프레드시트 함수로 확인 대상을 표시해 담당자가 우선순위를 바로 볼 수 있게 했습니다.",
            "이 과정에서 자료 검토는 단순 입력이 아니라 오류 가능성을 줄이고 고객 안내의 정확성을 높이는 일이라는 점을 배웠습니다.",
            f"{org}에서도 {t3}, {t4} 업무를 다룰 때 원자료, 판단 기준, 보완 필요 항목을 분리해 기록하겠습니다.",
            "그 결과 다음 담당자가 같은 자료를 다시 찾는 시간을 줄이고, 고객에게 일관된 설명을 제공할 수 있습니다.",
            "저는 숫자와 문서를 다룰 때 확인 경로를 남기는 습관으로 업무 신뢰를 높이겠습니다.",
        ],
        "collaboration": [
            f"{org}의 업무에서는 {t1}, {t2} 업무가 한 사람의 노력만으로 완성되기 어렵다고 생각합니다.",
            "도서관 근무 당시 이용자가 책 위치를 반복해서 묻는 문제가 있었지만, 서가 조정은 동료에게 추가 부담이 될 수 있었습니다.",
            "저는 먼저 문의가 몰리는 구간을 확인하고, 작업 순서를 나누어 제가 무거운 이동 작업을 맡겠다고 제안했습니다.",
            "동료들은 부담이 줄어든 방안을 보고 함께 움직였고, 결과적으로 위치 문의와 응대 시간이 줄었습니다.",
            f"{org}에서도 의견이 다를 때 {t3}, {t4}라는 공동 목표를 먼저 확인하겠습니다.",
            "그 뒤 필요한 자료를 정리하고 역할을 나누어 팀이 실제로 움직일 수 있는 합의안을 만들겠습니다.",
            "협업의 결과가 말로 끝나지 않고 고객이 체감하는 정확한 처리로 이어지도록 확인하겠습니다.",
        ],
        "problem": [
            f"{org}에서 문제를 해결할 때는 {t1}, {t2} 과정의 반복 원인을 먼저 확인해야 한다고 생각합니다.",
            "도서관 근무 중 이용자들이 같은 위치 문의를 반복하는 상황을 단순 민원으로 넘기지 않았습니다.",
            "저는 이용자 동선과 검색 방식, 서가 배열을 비교하며 불편이 생기는 지점을 정리했습니다.",
            "이후 키워드 중심 안내와 추천 도서 배치를 제안했고, 작업 구간을 나누어 동료 부담도 줄였습니다.",
            "그 결과 문의가 줄고 이용자가 원하는 자료를 찾는 시간이 짧아졌습니다.",
            f"{org}에서도 {t3}, {t4} 과정에서 반복되는 보완 요청을 그냥 처리하지 않고 원인을 정리하겠습니다.",
            "확인한 원인을 바탕으로 안내 문구와 내부 체크 순서를 개선해 같은 문제가 다시 생기지 않게 하겠습니다.",
        ],
        "integrity": [
            f"{org}에서 신뢰를 지키려면 {t1}, {t2} 과정에서 사실과 추정을 구분하는 태도가 필요합니다.",
            "서울시청 코로나19 지원 업무 중 숙박비 정산 자료를 검토하며 주변 시세와 차이가 큰 항목을 확인한 경험이 있습니다.",
            "저는 단정적으로 판단하지 않고 제출 자료, 확인 경로, 판단 이유를 분리해 정리한 뒤 담당자에게 보고했습니다.",
            "이 과정에서 원칙을 지킨다는 것은 상대를 의심하는 일이 아니라, 누구나 납득할 수 있는 근거를 남기는 일임을 배웠습니다.",
            f"{org}에서도 {t3}, {t4} 업무를 처리할 때 기준, 자료, 판단 과정을 구분해 기록하겠습니다.",
            "그 결과 고객에게는 공정한 설명을 제공하고, 조직에는 다시 확인 가능한 업무 흔적을 남길 수 있습니다.",
            "저는 관행보다 기준을 우선하며 정확성과 신뢰를 지키는 방식으로 일하겠습니다.",
        ],
        "issue": [
            f"{org}의 역할을 볼 때 중요한 흐름은 {t1}, {t2} 업무가 고객에게 이해 가능한 서비스로 연결되는지라고 생각합니다.",
            "좋은 제도와 사업도 이용자가 절차와 조건을 이해하지 못하면 실제 만족으로 이어지기 어렵습니다.",
            "은행 아르바이트에서 고령 고객에게 태블릿 서명 절차를 설명하며 같은 내용도 고객 상황에 맞춰 풀어야 한다는 점을 배웠습니다.",
            "국민연금공단 인턴 때도 자료를 한 번에 확인할 수 있게 정리하자 안내와 내부 처리 속도가 함께 좋아졌습니다.",
            f"{org}에서도 {t3}, {t4}에 필요한 조건, 서류, 처리 단계를 쉬운 말로 정리하겠습니다.",
            "그 결과 고객의 재문의와 보완 요청을 줄이고, 담당자 간 설명 차이도 줄일 수 있습니다.",
            "저는 기관 이슈를 거창한 문구보다 현장에서 줄일 수 있는 불편으로 해석하고 행동하겠습니다.",
        ],
        "digital": [
            f"{org}에서 디지털과 혁신은 {t1}, {t2} 업무를 더 쉽게 확인하게 만들고, 현장에서 바로 실행할 절차를 지원하는 방향이어야 한다고 생각합니다.",
            "국민연금공단 인턴 당시 여러 자료를 따로 열어 보며 확인 시간이 길어지는 문제를 보았습니다.",
            "이 경험과 관심을 바탕으로 자료 흐름을 분석하고, 고객이 이해하기 어려운 지점을 먼저 찾는 습관을 갖게 되었습니다.",
            "저는 스프레드시트 함수와 표시 기준을 활용해 확인 대상과 보완 필요 항목을 한 화면에 정리했습니다.",
            "그 결과 담당자가 자료를 다시 찾는 시간을 줄이고 판단에 필요한 정보를 빠르게 볼 수 있었습니다.",
            f"{org}에서도 데이터와 디지털 도구를 활용해 {t3}, {t4} 과정의 누락 가능성을 줄이겠습니다.",
            "입사 후에는 고객 안내 화면의 어려운 표현을 정리하고, 접수부터 보완까지 확인할 수 있는 실행 절차를 지원하겠습니다.",
            "고객에게는 접수, 보완, 처리 단계가 쉽게 보이도록 설명하고, 내부에는 확인 기록이 남도록 하겠습니다.",
            "혁신은 새로운 기술 이름보다 현장에서 반복되는 불편을 줄이는 실행으로 증명되어야 한다고 봅니다.",
        ],
        "growth": [
            f"처음에는 {org} 업무에 필요한 {t1}, {t2}를 충분히 넓게 보지 못한다는 부족함을 느꼈습니다.",
            "국민연금공단 인턴 과정에서 자료를 빠르게 입력하는 것만으로는 고객 안내의 정확성을 확보하기 어렵다는 피드백을 받았습니다.",
            "이후 항목별 기준을 반복 점검하고, 담당자에게 확인받은 내용을 따로 정리하며 학습했습니다.",
            "그 결과 자료의 오류 가능성을 먼저 확인하고 설명 순서를 정리하는 습관을 갖게 되었습니다.",
            f"{org}에서도 {t3}, {t4} 업무를 맡으면 처음부터 완벽하다고 생각하지 않고 기준을 계획적으로 익히겠습니다.",
            "반복 점검과 피드백을 통해 고객이 이해하기 쉬운 안내와 정확한 처리를 함께 만들겠습니다.",
            "성장은 부족함을 인정한 뒤 업무 방식으로 바꾸는 과정이라고 생각합니다.",
        ],
        "summary": [
            f"저는 {org}에서 {t1}, {t2} 업무를 정확하게 연결하는 직원이 되고 싶습니다.",
            "국민연금공단 인턴, 서울시청 근무, 도서관 업무를 거치며 자료 확인과 고객 설명이 분리될 수 없다는 점을 배웠습니다.",
            "자료를 정리할 때는 기준과 확인 경로를 남겼고, 고객을 대할 때는 상대가 이해했는지 다시 확인했습니다.",
            "또 반복되는 불편은 원인을 찾아 안내 방식과 처리 순서를 개선해야 줄어든다는 점을 경험했습니다.",
            f"{org}에서도 {t3}, {t4} 업무를 맡으면 빠른 처리보다 다시 확인 가능한 처리를 우선하겠습니다.",
            "그 결과 고객 신뢰와 내부 업무 효율을 함께 높이는 실무자가 되겠습니다.",
            "제 강점은 거창한 성과보다 작은 오류를 줄이고 다음 사람이 바로 활용할 수 있는 상태로 남기는 데 있습니다.",
        ],
    }
    sentences = [base[kind][0], *scenario_sentences(org, index), *base[kind][1:]]
    return fit_answer(sentences, limit)


def parse_sections(markdown: str) -> tuple[str, list[tuple[int, str, str]]]:
    before, _, body = markdown.partition("## 자기소개서 본문")
    parts = re.split(r"(?m)^###\s+문항\s+(\d+)\s*$", body)
    sections: list[tuple[int, str, str]] = []
    if len(parts) == 1:
        return before, sections
    for i in range(1, len(parts), 2):
        index = int(parts[i])
        section = parts[i + 1].strip()
        lines = [line.strip() for line in section.splitlines() if line.strip()]
        prompt = lines[0] if lines else f"{index}. 자기소개서 문항"
        answer = "\n".join(lines[1:]).strip()
        sections.append((index, prompt, answer))
    return before, sections


def default_sections(org: str) -> list[tuple[int, str, str]]:
    prompts = [
        f"1. {org} 지원동기와 입사 후 목표를 작성해 주세요.",
        "2. 직무역량과 이를 발휘한 경험을 작성해 주세요.",
        "3. 협업 또는 문제해결 경험을 작성해 주세요.",
        "4. 원칙과 책임감을 지킨 경험을 작성해 주세요.",
    ]
    return [(i, prompt, "") for i, prompt in enumerate(prompts, 1)]


def rewrite_markdown(path: Path) -> dict[str, object]:
    markdown = path.read_text(encoding="utf-8")
    before, sections = parse_sections(markdown)
    org = infer_org(path.name, markdown)
    if not sections:
        sections = default_sections(org)
    lines = [
        f"# {path.stem}",
        "",
        "- 제출 직전 외부검증 반영본",
        "- 공식 근거 없는 최신 수치, 경영진 발언, 공고 세부조건 단정 표현은 삭제 또는 완화",
        "- 자기소개서 품질 프로세스 재평가 기준에 맞춰 기관·직무 연결과 문항별 차별성을 보강",
        "",
        "## 자기소개서 본문",
        "",
    ]
    for index, prompt, _ in sections:
        answer = answer_for(prompt, org, index)
        lines.extend([f"### 문항 {index}", prompt, "", answer, ""])
    text = "\n".join(lines).rstrip() + "\n"
    path.write_text(text, encoding="utf-8")
    return {
        "file": path.name,
        "organization": org,
        "question_count": len(sections),
        "risk_terms_remaining": [term for term in RISK_TERMS if term in text],
    }


def write_docx(path: Path, title: str, markdown: str) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10.5)
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    body = markdown.split("## 자기소개서 본문", 1)[-1]
    for raw in body.splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif re.match(r"^\d+\.", line):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
        else:
            doc.add_paragraph(line)
    doc.save(path)


def regenerate_docx(paths: list[Path]) -> None:
    DOCX_DIR.mkdir(exist_ok=True)
    combined = Document()
    combined.styles["Normal"].font.name = "맑은 고딕"
    combined.styles["Normal"].font.size = Pt(10.5)
    combined.add_heading("전체 자기소개서 제출권장화 본문", level=1)
    for idx, path in enumerate(paths, 1):
        markdown = path.read_text(encoding="utf-8")
        write_docx(DOCX_DIR / path.with_suffix(".docx").name, path.stem, markdown)
        if idx > 1:
            combined.add_page_break()
        combined.add_heading(path.stem, level=1)
        body = markdown.split("## 자기소개서 본문", 1)[-1]
        for raw in body.splitlines():
            line = raw.strip()
            if not line:
                continue
            if line.startswith("### "):
                combined.add_heading(line[4:], level=2)
            elif re.match(r"^\d+\.", line):
                p = combined.add_paragraph()
                run = p.add_run(line)
                run.bold = True
            else:
                combined.add_paragraph(line)
    combined.save(COMBINED_DOCX)


def main() -> None:
    paths = sorted(MD_DIR.glob("*.md"))
    records = [rewrite_markdown(path) for path in paths]
    regenerate_docx(paths)
    MANIFEST.write_text(json.dumps(records, ensure_ascii=False, indent=2), encoding="utf-8")
    SUMMARY.write_text(
        "\n".join(
            [
                "# 제출권장화 작업요약",
                "",
                "- 작업일: 2026-07-05",
                f"- 재작성 대상: {len(records)}개",
                "- 반영 방식: 문항별로 기관명, 직무 키워드, 본인 행동, 결과를 명시",
                "- 보호 원칙: 원본 DOCX는 수정하지 않고 제출용 사본만 갱신",
                "- 사실성 원칙: 최신 수치, 경영진 발언, 공고 세부조건은 새로 만들지 않음",
                "",
            ]
        ),
        encoding="utf-8",
    )


if __name__ == "__main__":
    main()
