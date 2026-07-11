from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


ROOT = Path(__file__).resolve().parent
SOURCE_DIR = ROOT / "revised_drafts"
OUT_MD_DIR = ROOT / "submission_ready_drafts"
OUT_DOCX_DIR = ROOT / "submission_ready_drafts_docx"
MANIFEST_PATH = ROOT / "submission_ready_manifest.json"
COMBINED_DOCX = ROOT / "all_self_intro_submission_ready_20260705.docx"
SUMMARY_PATH = ROOT / "작업요약_외부검증반영.md"

OUT_MD_DIR.mkdir(exist_ok=True)
OUT_DOCX_DIR.mkdir(exist_ok=True)

RISKY_REPLACEMENTS = [
    (r"허위\s*청구를?\s*적발했습니다", "이상 징후를 확인해 담당자에게 보고했습니다"),
    (r"허위\s*청구를?\s*적발", "이상 징후를 확인해 보고"),
    (r"허위\s*청구", "이상 징후"),
    (r"허위\s*증빙", "증빙상 이상 징후"),
    (r"적발했습니다", "확인해 보고했습니다"),
    (r"적발", "확인"),
    (r"1억\s*원의?\s*예산\s*누수를?\s*막았습니다", "예산 집행상 추가 확인이 필요한 부분을 정리했습니다"),
    (r"1억\s*원의?\s*예산\s*누수를?\s*막아낸", "예산 집행상 추가 확인이 필요한 부분을 정리한"),
    (r"1억\s*원을?\s*지켜낸", "예산 집행상 추가 확인이 필요한 부분을 정리한"),
    (r"사상\s*최대", "관련 수요가 커지는 상황"),
    (r"사상\s*처음", "관련 수요가 커지는 상황"),
    (r"전년\s*대비\s*\d+%?\s*급증", "관련 수요가 늘어나는 흐름"),
    (r"급증", "증가"),
    (r"은행장[은는이가]?\s*[^.]*\.", ""),
    (r"사장[은는이가]?\s*[^.]*\.", ""),
    (r"경영진[은는이가]?\s*[^.]*\.", ""),
    (r"공식\s*발표\s*수치상", "공식 자료상"),
    (r"반려율\s*0%", "서류 보완을 줄이는 것"),
    (r"오류\s*0건", "오류를 줄이는 것"),
]

RISKY_PATTERNS = [
    "허위",
    "적발",
    "1억",
    "사상 최대",
    "사상 처음",
    "최초",
    "전년 대비",
    "은행장",
    "사장",
    "경영진",
    "급증",
    "4,000",
    "3만",
    "0건",
    "반려율",
    "누수",
]


def apply_external_verification(text: str) -> str:
    result = text
    for pattern, repl in RISKY_REPLACEMENTS:
        result = re.sub(pattern, repl, result)
    result = re.sub(r"\n{3,}", "\n\n", result)
    result = re.sub(r"[ \t]{2,}", " ", result)
    return result.strip()


def extract_submission_body(text: str) -> str:
    marker = "## 개선본"
    if marker in text:
        body = text.split(marker, 1)[1]
    else:
        body = text
    body = apply_external_verification(body)
    return body.strip()


def add_md_header(source_name: str, body: str) -> str:
    title = source_name.removesuffix(".md")
    return (
        f"# {title}\n\n"
        "- 제출 직전 외부검증 반영본\n"
        "- 공식 근거 없는 최신 수치, 경영진 발언, 공고 세부조건 단정 표현은 삭제 또는 완화\n"
        "- 원본 DOCX는 수정하지 않고 제출용 사본만 생성\n\n"
        "## 자기소개서 본문\n\n"
        f"{body}\n"
    )


def write_docx(path: Path, title: str, body: str) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10.5)

    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for raw in body.splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif re.match(r"^\d+\.", line):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
        else:
            doc.add_paragraph(line)
    doc.save(path)


def count_no_space(text: str) -> int:
    return len(re.sub(r"\s+", "", text))


def has_risky(text: str) -> list[str]:
    return [p for p in RISKY_PATTERNS if p in text]


def main() -> None:
    manifest = []
    combined = Document()
    combined.styles["Normal"].font.name = "맑은 고딕"
    combined.styles["Normal"].font.size = Pt(10.5)
    combined.add_heading("전체 자기소개서 제출용 외부검증 반영본", level=1)

    for idx, src in enumerate(sorted(SOURCE_DIR.glob("*.md")), 1):
        source_text = src.read_text(encoding="utf-8")
        body = extract_submission_body(source_text)
        final_md = add_md_header(src.name, body)
        out_md = OUT_MD_DIR / src.name
        out_docx = OUT_DOCX_DIR / src.with_suffix(".docx").name

        out_md.write_text(final_md, encoding="utf-8")
        write_docx(out_docx, src.stem, body)

        if idx > 1:
            combined.add_page_break()
        combined.add_heading(src.stem, level=1)
        for raw in body.splitlines():
            line = raw.strip()
            if not line:
                continue
            if line.startswith("### "):
                combined.add_heading(line[4:], level=2)
            elif re.match(r"^\d+\.", line):
                p = combined.add_paragraph()
                run = p.add_run(line)
                run.bold = True
            else:
                combined.add_paragraph(line)

        manifest.append(
            {
                "no": idx,
                "source": str(src.relative_to(ROOT)),
                "submission_md": str(out_md.relative_to(ROOT)),
                "submission_docx": str(out_docx.relative_to(ROOT)),
                "chars_no_space": count_no_space(body),
                "risky_terms_remaining": has_risky(body),
            }
        )

    combined.save(COMBINED_DOCX)
    MANIFEST_PATH.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")

    risky_remaining = sum(1 for item in manifest if item["risky_terms_remaining"])
    SUMMARY_PATH.write_text(
        "\n".join(
            [
                "# 외부검증 반영 작업요약",
                "",
                "- 작업일: 2026-07-05",
                f"- 제출용 Markdown 생성: {len(manifest)}개",
                f"- 제출용 Word 생성: {len(manifest)}개",
                "- 통합 Word 생성: all_self_intro_submission_ready_20260705.docx",
                f"- 고위험 표현 잔여 문서: {risky_remaining}개",
                "- 원본 DOCX와 기존 revised_drafts 묶음은 수정하지 않음",
                "",
                "## 반영 기준",
                "",
                "- 공식 근거 없는 최신 수치, 경영진 발언, 공고 세부조건 단정 표현 삭제 또는 완화",
                "- 본인 경험의 금액·건수·적발 표현은 확인·정리·보고 중심으로 완화",
                "- 프레임 폴더의 작성 원칙은 두괄식, 상황-행동-결과, 직무 연결에만 반영",
                "",
            ]
        ),
        encoding="utf-8",
    )


if __name__ == "__main__":
    main()
