from __future__ import annotations

import re
import sys
import zipfile
from pathlib import Path

from docx import Document


COUNT_RE = re.compile(r"공백 제외 (\d+)자$")


def count_without_spaces(text: str) -> int:
    return len(re.sub(r"\s+", "", text))


def main() -> None:
    source_path = Path(sys.argv[1])
    final_path = Path(sys.argv[2])

    with zipfile.ZipFile(final_path) as archive:
        bad_member = archive.testzip()
        if bad_member:
            raise RuntimeError(f"손상된 OOXML 항목: {bad_member}")
        required = {"[Content_Types].xml", "word/document.xml", "word/styles.xml"}
        missing = required.difference(archive.namelist())
        if missing:
            raise RuntimeError(f"필수 OOXML 항목 누락: {sorted(missing)}")

    source = Document(source_path)
    final = Document(final_path)
    source_paragraphs = [p for p in source.paragraphs if p.text]
    final_paragraphs = [p for p in final.paragraphs if p.text]
    if len(source_paragraphs) != 10 or len(final_paragraphs) != 10:
        raise RuntimeError("문단 구조가 예상과 다릅니다.")
    if len(source.sections) != len(final.sections) or len(source.tables) != len(final.tables):
        raise RuntimeError("섹션 또는 표 구조가 달라졌습니다.")

    for index, (before, after) in enumerate(zip(source_paragraphs, final_paragraphs)):
        before_style = before.style.name if before.style else None
        after_style = after.style.name if after.style else None
        if before_style != after_style:
            raise RuntimeError(f"문단 {index}의 스타일이 달라졌습니다.")

    counts: list[int] = []
    answers: list[str] = []
    for section_index in range(5):
        heading = final_paragraphs[section_index * 2].text
        answer = final_paragraphs[section_index * 2 + 1].text
        match = COUNT_RE.search(heading)
        count = count_without_spaces(answer)
        if not match or int(match.group(1)) != count:
            raise RuntimeError(f"문항 {section_index + 1} 글자 수 표기가 맞지 않습니다.")
        if count > 500:
            raise RuntimeError(f"문항 {section_index + 1}이 500자를 초과했습니다.")
        counts.append(count)
        answers.append(answer)

    full_text = "\n".join(answers)
    required_facts = ["허위 거래 20건", "예산 4천만 원", "일정을 1주일 앞당겼"]
    for fact in required_facts:
        if fact not in full_text:
            raise RuntimeError(f"확인된 핵심 사실 누락: {fact}")

    forbidden_unverified = ["약 40%", "문의가 30%", "20% 늘", "상인 50명", "시장 5곳", "목표 150건"]
    for phrase in forbidden_unverified:
        if phrase in full_text:
            raise RuntimeError(f"미확인 수치가 포함됐습니다: {phrase}")

    print("zip_integrity=ok")
    print("structure=ok")
    print("styles_preserved=ok")
    print("counts=" + ",".join(map(str, counts)))
    print("fact_guard=ok")


if __name__ == "__main__":
    main()
