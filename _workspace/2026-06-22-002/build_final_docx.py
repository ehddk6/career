from __future__ import annotations

import json
import re
import sys
from pathlib import Path

from docx import Document


def replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def count_without_spaces(text: str) -> int:
    return len(re.sub(r"\s+", "", text))


def main() -> None:
    source = Path(sys.argv[1])
    answers_path = Path(sys.argv[2])
    output = Path(sys.argv[3])
    answers = json.loads(answers_path.read_text(encoding="utf-8"))
    if len(answers) != 5:
        raise RuntimeError("답변은 정확히 5개여야 합니다.")

    document = Document(source)
    paragraphs = [paragraph for paragraph in document.paragraphs if paragraph.text]
    if len(paragraphs) != 10:
        raise RuntimeError(f"비어 있지 않은 문단이 10개가 아니라 {len(paragraphs)}개입니다.")

    for index, answer in enumerate(answers):
        count = count_without_spaces(answer)
        if count > 500:
            raise RuntimeError(f"문항 {index + 1}이 공백 제외 500자를 초과했습니다: {count}")
        heading_prefix = paragraphs[index * 2].text.split("—", 1)[0].rstrip()
        replace_paragraph_text(paragraphs[index * 2], f"{heading_prefix} — 공백 제외 {count}자")
        replace_paragraph_text(paragraphs[index * 2 + 1], answer)
        print(f"section={index + 1} chars_without_spaces={count} chars_with_spaces={len(answer)}")

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    print(f"saved={output}")


if __name__ == "__main__":
    main()
