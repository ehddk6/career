from __future__ import annotations

import re
import sys
import zipfile
from pathlib import Path

from docx import Document


COUNT_RE = re.compile(r"공백 포함 (\d+)자$")
NUMBER_RE = re.compile(r"\d+(?:[.,]\d+)?(?:주일|년|월|일|개|명|%|원)?")


def body_paragraphs(document: Document) -> list[str]:
    return [paragraph.text for paragraph in document.paragraphs if paragraph.text][1::2]


def main() -> None:
    original_path = Path(sys.argv[1])
    edited_path = Path(sys.argv[2])

    with zipfile.ZipFile(edited_path) as archive:
        bad_member = archive.testzip()
        if bad_member:
            raise RuntimeError(f"손상된 ZIP 항목: {bad_member}")
        required = {"[Content_Types].xml", "word/document.xml", "word/styles.xml"}
        missing = required.difference(archive.namelist())
        if missing:
            raise RuntimeError(f"필수 OOXML 항목 누락: {sorted(missing)}")

    original = Document(original_path)
    edited = Document(edited_path)
    original_nonempty = [p for p in original.paragraphs if p.text]
    edited_nonempty = [p for p in edited.paragraphs if p.text]
    if len(original_nonempty) != len(edited_nonempty) or len(edited_nonempty) != 10:
        raise RuntimeError("비어 있지 않은 문단 수가 보존되지 않았습니다.")

    if len(original.sections) != len(edited.sections):
        raise RuntimeError("섹션 수가 달라졌습니다.")
    if len(original.tables) != len(edited.tables):
        raise RuntimeError("표 수가 달라졌습니다.")

    for index, (before, after) in enumerate(zip(original_nonempty, edited_nonempty)):
        before_style = before.style.name if before.style else None
        after_style = after.style.name if after.style else None
        if before_style != after_style:
            raise RuntimeError(f"문단 {index} 스타일이 달라졌습니다.")

    edited_bodies = body_paragraphs(edited)
    original_bodies = body_paragraphs(original)
    for section_index, body in enumerate(edited_bodies):
        heading = edited_nonempty[section_index * 2].text
        match = COUNT_RE.search(heading)
        if not match or int(match.group(1)) != len(body):
            raise RuntimeError(f"문항 {section_index + 1}의 글자 수 표기가 맞지 않습니다.")
        if len(body) > 500:
            raise RuntimeError(f"문항 {section_index + 1}이 500자를 초과했습니다.")

    original_numbers = NUMBER_RE.findall("\n".join(original_bodies))
    edited_numbers = NUMBER_RE.findall("\n".join(edited_bodies))
    if original_numbers != edited_numbers:
        raise RuntimeError(
            f"본문 수치가 달라졌습니다: original={original_numbers}, edited={edited_numbers}"
        )

    required_terms = ["농협", "기초연금", "엑셀", "의료인력", "전통시장"]
    edited_text = "\n".join(edited_bodies)
    for term in required_terms:
        if term not in edited_text:
            raise RuntimeError(f"필수 고유 용어가 누락됐습니다: {term}")

    print("zip_integrity=ok")
    print(f"paragraphs={len(edited_nonempty)} sections={len(edited.sections)} tables={len(edited.tables)}")
    print("section_chars=" + ",".join(str(len(body)) for body in edited_bodies))
    print("numbers_preserved=ok")
    print("styles_preserved=ok")


if __name__ == "__main__":
    main()
