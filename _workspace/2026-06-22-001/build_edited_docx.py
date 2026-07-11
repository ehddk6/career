from __future__ import annotations

import sys
from pathlib import Path

from docx import Document


def replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def main() -> None:
    source = Path(sys.argv[1])
    final_text = Path(sys.argv[2])
    output = Path(sys.argv[3])

    lines = final_text.read_text(encoding="utf-8").splitlines()
    nonempty = [line for line in lines if line.strip()]
    if len(nonempty) != 10:
        raise RuntimeError(f"예상한 비어 있지 않은 문단 10개가 아니라 {len(nonempty)}개입니다.")

    document = Document(source)
    paragraphs = [paragraph for paragraph in document.paragraphs if paragraph.text]
    if len(paragraphs) != 10:
        raise RuntimeError(f"원본의 비어 있지 않은 문단이 10개가 아니라 {len(paragraphs)}개입니다.")

    for section_index in range(5):
        heading = nonempty[section_index * 2]
        body = nonempty[section_index * 2 + 1]
        heading_prefix = heading.split("—", 1)[0].rstrip()
        counted_heading = f"{heading_prefix} — 공백 포함 {len(body)}자"
        replace_paragraph_text(paragraphs[section_index * 2], counted_heading)
        replace_paragraph_text(paragraphs[section_index * 2 + 1], body)

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    print(f"saved={output}")
    for section_index in range(5):
        body = nonempty[section_index * 2 + 1]
        print(f"section={section_index + 1} chars={len(body)}")


if __name__ == "__main__":
    main()
