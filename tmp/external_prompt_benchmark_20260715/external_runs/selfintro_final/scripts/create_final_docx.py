from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "output" / "08_final_selfintro.json"
TARGET = ROOT / "output" / "08_final_selfintro.docx"


def set_run_font(run, name: str = "맑은 고딕", size: float = 10.5) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_margins(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_mm: list[float]) -> None:
    widths_dxa = [round(value / 25.4 * 1440) for value in widths_mm]
    total = sum(widths_dxa)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    for col, width_mm in zip(table.columns, widths_mm):
        col.width = Mm(width_mm)
    for row in table.rows:
        for cell, width_dxa in zip(row.cells, widths_dxa):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(width_dxa))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])
    set_run_font(run, size=8)
    run.font.color.rgb = RGBColor(117, 117, 117)


def add_answer(document: Document, text: str) -> None:
    for block in text.replace("\r\n", "\n").replace("\r", "\n").split("\n\n"):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(8)
        paragraph.paragraph_format.line_spacing = 1.333
        paragraph.paragraph_format.first_line_indent = Mm(4)
        paragraph.paragraph_format.widow_control = True
        run = paragraph.add_run(block.strip())
        set_run_font(run, size=10.3)
        run.font.color.rgb = RGBColor(32, 32, 32)


def build() -> None:
    payload = json.loads(SOURCE.read_text(encoding="utf-8"))
    document = Document()
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(17)
    section.bottom_margin = Mm(16)
    section.left_margin = Mm(19)
    section.right_margin = Mm(19)
    section.header_distance = Mm(8)
    section.footer_distance = Mm(8)

    normal = document.styles["Normal"]
    normal.font.name = "맑은 고딕"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.font.size = Pt(10.3)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("신용보증기금 체험형 청년인턴1(보증)")
    set_run_font(run, size=16)
    run.bold = True
    run.font.color.rgb = RGBColor(20, 62, 104)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(10)
    run = subtitle.add_run("자기소개서 · 최종 제출본")
    set_run_font(run, size=10)
    run.font.color.rgb = RGBColor(90, 90, 90)

    summary = document.add_table(rows=2, cols=4)
    summary.autofit = False
    set_table_geometry(summary, [43, 43, 43, 43])
    labels = ["문항 1", "문항 2", "문항 3", "문항 4"]
    for index, label in enumerate(labels):
        cell = summary.cell(0, index)
        set_cell_shading(cell, "DCE6F1")
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(label)
        set_run_font(run, size=8.5)
        run.bold = True
        count_cell = summary.cell(1, index)
        paragraph = count_cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(f"{payload['counts'][f'q{index + 1}']:,}자")
        set_run_font(run, size=8.5)
    document.add_paragraph().paragraph_format.space_after = Pt(0)

    for number in range(1, 5):
        question = payload["questions"][f"q{number}"]
        answer = payload["answers"][f"q{number}"]

        heading = document.add_paragraph()
        heading.paragraph_format.keep_with_next = True
        heading.paragraph_format.space_before = Pt(8 if number == 1 else 12)
        heading.paragraph_format.space_after = Pt(6)
        run = heading.add_run(f"문항 {number}")
        set_run_font(run, size=13)
        run.bold = True
        run.font.color.rgb = RGBColor(46, 116, 181)

        prompt = document.add_paragraph()
        prompt.paragraph_format.keep_with_next = True
        prompt.paragraph_format.space_after = Pt(5)
        prompt.paragraph_format.line_spacing = 1.15
        run = prompt.add_run(question)
        set_run_font(run, size=9)
        run.bold = True
        run.font.color.rgb = RGBColor(85, 85, 85)

        add_answer(document, answer)

        count = document.add_paragraph()
        count.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        count.paragraph_format.space_after = Pt(2)
        run = count.add_run(f"공백 포함 · 줄바꿈 제외 {payload['counts'][f'q{number}']:,}자")
        set_run_font(run, size=8)
        run.font.color.rgb = RGBColor(110, 110, 110)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)

    properties = document.core_properties
    properties.title = "신용보증기금 체험형 청년인턴1(보증) 자기소개서"
    properties.subject = "최종 제출본"
    properties.author = "지원자"
    properties.keywords = "신용보증기금, 청년인턴, 보증, 자기소개서"

    TARGET.parent.mkdir(parents=True, exist_ok=True)
    document.save(TARGET)


if __name__ == "__main__":
    build()
