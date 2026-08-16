from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


QUESTIONS = [
    "신용보증기금 청년인턴에 지원하게 된 동기와 신용보증기금의 역할에 대한 본인의 이해를 바탕으로 인턴 근무 중 배우고, 기여하고 싶은 부분을 기술하여 주십시오.",
    "지원자가 새로운 조직에 적응하기 위해 중요하게 생각하는 태도가 무엇이며, 그 태도를 실제 근무과정에서 실천하기 위해 어떤 노력을 할 것인지 기술하여 주십시오.",
    "신용보증기금의 청년인턴으로 근무한다고 가정할 때, 실제 근무 시 지원자의 업무수행계획을 기술하여 주십시오.",
    "최근 중소기업에 큰 영향을 미치는 경제·사회 이슈를 하나 선택하여 그 이유를 서술하여 주시기 바랍니다. 또한, 해당 이슈와 관련하여 영향을 받는 중소기업을 정책 금융기관이 지원할 수 있는 방안과 지원 과정에서 유의할 점을 서술하여 주시기 바랍니다.",
]

LIMITS = [600, 600, 600, 1500]


def set_run_font(run, name: str, size: float, *, bold: bool = False, color: str | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_end])
    set_run_font(run, "맑은 고딕", 8.5, color="777777")


def parse_answers(markdown: str) -> list[str]:
    parts = re.split(r"^##\s+\d+\..*$", markdown, flags=re.MULTILINE)[1:]
    answers = [part.strip() for part in parts]
    if len(answers) != 4:
        raise ValueError(f"문항 4개가 필요하지만 {len(answers)}개를 찾았습니다.")
    return answers


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)
    section.different_first_page_header_footer = False

    normal = doc.styles["Normal"]
    normal.font.name = "맑은 고딕"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.55
    normal.paragraph_format.space_after = Pt(0)

    for sec in doc.sections:
        add_page_number(sec.footer.paragraphs[0])


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    p.paragraph_format.keep_with_next = True
    run = p.add_run("자기소개서")
    set_run_font(run, "맑은 고딕", 24, bold=True, color="1F4E79")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    p.paragraph_format.keep_with_next = True
    run = p.add_run("신용보증기금 체험형 청년인턴1(보증)")
    set_run_font(run, "맑은 고딕", 10.5, color="666666")


def add_question(doc: Document, index: int, answer: str) -> None:
    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(10 if index != 1 else 0)
    heading.paragraph_format.space_after = Pt(7)
    heading.paragraph_format.keep_with_next = True
    heading.paragraph_format.line_spacing = 1.25
    run = heading.add_run(f"{index}. {QUESTIONS[index - 1]}")
    set_run_font(run, "맑은 고딕", 13.5, bold=True, color="2E75B6")

    count = len(answer)
    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(7)
    meta.paragraph_format.keep_with_next = True
    run = meta.add_run(f"제한: {LIMITS[index - 1]:,}자 (공백 포함) · 현재: {count:,}자")
    set_run_font(run, "맑은 고딕", 8.5, color="666666")

    for paragraph_text in answer.split("\n\n"):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.52
        p.paragraph_format.space_after = Pt(6 if index == 4 else 0)
        p.paragraph_format.widow_control = True
        run = p.add_run(paragraph_text.replace("\n", " "))
        set_run_font(run, "맑은 고딕", 10.5)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input_md", type=Path)
    parser.add_argument("output_docx", type=Path)
    args = parser.parse_args()

    markdown = args.input_md.read_text(encoding="utf-8")
    answers = parse_answers(markdown)
    doc = Document()
    configure_document(doc)
    add_title(doc)
    for index, answer in enumerate(answers, start=1):
        add_question(doc, index, answer)

    doc.core_properties.title = "신용보증기금 체험형 청년인턴1(보증) 자기소개서"
    doc.core_properties.subject = "자기소개서"
    doc.core_properties.author = "지원자"
    args.output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.output_docx)


if __name__ == "__main__":
    main()
