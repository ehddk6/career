from docx import Document
from docx.shared import Inches, Pt, RGBColor

from career_pipeline.models import DraftResponse, Question
from career_pipeline.rendering import render_draft_docx, render_draft_markdown


def test_markdown_and_docx_contain_all_questions(tmp_path):
    questions = [
        Question(1, "지원동기", 600),
        Question(2, "문제해결", 600),
    ]
    responses = [
        DraftResponse(1, "지원 답변", ("경험정리/a.docx",)),
        DraftResponse(2, "개선 답변", ("경험정리/b.docx",)),
    ]

    markdown = render_draft_markdown(questions, responses)
    output = tmp_path / "draft.docx"
    render_draft_docx(questions, responses, output)

    assert "## 1. 지원동기" in markdown
    assert "개선 답변" in markdown
    text = "\n".join(paragraph.text for paragraph in Document(output).paragraphs)
    assert "지원 답변" in text
    assert "개선 답변" in text


def test_docx_uses_standard_business_brief_geometry_and_styles(tmp_path):
    output = tmp_path / "draft.docx"
    render_draft_docx(
        [Question(1, "지원동기", 600)],
        [DraftResponse(1, "지원 답변", ("a.docx",))],
        output,
    )

    document = Document(output)
    section = document.sections[0]
    normal = document.styles["Normal"]
    heading = document.styles["Heading 1"]

    assert section.page_width == Inches(8.5)
    assert section.page_height == Inches(11)
    assert section.top_margin == Inches(1)
    assert section.right_margin == Inches(1)
    assert section.bottom_margin == Inches(1)
    assert section.left_margin == Inches(1)
    assert normal.font.name == "Calibri"
    assert normal.font.size == Pt(11)
    assert normal.paragraph_format.space_after == Pt(6)
    assert normal.paragraph_format.line_spacing == 1.1
    assert heading.font.size == Pt(16)
    assert heading.font.color.rgb == RGBColor(0x2E, 0x74, 0xB5)
    assert heading.paragraph_format.space_before == Pt(16)
    assert heading.paragraph_format.space_after == Pt(8)
