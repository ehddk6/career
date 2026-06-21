from pathlib import Path

from docx import Document
from openpyxl import Workbook

from career_pipeline.extractors import extract_path, split_text_blocks
from career_pipeline.models import SourceRecord


def record(path: Path) -> SourceRecord:
    return SourceRecord(
        path,
        path.name,
        path.suffix.lower(),
        path.stat().st_size,
        "hash",
        "use",
    )


def test_extracts_docx_paragraphs_and_tables(tmp_path: Path):
    path = tmp_path / "draft.docx"
    document = Document()
    document.add_paragraph("첫 번째 문장입니다. 두 번째 문장입니다.")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "항목"
    table.cell(0, 1).text = "답변 내용"
    document.save(path)

    result = extract_path(record(path))

    assert result.paragraphs == (
        "첫 번째 문장입니다.",
        "두 번째 문장입니다.",
        "[표 1 R1C1] 항목",
        "[표 1 R1C2] 답변 내용",
    )


def test_extracts_xlsx_cells_and_sheet_names(tmp_path: Path):
    path = tmp_path / "interview.xlsx"
    workbook = Workbook()
    workbook.active.title = "가이드"
    workbook.active["B2"] = "면접 전략"
    workbook.active["C4"] = "지원동기"
    workbook.save(path)

    result = extract_path(record(path))

    assert "[시트] 가이드" in result.text
    assert "면접 전략" in result.text
    assert "지원동기" in result.text
    assert "[가이드!B2] 면접 전략" in result.paragraphs
    assert "[가이드!C4] 지원동기" in result.paragraphs


def test_extracts_utf8_bom_text(tmp_path: Path):
    path = tmp_path / "notes.txt"
    path.write_text("첫 경험입니다. 둘째 근거입니다.", encoding="utf-8-sig")
    assert extract_path(record(path)).paragraphs == (
        "첫 경험입니다.",
        "둘째 근거입니다.",
    )


def test_rejects_unsupported_file(tmp_path: Path):
    path = tmp_path / "image.jpg"
    path.write_bytes(b"jpg")

    try:
        extract_path(record(path))
    except ValueError as error:
        assert "unsupported extension" in str(error)
    else:
        raise AssertionError("unsupported file was accepted")


def test_splits_large_extracted_pages_into_logical_blocks():
    text = "첫 번째 문장입니다.\n두 번째 문장입니다. 세 번째 문장입니다."

    assert split_text_blocks(text) == [
        "첫 번째 문장입니다.",
        "두 번째 문장입니다.",
        "세 번째 문장입니다.",
    ]
