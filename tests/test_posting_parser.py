from hashlib import sha256
from pathlib import Path

from docx import Document

from career_pipeline.models import Question
from career_pipeline.posting_loader import load_posting_source
from career_pipeline.posting_parser import parse_posting, reconcile_questions
from career_pipeline.posting_schema import LoadedPosting, PostingSourceMetadata
from career_pipeline.questions import extract_questions


FIXTURES = Path(__file__).parent / "fixtures"


def loaded_html_fixture(name: str) -> LoadedPosting:
    content = (FIXTURES / name).read_bytes()
    return LoadedPosting(
        PostingSourceMetadata(
            kind="url",
            location="https://khug.or.kr/posting",
            retrieved_at="2026-06-21T12:00:00+09:00",
            content_sha256=sha256(content).hexdigest(),
            official_status="verified_domain",
            content_type="text/html; charset=utf-8",
        ),
        ".html",
        content,
    )


def test_parse_posting_extracts_role_duty_requirement_and_question():
    loaded = loaded_html_fixture("hug_posting_excerpt.html")

    analysis = parse_posting(loaded, target="HUG 금융·기금(강원)")

    assert analysis.role == "금융·기금(강원)"
    assert analysis.duties == ("도시재생 금융지원 관련 안내 등 업무 보조",)
    assert analysis.requirements == ("공고일 기준 지원자격을 충족한 자",)
    assert analysis.questions[0].character_limit == 600
    assert "organization" in analysis.uncertainties


def test_parse_posting_extracts_docx_blocks(tmp_path: Path):
    path = tmp_path / "posting.docx"
    document = Document()
    for text in ("채용분야", "행정지원", "담당업무", "민원 안내"):
        document.add_paragraph(text)
    document.save(path)

    analysis = parse_posting(
        load_posting_source(path, official_source=True), target="기관 행정지원"
    )

    assert analysis.role == "행정지원"
    assert analysis.duties == ("민원 안내",)


def test_parse_posting_records_missing_required_fields_and_unclassified_text():
    content = "<h1>공고</h1><h2>채용분야</h2><p>행정</p><p>기타 안내</p>".encode()
    loaded = LoadedPosting(
        PostingSourceMetadata(
            "url",
            "https://example.or.kr/posting",
            "2026-06-21T12:00:00+09:00",
            sha256(content).hexdigest(),
            "verified_domain",
            "text/html",
        ),
        ".html",
        content,
    )

    analysis = parse_posting(loaded, target="기관 행정")

    assert "organization" in analysis.uncertainties
    assert "duties" in analysis.uncertainties
    assert any("기타 안내" in item for item in analysis.uncertainties)


def test_reconcile_questions_reports_prompt_and_limit_mismatch():
    posting = (Question(1, "지원동기를 기술해 주십시오.", 600),)
    matching = (Question(1, " 지원동기를 기술해 주십시오 ", 600),)
    mismatch = (Question(1, "입사 후 목표를 기술해 주십시오.", 500),)

    assert reconcile_questions(posting, matching).matched

    result = reconcile_questions(posting, mismatch)
    assert not result.matched
    assert {row.reason for row in result.mismatches} == {
        "prompt_mismatch",
        "character_limit_mismatch",
    }


def test_reconcile_questions_uses_nonempty_side():
    questions = (Question(1, "지원동기를 기술해 주십시오.", 600),)

    assert reconcile_questions(questions, ()).questions == questions
    assert reconcile_questions((), questions).questions == questions


def test_extracts_numbered_limited_essay_that_ends_with_request_wording():
    questions = extract_questions(
        ("4. 최근 경제 이슈와 정책금융 지원 방안을 서술하여 주시기 바랍니다. (1500자 이내)",)
    )

    assert questions[0].character_limit == 1500


def test_extracts_upper_bound_and_joins_multi_sentence_prompt():
    questions = extract_questions((
        "4. 최근 경제 이슈를 하나 선택하여 이유를 서술하여 주시기 바랍니다.",
        "또한 정책금융 지원 방안과 유의점을 서술하여 주시기 바랍니다. (400자 이상 600자 이내)",
    ))

    assert questions[0].character_limit == 600
    assert "또한 정책금융" in questions[0].prompt
    assert not questions[0].prompt.startswith("4.")
    assert "()" not in questions[0].prompt

def test_parse_posting_falls_back_when_no_section_labels(tmp_path: Path):
    path = tmp_path / 'posting.docx'
    document = Document()
    for line in ('주택도시보증공사', '금융기금(강원)', '주택청약 접수 창구 관리', '정확성 소통력'):
        document.add_paragraph(line)
    document.save(path)

    analysis = parse_posting(
        load_posting_source(path, official_source=True), target='HUG'
    )

    assert analysis.organization == '주택도시보증공사'
    assert analysis.role == '금융기금(강원)'
    assert '주택청약' in analysis.duties[0]


def test_parse_posting_extracts_inline_labels_from_common_korean_format():
    content = """
    기관명: 한국주택금융공사
    채용분야: 행정직
    담당업무: 보증심사 자료 검토 및 민원 안내
    필요역량: 정확성, 고객 소통
    자기소개서: 지원동기를 600자 이내로 기술해 주십시오.
    """.encode("utf-8")
    loaded = LoadedPosting(
        PostingSourceMetadata(
            "url",
            "https://hf.go.kr/posting",
            "2026-06-21T12:00:00+09:00",
            sha256(content).hexdigest(),
            "verified_domain",
            "text/plain",
        ),
        ".txt",
        content,
    )

    analysis = parse_posting(loaded, target="한국주택금융공사 행정")

    assert analysis.organization == "한국주택금융공사"
    assert analysis.role == "행정직"
    assert analysis.duties == ("보증심사 자료 검토 및 민원 안내",)
    assert analysis.competencies == ("정확성, 고객 소통",)
    assert analysis.questions[0].character_limit == 600


def test_parse_posting_uses_target_when_public_page_omits_organization_and_has_table_header():
    content = """
    <h1>2026년도 하반기 체험형 청년인턴 채용</h1>
    <h2>채용분야</h2><p>근무기간</p><p>채용인원</p>
    <h2>주요업무</h2><p>신용보증 기한연장, 기업신용 상시관리</p>
    """.encode("utf-8")
    loaded = LoadedPosting(
        PostingSourceMetadata("url", "https://example.or.kr", "2026-07-11T09:00:00+09:00", sha256(content).hexdigest(), "verified_domain", "text/html"),
        ".html",
        content,
    )

    analysis = parse_posting(loaded, target="신용보증기금 체험형 청년인턴1(보증)")

    assert analysis.organization == "신용보증기금"
    assert analysis.role == "체험형 청년인턴1(보증)"
